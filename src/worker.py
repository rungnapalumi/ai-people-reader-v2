import os
import io
import json
import gc
import threading
import resource
import uuid
from datetime import datetime, timezone

from dotenv import load_dotenv
load_dotenv()
import time
import math
import tempfile
import logging
import shutil
import subprocess
import smtplib
import re
from dataclasses import dataclass
from typing import Dict, Any, Optional, List, Tuple
from email.mime.text import MIMEText

import boto3

# Heavy libs
import cv2
import numpy as np


def _configure_mediapipe_headless_env() -> None:
    if str(os.getenv("MEDIAPIPE_USE_GPU", "")).strip().lower() in ("1", "true", "yes", "on"):
        return
    os.environ.setdefault("CUDA_VISIBLE_DEVICES", "-1")
    os.environ.setdefault("TF_CPP_MIN_LOG_LEVEL", "2")
    os.environ.setdefault("GLOG_minloglevel", "2")
    os.environ.setdefault("TF_FORCE_CPU", "1")
    os.environ.setdefault("TF_ENABLE_ONEDNN_OPTS", "0")


_configure_mediapipe_headless_env()

try:
    # Try new MediaPipe API first (0.10.8+)
    from mediapipe.python.solutions import pose as mp_pose_module
    from mediapipe.python.solutions.pose import Pose, PoseLandmark
except ImportError:
    # Fall back to old API
    import mediapipe as mp
    mp_pose_module = mp.solutions.pose
    Pose = mp_pose_module.Pose
    PoseLandmark = mp_pose_module.PoseLandmark

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)s | %(message)s")

AWS_BUCKET = os.getenv("AWS_BUCKET") or os.getenv("S3_BUCKET")
AWS_REGION = os.getenv("AWS_REGION", "ap-southeast-1")

if not AWS_BUCKET:
    raise RuntimeError("Missing AWS_BUCKET (or S3_BUCKET)")

s3 = boto3.client("s3", region_name=AWS_REGION)
SES_REGION = os.getenv("SES_REGION", AWS_REGION).strip() or AWS_REGION
SES_FROM_EMAIL = (os.getenv("SES_FROM_EMAIL") or "").strip()
SMTP_HOST = (os.getenv("SMTP_HOST") or "").strip()
SMTP_PORT = int(os.getenv("SMTP_PORT", "465"))
SMTP_USERNAME = (os.getenv("SMTP_USERNAME") or "").strip()
SMTP_PASSWORD = (os.getenv("SMTP_PASSWORD") or "").strip()
SMTP_USE_TLS = str(os.getenv("SMTP_USE_TLS", "false")).strip().lower() in ("1", "true", "yes", "y")
SMTP_USE_SSL = str(os.getenv("SMTP_USE_SSL", "true")).strip().lower() in ("1", "true", "yes", "y")
SMTP_FROM_EMAIL = (os.getenv("SMTP_FROM_EMAIL") or SES_FROM_EMAIL).strip()

PENDING = "jobs/pending/"
PROCESSING = "jobs/processing/"
FINISHED = "jobs/finished/"
FAILED = "jobs/failed/"
PROCESSING_STALE_MINUTES = int(os.getenv("PROCESSING_STALE_MINUTES", "20"))
PROCESSING_RECOVERY_MAX_ITEMS = int(os.getenv("PROCESSING_RECOVERY_MAX_ITEMS", "50"))
PROCESSING_RECOVERY_INTERVAL_SECONDS = int(os.getenv("PROCESSING_RECOVERY_INTERVAL_SECONDS", "60"))
IDLE_HEARTBEAT_SECONDS = int(os.getenv("IDLE_HEARTBEAT_SECONDS", "60"))
MAX_VIDEO_JOB_RETRIES = int(os.getenv("MAX_VIDEO_JOB_RETRIES", "2"))
# Unique per-process identifier stamped into every processing JSON we claim. Used by
# orphan detection to tell "my live claim" from "a dead worker left this behind".
WORKER_ID = f"{int(time.time())}-{os.getpid()}-{uuid.uuid4().hex[:8]}"
# On startup the video worker is the sole claimer of processing/*.json (render.yaml: single
# instance). Any JSON already sitting in processing/ must therefore belong to a crashed
# predecessor — reclaim unconditionally instead of waiting PROCESSING_STALE_MINUTES, because
# the predecessor's heartbeat thread may have refreshed LastModified seconds before it died.
FORCE_RECOVER_ON_STARTUP = str(os.getenv("FORCE_RECOVER_ON_STARTUP", "true")).strip().lower() in ("1", "true", "yes", "on")
# Secondary orphan sweep: when the worker is idle (no claimed job in this process) but S3
# still shows processing>0, read each processing JSON and reclaim any whose worker_id does
# not match this process's WORKER_ID. Cheap (only runs while idle) and complementary to the
# age-based stale check.
ORPHAN_SWEEP_WHEN_IDLE = str(os.getenv("ORPHAN_SWEEP_WHEN_IDLE", "true")).strip().lower() in ("1", "true", "yes", "on")
# S3 lists pending in lexicographic order; many report jobs can fill the first N keys and hide dots/skeleton.
REPORT_JOB_MODES = frozenset({"report", "report_th_en", "report_generator"})
# Max *.json keys to *read* under jobs/pending/ per list_pending() call while searching for
# dots/skeleton. Report jobs are skipped for *selection* but each still requires one S3 read.
# Older bug: a too-small cap stopped the scan after N reads even when all N were mode=report,
# so dots/skeleton keys that sorted later were never seen (skeleton "never runs").
_PENDING_READS_ENV = os.getenv("WORKER_PENDING_MAX_JSON_READS") or os.getenv("WORKER_PENDING_MAX_SCAN") or "200000"
WORKER_PENDING_MAX_JSON_READS = max(2000, int(_PENDING_READS_ENV or "200000"))
# Back-compat alias for logs / env docs
WORKER_PENDING_MAX_SCAN = WORKER_PENDING_MAX_JSON_READS
# Modes this worker is allowed to claim and recover. Default keeps single-instance behaviour
# (one worker handles every video mode). Set to "dots" or "skeleton" on each Render service when
# splitting the video worker into two parallel instances — that way each worker only touches its
# own jobs and recovery sweeps don't yank the other worker's in-flight processing/<id>.json.
def _parse_handles_modes(value: str) -> frozenset:
    items = {seg.strip().lower() for seg in (value or "").split(",") if seg.strip()}
    return frozenset(items) if items else frozenset({"dots", "skeleton"})


WORKER_HANDLES_MODES = _parse_handles_modes(os.getenv("WORKER_HANDLES_MODES") or "dots,skeleton")
WORKER_HANDLES_ALL_VIDEO_MODES = WORKER_HANDLES_MODES == frozenset({"dots", "skeleton"})


def _job_mode_handled_here(mode: str) -> bool:
    """True if a pending/processing JSON's mode belongs to this worker. Always False for report modes."""
    m = (mode or "").strip().lower()
    if not m:
        # Unknown mode — only the back-compat "handles all video modes" worker should sweep it.
        return WORKER_HANDLES_ALL_VIDEO_MODES
    if m in REPORT_JOB_MODES:
        return False
    return m in WORKER_HANDLES_MODES

# Skeleton/dots ready emails: also notify these addresses (not shown in Streamlit). Set to "" to disable.
_DEFAULT_INTERNAL_ANALYSIS_NOTIFY_EMAILS = "rungnapa@imagematters.at,petchpat@gmail.com"


def get_internal_analysis_notify_recipients() -> List[str]:
    raw = os.environ.get("INTERNAL_ANALYSIS_NOTIFY_EMAILS")
    if raw is None:
        raw = _DEFAULT_INTERNAL_ANALYSIS_NOTIFY_EMAILS
    else:
        raw = str(raw).strip()
    if not raw:
        return []
    return parse_email_list(raw)


def merge_notification_recipients(primary: List[str]) -> List[str]:
    combined = list(primary) + get_internal_analysis_notify_recipients()
    seen: set = set()
    out: List[str] = []
    for email in combined:
        key = email.lower()
        if key in seen:
            continue
        if not is_valid_email(email):
            continue
        seen.add(key)
        out.append(email)
    return out


# -----------------------------
# S3 helpers
# -----------------------------
EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")


def s3_read_json(key: str) -> Dict[str, Any]:
    obj = s3.get_object(Bucket=AWS_BUCKET, Key=key)
    raw = obj["Body"].read().decode("utf-8")
    return json.loads(raw)


def s3_head_exists(key: str) -> bool:
    """True if object exists (HeadObject). Used to serialize skeleton after dots when multiple video replicas run."""
    k = str(key or "").strip()
    if not k:
        return False
    try:
        s3.head_object(Bucket=AWS_BUCKET, Key=k)
        return True
    except Exception:
        return False


def wait_for_dots_mp4_before_skeleton(job: Dict[str, Any], group_id: str) -> None:
    """
    When Render (or any host) runs **more than one** video-worker replica, pending dots + skeleton
    jobs can be claimed in parallel — skeleton may finish before dots.mp4 exists. People Reader expects
    dots before skeleton in UX; bundle email also expects both files. Poll S3 for the canonical dots key.
    """
    if str(os.getenv("SKELETON_WAIT_FOR_DOTS", "true")).strip().lower() not in ("1", "true", "yes", "on"):
        return
    if bool(job.get("skip_wait_for_dots")):
        return
    if job.get("require_dots_output") is False:
        return
    # Only People Reader (and explicit callers) set this — legacy skeleton jobs omit it and must not block here.
    if job.get("require_dots_output") is not True:
        return
    gid = str(group_id or "").strip()
    if not gid:
        return
    dots_key = str(job.get("dots_output_wait_key") or "").strip() or f"jobs/output/groups/{gid}/dots.mp4"
    timeout_s = max(60.0, float(os.getenv("SKELETON_WAIT_FOR_DOTS_TIMEOUT_SEC", "300") or "300"))
    poll_s = max(2.0, float(os.getenv("SKELETON_WAIT_FOR_DOTS_POLL_SEC", "5") or "5"))
    log_every = max(10.0, float(os.getenv("SKELETON_WAIT_FOR_DOTS_LOG_EVERY_SEC", "30") or "30"))
    deadline = time.time() + timeout_s
    last_log = 0.0
    if s3_head_exists(dots_key):
        return
    logging.info(
        "[skeleton] waiting for dots.mp4 before encoding (multi-replica safe) group_id=%s key=%s",
        gid,
        dots_key,
    )
    while time.time() < deadline:
        if s3_head_exists(dots_key):
            logging.info("[skeleton] dots ready; starting skeleton encode group_id=%s", gid)
            return
        now = time.time()
        if now - last_log >= log_every:
            logging.info("[skeleton] still waiting for dots.mp4 group_id=%s key=%s", gid, dots_key)
            last_log = now
        time.sleep(poll_s)
    raise RuntimeError(
        f"[skeleton] timed out waiting for dots.mp4 ({int(timeout_s)}s) — failing job "
        f"group_id={gid} key={dots_key}. Dots job may have failed or is still processing."
    )


def s3_write_json(key: str, payload: Dict[str, Any]) -> None:
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    s3.put_object(
        Bucket=AWS_BUCKET,
        Key=key,
        Body=body,
        ContentType="application/json; charset=utf-8",
    )


def s3_download_to_file(key: str, path: str) -> None:
    s3.download_file(AWS_BUCKET, key, path)


def s3_upload_file(path: str, key: str, content_type: str) -> None:
    s3.upload_file(path, AWS_BUCKET, key, ExtraArgs={"ContentType": content_type})


def s3_put_bytes(key: str, data: bytes, content_type: str) -> None:
    s3.put_object(Bucket=AWS_BUCKET, Key=key, Body=data, ContentType=content_type)


def presigned_get_url(key: str, expires: int = 86400) -> str:
    return s3.generate_presigned_url(
        ClientMethod="get_object",
        Params={"Bucket": AWS_BUCKET, "Key": key},
        ExpiresIn=max(300, int(expires)),
    )


def is_valid_email(value: str) -> bool:
    return bool(EMAIL_RE.match(str(value or "").strip()))


def parse_email_list(value: str) -> List[str]:
    seen: set = set()
    out: List[str] = []
    for token in re.split(r"[,\s;]+", str(value or "").strip()):
        email = token.strip()
        if not email:
            continue
        key = email.lower()
        if key in seen:
            continue
        if is_valid_email(email):
            seen.add(key)
            out.append(email)
    return out


def _send_email_via_ses(to_email: str, subject: str, body: str) -> Tuple[bool, str]:
    if not SES_FROM_EMAIL:
        return False, "ses_from_not_configured"
    try:
        ses = boto3.client("ses", region_name=SES_REGION)
        ses.send_email(
            Source=SES_FROM_EMAIL,
            Destination={"ToAddresses": [to_email]},
            Message={
                "Subject": {"Data": subject, "Charset": "UTF-8"},
                "Body": {"Text": {"Data": body, "Charset": "UTF-8"}},
            },
        )
        return True, "sent_via_ses"
    except Exception as e:
        return False, f"ses_failed:{e}"


def _send_email_via_smtp(to_email: str, subject: str, body: str) -> Tuple[bool, str]:
    if not (SMTP_HOST and SMTP_USERNAME and SMTP_PASSWORD and SMTP_FROM_EMAIL):
        return False, "smtp_not_configured"
    try:
        msg = MIMEText(body, _charset="utf-8")
        msg["Subject"] = subject
        msg["From"] = SMTP_FROM_EMAIL
        msg["To"] = to_email
        if SMTP_USE_SSL:
            with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT, timeout=20) as server:
                server.login(SMTP_USERNAME, SMTP_PASSWORD)
                server.sendmail(SMTP_FROM_EMAIL, [to_email], msg.as_string())
        else:
            with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=20) as server:
                if SMTP_USE_TLS:
                    server.starttls()
                server.login(SMTP_USERNAME, SMTP_PASSWORD)
                server.sendmail(SMTP_FROM_EMAIL, [to_email], msg.as_string())
        return True, "sent_via_smtp"
    except Exception as e:
        return False, f"smtp_failed:{e}"


def send_mode_ready_email(job: Dict[str, Any], result: Dict[str, Any]) -> Tuple[bool, str]:
    mode = str(result.get("mode") or "").strip().lower()
    if mode not in ("dots", "skeleton"):
        return False, "skip_non_video_mode"
    # Report worker sends skeleton link with PDFs when both are queued (TTB/LPA/etc.); avoid duplicate mail.
    if bool(job.get("suppress_completion_email") or job.get("skip_mode_ready_email")):
        return False, "skip_suppress_completion_email"
    recipients = merge_notification_recipients(
        parse_email_list(str(job.get("notify_email") or job.get("employee_email") or ""))
    )
    if not recipients:
        return False, "skip_no_valid_recipients"
    output_key = str(result.get("output_key") or "").strip()
    if not output_key:
        return False, "skip_missing_output_key"
    try:
        url = presigned_get_url(output_key, expires=86400)
    except Exception as e:
        return False, f"skip_presign_failed:{e}"

    group_id = str(job.get("group_id") or "").strip()
    job_id = str(job.get("job_id") or "").strip()
    mode_title = "Dots Video" if mode == "dots" else "Skeleton Video"
    subject = f"AI People Reader - {mode_title} Ready ({group_id or job_id})"
    body = (
        f"Your {mode_title.lower()} is ready.\n\n"
        f"Group ID: {group_id}\n"
        f"Job ID: {job_id}\n"
        f"Download link (valid for 24 hours):\n{url}\n"
    )

    statuses: List[str] = []
    sent_any = False
    for to_email in recipients:
        sent, status = _send_email_via_ses(to_email, subject, body)
        if not sent:
            sent, status = _send_email_via_smtp(to_email, subject, body)
        statuses.append(f"{to_email}:{status}")
        sent_any = sent_any or sent
    return sent_any, " | ".join(statuses)


def list_pending(limit: int = 200, max_scan: Optional[int] = None) -> List[str]:
    """
    List pending keys; dots first, then skeleton, then any other non-report video modes.

    Report jobs are skipped (report_worker). Skeleton is ordered before other video modes so
    People Reader bundle email is not blocked behind unrelated legacy jobs in the same batch.

    Important: S3 lists keys lexicographically. Many mode=report JSONs may appear *before*
    dots/skeleton for the same submission. We must keep reading past those reports until we
    find video jobs or hit WORKER_PENDING_MAX_JSON_READS (each key still costs one S3 GET).
    """
    hard_cap = max(max_scan if max_scan is not None else WORKER_PENDING_MAX_JSON_READS, 1)
    dots_keys: List[str] = []
    skeleton_keys: List[str] = []
    other_keys: List[str] = []
    total_reads = 0
    report_skips = 0
    paginator = s3.get_paginator("list_objects_v2")
    for page in paginator.paginate(Bucket=AWS_BUCKET, Prefix=PENDING):
        for item in page.get("Contents", []):
            key = str(item.get("Key") or "")
            if not key.endswith(".json"):
                continue
            total_reads += 1
            if total_reads > hard_cap:
                break
            try:
                job = s3_read_json(key)
                mode = str(job.get("mode") or "").strip().lower()
                if mode in REPORT_JOB_MODES:
                    report_skips += 1
                    continue
                # Skip modes that another worker instance is responsible for (e.g. when this
                # process is the dots-only worker and the key is mode=skeleton).
                if not _job_mode_handled_here(mode):
                    continue
                if mode == "dots":
                    dots_keys.append(key)
                elif mode == "skeleton":
                    skeleton_keys.append(key)
                else:
                    other_keys.append(key)
            except Exception:
                # On read failure default to handling the key only if this worker is the
                # back-compat "handles all" instance; otherwise skip so the specialised peer can claim it.
                if WORKER_HANDLES_ALL_VIDEO_MODES:
                    other_keys.append(key)
            picked = len(dots_keys) + len(skeleton_keys) + len(other_keys)
            if picked >= limit:
                return (dots_keys + skeleton_keys + other_keys)[:limit]
        if total_reads > hard_cap:
            break

    merged = (dots_keys + skeleton_keys + other_keys)[:limit]
    if not merged and total_reads >= hard_cap:
        logging.warning(
            "list_pending: read %s pending JSON keys (hard_cap=%s, report_skips=%s), found no dots/skeleton. "
            "Increase WORKER_PENDING_MAX_JSON_READS (or legacy WORKER_PENDING_MAX_SCAN) if jobs exist deeper in the queue.",
            total_reads,
            hard_cap,
            report_skips,
        )
    elif not dots_keys and (skeleton_keys or other_keys):
        logging.info(
            "list_pending: no dots in this batch (skeleton=%s other=%s of %s JSON reads, report_skips=%s)",
            len(skeleton_keys),
            len(other_keys),
            total_reads,
            report_skips,
        )
    return merged


def count_jobs(prefix: str) -> int:
    total = 0
    paginator = s3.get_paginator("list_objects_v2")
    for page in paginator.paginate(Bucket=AWS_BUCKET, Prefix=prefix):
        for item in page.get("Contents", []):
            key = str(item.get("Key") or "")
            if key.endswith(".json"):
                total += 1
    return total


def force_recover_all_processing_on_startup(max_items: int) -> int:
    """Unconditionally move every jobs/processing/*.json **of this worker's modes** back to pending.

    Originally this swept every processing JSON because the video worker was deployed as a
    single instance. With WORKER_HANDLES_MODES splitting dots and skeleton into separate
    services, this worker MUST only reclaim its own modes — otherwise the dots-worker boot
    would yank the skeleton-worker's in-flight processing/<id>.json (and vice versa), causing
    duplicate processing and duplicate emails.
    """
    max_items = max(1, int(max_items or 1))
    recovered = 0
    scanned = 0
    skipped_other_mode = 0
    paginator = s3.get_paginator("list_objects_v2")
    for page in paginator.paginate(Bucket=AWS_BUCKET, Prefix=PROCESSING):
        for item in page.get("Contents", []):
            key = str(item.get("Key") or "")
            if not key.endswith(".json"):
                continue
            scanned += 1
            # Only reclaim processing JSONs whose mode this worker is responsible for.
            # When WORKER_HANDLES_MODES is the default (both modes), this is a no-op cost wise
            # but adds one S3 read per processing key — acceptable tradeoff for safety.
            try:
                body = s3_read_json(key)
                mode = str(body.get("mode") or "").strip().lower()
            except Exception as e:
                logging.warning("[startup_force_recovery] could not read key=%s err=%s — skipping", key, e)
                continue
            if not _job_mode_handled_here(mode):
                skipped_other_mode += 1
                continue
            try:
                move_job(key, PENDING)
                recovered += 1
                logging.warning(
                    "[startup_force_recovery] reclaimed orphaned processing->pending key=%s mode=%s (worker_modes=%s)",
                    key,
                    mode or "<missing>",
                    sorted(WORKER_HANDLES_MODES),
                )
            except Exception as e:
                logging.warning("[startup_force_recovery] failed to move key=%s err=%s", key, e)
            if scanned >= max_items:
                if skipped_other_mode:
                    logging.info(
                        "[startup_force_recovery] skipped %s processing JSON(s) belonging to other worker modes",
                        skipped_other_mode,
                    )
                return recovered
    if skipped_other_mode:
        logging.info(
            "[startup_force_recovery] skipped %s processing JSON(s) belonging to other worker modes",
            skipped_other_mode,
        )
    return recovered


def recover_orphaned_processing_by_worker_id(current_worker_id: str, max_items: int) -> int:
    """Reclaim any processing JSON whose worker_id differs from this process's WORKER_ID.

    Run only while this worker holds no active claim (idle sweep). A JSON without a
    worker_id field is treated as legacy/orphan and reclaimed too. With mode-split workers,
    only modes this worker handles are eligible.
    """
    max_items = max(1, int(max_items or 1))
    recovered = 0
    scanned = 0
    paginator = s3.get_paginator("list_objects_v2")
    for page in paginator.paginate(Bucket=AWS_BUCKET, Prefix=PROCESSING):
        for item in page.get("Contents", []):
            key = str(item.get("Key") or "")
            if not key.endswith(".json"):
                continue
            scanned += 1
            try:
                body = s3_read_json(key)
            except Exception as e:
                logging.warning("[orphan_sweep] could not read key=%s err=%s", key, e)
                continue
            mode = str(body.get("mode") or "").strip().lower()
            if not _job_mode_handled_here(mode):
                continue
            owner = str(body.get("worker_id") or "").strip()
            if owner == current_worker_id:
                continue
            try:
                move_job(key, PENDING)
                recovered += 1
                logging.warning(
                    "[orphan_sweep] reclaimed key=%s mode=%s owner=%r (current worker_id=%s)",
                    key,
                    mode or "<missing>",
                    owner or "<missing>",
                    current_worker_id,
                )
            except Exception as e:
                logging.warning("[orphan_sweep] failed to move key=%s err=%s", key, e)
            if scanned >= max_items:
                return recovered
    return recovered


def recover_stale_processing_jobs(stale_minutes: int, max_items: int) -> int:
    """Reclaim processing JSONs older than `stale_minutes`, restricted to this worker's modes."""
    stale_minutes = max(1, int(stale_minutes or 1))
    max_items = max(1, int(max_items or 1))
    now_ts = time.time()
    recovered = 0
    scanned = 0
    paginator = s3.get_paginator("list_objects_v2")
    for page in paginator.paginate(Bucket=AWS_BUCKET, Prefix=PROCESSING):
        for item in page.get("Contents", []):
            key = str(item.get("Key") or "")
            if not key.endswith(".json"):
                continue
            scanned += 1
            last_modified = item.get("LastModified")
            if not last_modified:
                continue
            age_minutes = (now_ts - float(last_modified.timestamp())) / 60.0
            if age_minutes < float(stale_minutes):
                continue
            # Read mode before moving so we don't reclaim another worker's stale-but-still-its-own job.
            try:
                body = s3_read_json(key)
                mode = str(body.get("mode") or "").strip().lower()
            except Exception as e:
                logging.warning("[startup_recovery] could not read stale key=%s err=%s — skipping", key, e)
                continue
            if not _job_mode_handled_here(mode):
                continue
            try:
                move_job(key, PENDING)
                recovered += 1
                logging.warning(
                    "[startup_recovery] moved stale processing->pending key=%s mode=%s age_minutes=%.1f",
                    key,
                    mode or "<missing>",
                    age_minutes,
                )
            except Exception as e:
                logging.warning("[startup_recovery] failed to move stale key=%s err=%s", key, e)
            if scanned >= max_items:
                return recovered
    return recovered


def claim_job(pending_key: str) -> Optional[str]:
    """
    Atomic-ish claim: copy pending -> processing then delete pending.
    If copy/delete fails (race), return None.
    """
    name = pending_key.split("/")[-1]
    processing_key = PROCESSING + name
    try:
        s3.copy_object(
            Bucket=AWS_BUCKET,
            CopySource={"Bucket": AWS_BUCKET, "Key": pending_key},
            Key=processing_key,
            ContentType="application/json; charset=utf-8",
            MetadataDirective="REPLACE",
        )
        s3.delete_object(Bucket=AWS_BUCKET, Key=pending_key)
        return processing_key
    except Exception as exc:
        logging.warning("claim_job failed pending_key=%s err=%s", pending_key, exc)
        return None


def move_job(current_key: str, dest_prefix: str) -> str:
    name = current_key.split("/")[-1]
    dest_key = dest_prefix + name
    s3.copy_object(
        Bucket=AWS_BUCKET,
        CopySource={"Bucket": AWS_BUCKET, "Key": current_key},
        Key=dest_key,
        ContentType="application/json; charset=utf-8",
        MetadataDirective="REPLACE",
    )
    s3.delete_object(Bucket=AWS_BUCKET, Key=current_key)
    return dest_key


# -----------------------------
# Video processing
# -----------------------------
def open_video(path: str) -> cv2.VideoCapture:
    cap = cv2.VideoCapture(path)
    if not cap.isOpened():
        raise RuntimeError("Cannot open video")
    return cap


def write_mp4(out_path: str, fps: float, w: int, h: int) -> cv2.VideoWriter:
    fourcc = cv2.VideoWriter_fourcc(*"mp4v")
    vw = cv2.VideoWriter(out_path, fourcc, fps if fps > 0 else 25.0, (w, h))
    if not vw.isOpened():
        raise RuntimeError("Cannot open VideoWriter (mp4v)")
    return vw


def validate_video_file(path: str) -> None:
    """Fail fast if output video is empty/corrupt before upload."""
    if not os.path.exists(path):
        raise RuntimeError(f"Output video file does not exist: {path}")
    size = os.path.getsize(path)
    if size <= 0:
        raise RuntimeError(f"Output video file is empty: {path}")

    cap = cv2.VideoCapture(path)
    if not cap.isOpened():
        raise RuntimeError(f"Cannot open output video: {path}")
    frames_meta = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
    fps = float(cap.get(cv2.CAP_PROP_FPS) or 0.0)

    if frames_meta <= 0:
        n = 0
        while n < 50_000:
            ok, _ = cap.read()
            if not ok:
                break
            n += 1
        cap.release()
        if n <= 0:
            raise RuntimeError(f"Output video has no frames: {path}")
        logging.info("validate_video_file: counted %s frames (metadata was 0) %s", n, path)
        return

    cap.release()
    if fps <= 0:
        logging.warning(
            "validate_video_file: metadata fps missing for %s (frames=%s); skipping strict fps/duration check",
            path,
            frames_meta,
        )
        return
    duration = frames_meta / fps
    if duration <= 0.2:
        raise RuntimeError(f"Output video duration too short ({duration:.3f}s): {path}")


def _run_ffmpeg_transcode(cmd: List[str], timeout: int = 600) -> None:
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout)
    except subprocess.TimeoutExpired:
        raise RuntimeError(f"ffmpeg transcode timed out after {timeout}s")
    if proc.returncode != 0:
        raise RuntimeError(f"ffmpeg transcode failed: {proc.stderr[-400:]}")


def transcode_dots_mp4(input_path: str, output_path: str) -> None:
    """Ultra-compatible dots profile for browser playback from shared links."""
    ffmpeg_bin = shutil.which("ffmpeg")
    if not ffmpeg_bin:
        raise RuntimeError("ffmpeg not found. Install ffmpeg to enable browser-compatible MP4 output.")

    cmd = [
        ffmpeg_bin,
        "-y",
        "-i",
        input_path,
        "-f",
        "lavfi",
        "-i",
        "anullsrc=channel_layout=stereo:sample_rate=48000",
        "-map",
        "0:v:0",
        "-map",
        "1:a:0",
        "-vf",
        (
            "scale='if(gt(iw,854),854,iw)':'if(gt(ih,480),480,ih)':"
            "force_original_aspect_ratio=decrease,"
            "scale=trunc(iw/2)*2:trunc(ih/2)*2,"
            "fps=24,format=yuv420p"
        ),
        "-c:v",
        "libx264",
        "-profile:v",
        "baseline",
        "-level",
        "3.0",
        "-x264-params",
        "bframes=0:ref=1:cabac=0:keyint=48:min-keyint=48:scenecut=0",
        "-preset",
        "veryfast",
        "-crf",
        "28",
        "-maxrate",
        "1200k",
        "-bufsize",
        "2400k",
        "-g",
        "48",
        "-keyint_min",
        "48",
        "-sc_threshold",
        "0",
        "-movflags",
        "+faststart",
        "-vsync",
        "cfr",
        "-c:a",
        "aac",
        "-b:a",
        "64k",
        "-ar",
        "48000",
        "-ac",
        "2",
        "-shortest",
        "-r",
        "24",
        output_path,
    ]
    _run_ffmpeg_transcode(cmd)
    validate_video_file(output_path)


def transcode_skeleton_mp4(input_path: str, output_path: str) -> None:
    """Ultra-compatible skeleton profile for browser playback from links."""
    ffmpeg_bin = shutil.which("ffmpeg")
    if not ffmpeg_bin:
        raise RuntimeError("ffmpeg not found. Install ffmpeg to enable browser-compatible MP4 output.")

    cmd = [
        ffmpeg_bin,
        "-y",
        "-i",
        input_path,
        "-f",
        "lavfi",
        "-i",
        "anullsrc=channel_layout=stereo:sample_rate=48000",
        "-map",
        "0:v:0",
        "-map",
        "1:a:0",
        "-vf",
        (
            "scale='if(gt(iw,854),854,iw)':'if(gt(ih,480),480,ih)':"
            "force_original_aspect_ratio=decrease,"
            "scale=trunc(iw/2)*2:trunc(ih/2)*2,"
            "fps=24,format=yuv420p"
        ),
        "-c:v",
        "libx264",
        "-profile:v",
        "baseline",
        "-level",
        "3.0",
        "-x264-params",
        "bframes=0:ref=1:cabac=0:keyint=48:min-keyint=48:scenecut=0:colorprim=bt709:transfer=bt709:colormatrix=bt709",
        "-preset",
        "veryfast",
        "-crf",
        "28",
        "-maxrate",
        "1200k",
        "-bufsize",
        "2400k",
        "-g",
        "48",
        "-keyint_min",
        "48",
        "-sc_threshold",
        "0",
        "-movflags",
        "+faststart",
        "-vsync",
        "cfr",
        "-colorspace",
        "bt709",
        "-color_primaries",
        "bt709",
        "-color_trc",
        "bt709",
        "-c:a",
        "aac",
        "-b:a",
        "64k",
        "-ar",
        "48000",
        "-ac",
        "2",
        "-shortest",
        "-r",
        "24",
        output_path,
    ]
    _run_ffmpeg_transcode(cmd)
    validate_video_file(output_path)


# -----------------------------
# Dots / Skeleton overlay (simple, stable)
# -----------------------------
POSE_LANDMARK_IDS = [
    PoseLandmark.NOSE,
    PoseLandmark.LEFT_EYE,
    PoseLandmark.RIGHT_EYE,
    PoseLandmark.LEFT_SHOULDER,
    PoseLandmark.RIGHT_SHOULDER,
    PoseLandmark.LEFT_ELBOW,
    PoseLandmark.RIGHT_ELBOW,
    PoseLandmark.LEFT_WRIST,
    PoseLandmark.RIGHT_WRIST,
    PoseLandmark.LEFT_HIP,
    PoseLandmark.RIGHT_HIP,
    PoseLandmark.LEFT_KNEE,
    PoseLandmark.RIGHT_KNEE,
    PoseLandmark.LEFT_ANKLE,
    PoseLandmark.RIGHT_ANKLE,
]

# Skeleton: body only (no face). Neon style uses SKELETON_JOINT_IDS for joint highlights.
SKELETON_JOINT_IDS = [
    PoseLandmark.LEFT_SHOULDER,
    PoseLandmark.RIGHT_SHOULDER,
    PoseLandmark.LEFT_ELBOW,
    PoseLandmark.RIGHT_ELBOW,
    PoseLandmark.LEFT_WRIST,
    PoseLandmark.RIGHT_WRIST,
    PoseLandmark.LEFT_HIP,
    PoseLandmark.RIGHT_HIP,
    PoseLandmark.LEFT_KNEE,
    PoseLandmark.RIGHT_KNEE,
    PoseLandmark.LEFT_ANKLE,
    PoseLandmark.RIGHT_ANKLE,
]

SKELETON_EDGES = [
    (PoseLandmark.LEFT_SHOULDER, PoseLandmark.RIGHT_SHOULDER),
    (PoseLandmark.LEFT_SHOULDER, PoseLandmark.LEFT_ELBOW),
    (PoseLandmark.LEFT_ELBOW, PoseLandmark.LEFT_WRIST),
    (PoseLandmark.RIGHT_SHOULDER, PoseLandmark.RIGHT_ELBOW),
    (PoseLandmark.RIGHT_ELBOW, PoseLandmark.RIGHT_WRIST),
    (PoseLandmark.LEFT_HIP, PoseLandmark.RIGHT_HIP),
    (PoseLandmark.LEFT_SHOULDER, PoseLandmark.LEFT_HIP),
    (PoseLandmark.RIGHT_SHOULDER, PoseLandmark.RIGHT_HIP),
    (PoseLandmark.LEFT_HIP, PoseLandmark.LEFT_KNEE),
    (PoseLandmark.LEFT_KNEE, PoseLandmark.LEFT_ANKLE),
    (PoseLandmark.RIGHT_HIP, PoseLandmark.RIGHT_KNEE),
    (PoseLandmark.RIGHT_KNEE, PoseLandmark.RIGHT_ANKLE),
]


def _lm_to_px(lm, w: int, h: int) -> Tuple[int, int]:
    return int(lm.x * w), int(lm.y * h)


def _lm_to_px_spread(lm, w: int, h: int, spread: float) -> Tuple[int, int]:
    """Landmarks scaled toward frame center so overlay is smaller on screen (same idea as dots)."""
    nx = 0.5 + (float(lm.x) - 0.5) * spread
    ny = 0.5 + (float(lm.y) - 0.5) * spread
    return int(nx * w), int(ny * h)


def _frame_for_pose(frame: np.ndarray, max_width: int = 640) -> np.ndarray:
    """
    Downscale frame for MediaPipe inference to reduce CPU cost.
    Keep aspect ratio so normalized landmark coordinates remain consistent.
    """
    if frame is None:
        return frame
    h, w = frame.shape[:2]
    if w <= max_width or max_width <= 0:
        return frame
    new_h = max(1, int(round(h * (float(max_width) / float(w)))))
    return cv2.resize(frame, (max_width, new_h), interpolation=cv2.INTER_AREA)


def _prepare_pose_input_uniform(
    frame_bgr: np.ndarray, max_side: int
) -> Tuple[np.ndarray, int, int]:
    """Uniform resize for pose; landmarks stay 1:1 with output when aspect is preserved."""
    if frame_bgr is None:
        raise ValueError("empty frame")
    fh, fw = frame_bgr.shape[:2]
    if max_side <= 0 or max(fw, fh) <= max_side:
        return cv2.cvtColor(frame_bgr, cv2.COLOR_BGR2RGB), fw, fh
    scale = max_side / float(max(fw, fh))
    pw = max(1, int(round(fw * scale)))
    ph = max(1, int(round(fh * scale)))
    small = cv2.resize(frame_bgr, (pw, ph), interpolation=cv2.INTER_AREA)
    return cv2.cvtColor(small, cv2.COLOR_BGR2RGB), pw, ph


def generate_dots_video(input_path: str, out_path: str) -> None:
    """Generate dot motion video with bright white dots on black background (all landmarks)."""
    cap = open_video(input_path)
    fps = cap.get(cv2.CAP_PROP_FPS) or 25.0
    w = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
    h = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))

    vw = write_mp4(out_path, fps, w, h)

    # Draw at 2x then downscale. Larger radius + LINEAR downscale (not AREA) so dots stay
    # bright and visible on black; AREA averaging was shrinking/dimming tiny dots.
    scale = 2
    w2, h2 = w * scale, h * scale
    # Shrink pose toward frame center so the figure does not dominate the frame (was "too big").
    dots_pose_spread = 0.76  # 1.0 = full landmark span; lower = smaller on canvas
    dot_radius = 4  # at 2x canvas; keep visible but not oversized
    dot_glow = (220, 220, 220)  # subtle halo for contrast on very dark backgrounds
    dot_core = (255, 255, 255)

    process_every_n = 2 if fps >= 20 else 1
    last_landmarks = None

    with Pose(static_image_mode=False, model_complexity=1, enable_segmentation=False) as pose:
        frame_idx = 0
        while True:
            ok, frame = cap.read()
            if not ok:
                break
            frame_idx += 1

            output = np.zeros((h2, w2, 3), dtype=np.uint8)

            if frame_idx % process_every_n == 0:
                pose_frame = _frame_for_pose(frame, max_width=640)
                rgb = cv2.cvtColor(pose_frame, cv2.COLOR_BGR2RGB)
                res = pose.process(rgb)
                last_landmarks = res.pose_landmarks.landmark if res.pose_landmarks else None

            if last_landmarks:
                for lm in last_landmarks:
                    nx = 0.5 + (lm.x - 0.5) * dots_pose_spread
                    ny = 0.5 + (lm.y - 0.5) * dots_pose_spread
                    cx, cy = int(nx * w2), int(ny * h2)
                    if 0 <= cx < w2 and 0 <= cy < h2:
                        cv2.circle(output, (cx, cy), dot_radius + 1, dot_glow, -1)
                        cv2.circle(output, (cx, cy), dot_radius, dot_core, -1)

            output = cv2.resize(output, (w, h), interpolation=cv2.INTER_LINEAR)
            vw.write(output)

    vw.release()
    cap.release()


def generate_skeleton_video(input_path: str, out_path: str) -> None:
    cap = open_video(input_path)
    fps = cap.get(cv2.CAP_PROP_FPS) or 25.0
    w = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
    h = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))

    vw = write_mp4(out_path, fps, w, h)

    # ----- Fixed colors (never vary per video) -----
    # Cyan outer, white inner. Drawn DIRECTLY on the frame (no alpha blend, no Gaussian
    # blur) so the rendered color is exactly these BGR values regardless of what is
    # underneath. Previous "glow" pipeline composited a blurred cyan layer via
    # cv2.addWeighted(frame, 1.0, glow, 0.52, 0.0) which made the perceived color depend
    # on background brightness — bright videos saturated to white, dark videos stayed cyan.
    CYAN_BGR = (255, 255, 0)
    WHITE_BGR = (255, 255, 255)

    skeleton_pose_spread = 1.0
    vis_min = 0.35

    # ----- Laser-line skeleton -----
    # Goal: a crisp, bright "laser" stroke that reads the same on 1080p and
    # 540p clips. The previous resolution-aware scaling (min_side / 380, then
    # / 240) made the strokes plump on vertical phone footage. Switch to a
    # constant minimal double-stroke:
    #   outer cyan halo: 2 px
    #   inner white core: 1 px
    # Total visual width ~3 px with anti-aliasing — the thinnest "rim + core"
    # we can render while still keeping the cyan glow visible. Both values are
    # still overridable via SKELETON_OUTER_THICKNESS / SKELETON_INNER_THICKNESS
    # if a particular deployment wants a beefier line.
    _min_side = max(1, min(w, h))
    _auto_outer = 2  # constant 2-px cyan halo regardless of resolution
    try:
        outer_thick = max(1, int(os.getenv("SKELETON_OUTER_THICKNESS", "0") or "0") or _auto_outer)
    except ValueError:
        outer_thick = _auto_outer
    inner_thick = max(1, int(os.getenv("SKELETON_INNER_THICKNESS", "0") or "0") or max(1, outer_thick - 1))
    # Joint dots: tiny, so they read as bright pinpoints (laser endpoints)
    # rather than blobs. Cyan disc one pixel larger than the white core.
    outer_joint_r = max(2, outer_thick)
    inner_joint_r = max(1, inner_thick)

    # ----- Optional legacy glow mode, off by default -----
    # SKELETON_STYLE=glow restores the previous blurred-cyan compositing if ever wanted.
    skeleton_style = str(os.getenv("SKELETON_STYLE", "sharp")).strip().lower()
    glow_sigma = 2.0
    glow_alpha = 0.52
    kblur = max(3, int(round(glow_sigma * 4)) | 1)

    pose_max_side = max(0, int(os.getenv("SKELETON_POSE_MAX_SIDE", "640") or "640"))
    model_cx = max(0, min(2, int(os.getenv("SKELETON_MODEL_COMPLEXITY", "1") or "1")))
    last_landmarks = None
    process_every_n = 2 if fps >= 20 else 1

    logging.info(
        "[skeleton] render cfg: style=%s outer_thick=%s inner_thick=%s outer_r=%s inner_r=%s w=%s h=%s fps=%.1f",
        skeleton_style, outer_thick, inner_thick, outer_joint_r, inner_joint_r, w, h, float(fps or 0.0),
    )

    with Pose(
        static_image_mode=False,
        model_complexity=model_cx,
        enable_segmentation=False,
    ) as pose:
        frame_n = 0
        while True:
            ok, frame = cap.read()
            if not ok:
                break
            frame_n += 1
            if frame_n % 450 == 0:
                logging.info("[skeleton] encoding progress frames=%s (model_complexity=%s)", frame_n, model_cx)

            fh, fw = frame.shape[:2]
            if fw != w or fh != h:
                frame = cv2.resize(frame, (w, h), interpolation=cv2.INTER_LINEAR)

            if frame_n % process_every_n == 0:
                rgb, _pw, _ph = _prepare_pose_input_uniform(frame, pose_max_side)
                res = pose.process(rgb)
                last_landmarks = res.pose_landmarks.landmark if res.pose_landmarks else last_landmarks

            if last_landmarks:
                lms = last_landmarks

                if skeleton_style == "glow":
                    # Legacy blurred-glow compositor (kept for parity / opt-in via env).
                    glow_layer = np.zeros_like(frame, dtype=np.uint8)
                    for a, b in SKELETON_EDGES:
                        la, lb = lms[a], lms[b]
                        if la.visibility < vis_min or lb.visibility < vis_min:
                            continue
                        xa, ya = _lm_to_px_spread(la, w, h, skeleton_pose_spread)
                        xb, yb = _lm_to_px_spread(lb, w, h, skeleton_pose_spread)
                        cv2.line(glow_layer, (xa, ya), (xb, yb), CYAN_BGR, outer_thick + 4, cv2.LINE_AA)
                    for pid in SKELETON_JOINT_IDS:
                        lm = lms[pid]
                        if lm.visibility < vis_min:
                            continue
                        x, y = _lm_to_px_spread(lm, w, h, skeleton_pose_spread)
                        cv2.circle(glow_layer, (x, y), outer_joint_r + 2, CYAN_BGR, -1, lineType=cv2.LINE_AA)
                    glow_soft = cv2.GaussianBlur(glow_layer, (kblur, kblur), glow_sigma)
                    frame = cv2.addWeighted(frame, 1.0, glow_soft, glow_alpha, 0.0)

                # Sharp cyan outer stroke — always pure (255,255,0) BGR.
                for a, b in SKELETON_EDGES:
                    la, lb = lms[a], lms[b]
                    if la.visibility < vis_min or lb.visibility < vis_min:
                        continue
                    xa, ya = _lm_to_px_spread(la, w, h, skeleton_pose_spread)
                    xb, yb = _lm_to_px_spread(lb, w, h, skeleton_pose_spread)
                    cv2.line(frame, (xa, ya), (xb, yb), CYAN_BGR, outer_thick, cv2.LINE_AA)
                for pid in SKELETON_JOINT_IDS:
                    lm = lms[pid]
                    if lm.visibility < vis_min:
                        continue
                    x, y = _lm_to_px_spread(lm, w, h, skeleton_pose_spread)
                    cv2.circle(frame, (x, y), outer_joint_r, CYAN_BGR, -1, lineType=cv2.LINE_AA)

                # Sharp white inner stroke — always pure (255,255,255) BGR.
                for a, b in SKELETON_EDGES:
                    la, lb = lms[a], lms[b]
                    if la.visibility < vis_min or lb.visibility < vis_min:
                        continue
                    xa, ya = _lm_to_px_spread(la, w, h, skeleton_pose_spread)
                    xb, yb = _lm_to_px_spread(lb, w, h, skeleton_pose_spread)
                    cv2.line(frame, (xa, ya), (xb, yb), WHITE_BGR, inner_thick, cv2.LINE_AA)
                for pid in SKELETON_JOINT_IDS:
                    lm = lms[pid]
                    if lm.visibility < vis_min:
                        continue
                    x, y = _lm_to_px_spread(lm, w, h, skeleton_pose_spread)
                    cv2.circle(frame, (x, y), inner_joint_r, WHITE_BGR, -1, lineType=cv2.LINE_AA)

            vw.write(frame)

    vw.release()
    cap.release()


# -----------------------------
# First Impression (REAL from video via MediaPipe Pose)
# -----------------------------
@dataclass
class FirstImpressionResult:
    eye_contact_pct: float
    upright_pct: float
    stance_stability: float
    notes: Dict[str, str]


def analyze_first_impression(input_path: str, sample_every_n: int = 3, max_frames: int = 300) -> FirstImpressionResult:
    """
    Simple + stable:
    - Eye Contact proxy: nose centered between eyes & eyes visible
    - Uprightness: torso vector close to vertical
    - Stance stability: ankles distance variance (lower variance = more stable)
    """
    cap = open_video(input_path)
    w = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
    h = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))

    total = 0
    eye_ok = 0
    upright_ok = 0
    ankle_dist: List[float] = []

    with Pose(static_image_mode=False, model_complexity=1, enable_segmentation=False) as pose:
        i = 0
        while True:
            ok, frame = cap.read()
            if not ok:
                break
            i += 1
            if i % sample_every_n != 0:
                continue

            rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            res = pose.process(rgb)
            if not res.pose_landmarks:
                continue

            lms = res.pose_landmarks.landmark
            nose = lms[PoseLandmark.NOSE]
            leye = lms[PoseLandmark.LEFT_EYE]
            reye = lms[PoseLandmark.RIGHT_EYE]
            lsh = lms[PoseLandmark.LEFT_SHOULDER]
            rsh = lms[PoseLandmark.RIGHT_SHOULDER]
            lhip = lms[PoseLandmark.LEFT_HIP]
            rhip = lms[PoseLandmark.RIGHT_HIP]
            lank = lms[PoseLandmark.LEFT_ANKLE]
            rank = lms[PoseLandmark.RIGHT_ANKLE]

            # count only if main landmarks are visible enough
            if min(nose.visibility, leye.visibility, reye.visibility, lsh.visibility, rsh.visibility, lhip.visibility, rhip.visibility) < 0.5:
                continue

            total += 1

            # Eye contact proxy: nose x between eye x
            minx = min(leye.x, reye.x)
            maxx = max(leye.x, reye.x)
            if (minx <= nose.x <= maxx):
                eye_ok += 1

            # Uprightness: torso vector midHip->midShoulder angle to vertical
            mid_sh = np.array([(lsh.x + rsh.x) / 2.0, (lsh.y + rsh.y) / 2.0])
            mid_hip = np.array([(lhip.x + rhip.x) / 2.0, (lhip.y + rhip.y) / 2.0])
            v = mid_sh - mid_hip  # up direction (y smaller is up in image coords, but angle works)
            # angle to vertical axis (0, -1) in image space
            vert = np.array([0.0, -1.0])
            v_norm = np.linalg.norm(v) + 1e-9
            cosang = float(np.dot(v / v_norm, vert))
            ang = math.degrees(math.acos(max(-1.0, min(1.0, cosang))))
            if ang <= 15.0:
                upright_ok += 1

            # stance: ankle distance (normalized)
            if min(lank.visibility, rank.visibility) >= 0.5:
                dx = (lank.x - rank.x)
                dy = (lank.y - rank.y)
                ankle_dist.append(math.sqrt(dx*dx + dy*dy))

            if total >= max_frames:
                break

    cap.release()

    if total == 0:
        return FirstImpressionResult(
            eye_contact_pct=0.0,
            upright_pct=0.0,
            stance_stability=0.0,
            notes={"error": "insufficient_pose_frames"},
        )

    eye_pct = 100.0 * (eye_ok / total)
    upright_pct = 100.0 * (upright_ok / total)

    if len(ankle_dist) >= 10:
        std = float(np.std(np.array(ankle_dist)))
        # convert std to stability score (0..100) : lower std => higher score
        stability = max(0.0, min(100.0, 100.0 * (1.0 - (std / 0.20))))  # heuristic
    else:
        stability = 0.0

    notes = {}
    return FirstImpressionResult(
        eye_contact_pct=eye_pct,
        upright_pct=upright_pct,
        stance_stability=stability,
        notes=notes,
    )


# -----------------------------
# DOCX reports (EN/TH)
# -----------------------------
def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def _fmt_band(pct: float) -> str:
    if pct >= 70:
        return "Strong"
    if pct >= 40:
        return "Moderate"
    return "Needs improvement"


def _fmt_band_th(pct: float) -> str:
    if pct >= 70:
        return "ดีมาก"
    if pct >= 40:
        return "ปานกลาง"
    return "ควรพัฒนา"


def build_report_en(fi: FirstImpressionResult, meta: Dict[str, Any]) -> bytes:
    doc = Document()
    _add_title(doc, "AI People Reader — Report (EN)")

    doc.add_paragraph(f"Group: {meta.get('group_id','')}")
    doc.add_paragraph(f"User: {meta.get('user_name','')}")
    doc.add_paragraph(" ")

    _add_h2(doc, "First Impression (from uploaded video)")
    doc.add_paragraph(f"Eye Contact: {fi.eye_contact_pct:.1f}%  — {_fmt_band(fi.eye_contact_pct)}")
    doc.add_paragraph(f"Uprightness (Posture & Upper-Body Alignment): {fi.upright_pct:.1f}%  — {_fmt_band(fi.upright_pct)}")
    doc.add_paragraph(f"Stance (Lower-Body Stability & Grounding): {fi.stance_stability:.1f}/100  — {_fmt_band(fi.stance_stability)}")

    if fi.notes:
        _add_h2(doc, "Notes")
        for k, v in fi.notes.items():
            doc.add_paragraph(f"- {k}: {v}")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_report_th(fi: FirstImpressionResult, meta: Dict[str, Any]) -> bytes:
    doc = Document()
    _add_title(doc, "AI People Reader — รายงาน (TH)")

    doc.add_paragraph(f"Group: {meta.get('group_id','')}")
    doc.add_paragraph(f"ผู้ใช้: {meta.get('user_name','')}")
    doc.add_paragraph(" ")

    _add_h2(doc, "First Impression (วิเคราะห์จากวิดีโอจริงที่อัปโหลด)")
    doc.add_paragraph(f"Eye Contact: {fi.eye_contact_pct:.1f}%  — {_fmt_band_th(fi.eye_contact_pct)}")
    doc.add_paragraph(f"Uprightness (แนวลำตัว/การวางช่วงบน): {fi.upright_pct:.1f}%  — {_fmt_band_th(fi.upright_pct)}")
    doc.add_paragraph(f"Stance (ความมั่นคงช่วงล่าง/การยืน): {fi.stance_stability:.1f}/100  — {_fmt_band_th(fi.stance_stability)}")

    if fi.notes:
        _add_h2(doc, "หมายเหตุ")
        for k, v in fi.notes.items():
            doc.add_paragraph(f"- {k}: {v}")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# -----------------------------
# Processing heartbeat (prevents stale recovery during long skeleton encode)
# -----------------------------
class ProcessingHeartbeat:
    """Background thread that periodically re-writes the processing JSON to S3,
    refreshing its LastModified timestamp so stale-recovery doesn't reclaim it."""

    def __init__(self, processing_key: str, job: Dict[str, Any], interval_s: float = 60.0):
        self._key = processing_key
        self._job = job
        self._interval = max(10.0, interval_s)
        self._stop = threading.Event()
        self._thread: Optional[threading.Thread] = None

    def start(self) -> None:
        self._stop.clear()
        self._thread = threading.Thread(target=self._run, daemon=True, name="heartbeat")
        self._thread.start()

    def stop(self) -> None:
        self._stop.set()
        if self._thread is not None:
            self._thread.join(timeout=5)

    def _run(self) -> None:
        while not self._stop.wait(timeout=self._interval):
            try:
                s3_write_json(self._key, self._job)
                logging.info("[heartbeat-thread] refreshed processing key=%s", self._key)
            except Exception as exc:
                logging.warning("[heartbeat-thread] failed to refresh key=%s err=%s", self._key, exc)


def _get_peak_rss_mb() -> float:
    """Return peak RSS in MB using resource module (macOS returns bytes, Linux returns KB)."""
    try:
        ru = resource.getrusage(resource.RUSAGE_SELF)
        import sys
        if sys.platform == "darwin":
            return ru.ru_maxrss / (1024 * 1024)
        return ru.ru_maxrss / 1024
    except Exception:
        return -1.0


# -----------------------------
# Job processor
# -----------------------------
def process_job(job: Dict[str, Any]) -> Dict[str, Any]:
    mode = (job.get("mode") or "").strip()
    group_id = job.get("group_id", "")
    user_name = job.get("user_name", "")
    input_key = job.get("input_key", "")

    if not input_key:
        raise RuntimeError("Missing input_key")

    with tempfile.TemporaryDirectory() as td:
        in_path = os.path.join(td, "input.mp4")
        s3_download_to_file(input_key, in_path)

        if mode == "dots":
            out_key = job["output_key"]
            raw_path = os.path.join(td, "dots_raw.mp4")
            out_path = os.path.join(td, "dots.mp4")
            generate_dots_video(in_path, raw_path)
            validate_video_file(raw_path)
            transcode_dots_mp4(raw_path, out_path)
            s3_upload_file(out_path, out_key, "video/mp4")
            return {"ok": True, "mode": "dots", "output_key": out_key}

        if mode == "skeleton":
            wait_for_dots_mp4_before_skeleton(job, str(group_id or "").strip())
            out_key = job["output_key"]
            raw_path = os.path.join(td, "skeleton_raw.mp4")
            out_path = os.path.join(td, "skeleton.mp4")
            generate_skeleton_video(in_path, raw_path)
            validate_video_file(raw_path)
            transcode_skeleton_mp4(raw_path, out_path)
            s3_upload_file(out_path, out_key, "video/mp4")
            return {"ok": True, "mode": "skeleton", "output_key": out_key}

        if mode == "report_bundle":
            out_en_key = job["output_en_key"]
            out_th_key = job["output_th_key"]
            debug_key = job.get("output_debug_key", "")

            logging.info("Report bundle: EN=%s TH=%s", out_en_key, out_th_key)

            fi = analyze_first_impression(in_path)

            logging.info("First impression: eye_contact=%.1f%% upright=%.1f%% stance=%.1f", 
                        fi.eye_contact_pct, fi.upright_pct, fi.stance_stability)

            meta = {"group_id": group_id, "user_name": user_name}
            en_bytes = build_report_en(fi, meta)
            th_bytes = build_report_th(fi, meta)

            logging.info("Built reports: EN=%d bytes TH=%d bytes", len(en_bytes), len(th_bytes))

            s3_put_bytes(out_en_key, en_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            s3_put_bytes(out_th_key, th_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            logging.info("Uploaded reports to S3")

            if debug_key:
                dbg = {
                    "group_id": group_id,
                    "user_name": user_name,
                    "first_impression": {
                        "eye_contact_pct": fi.eye_contact_pct,
                        "upright_pct": fi.upright_pct,
                        "stance_stability": fi.stance_stability,
                        "notes": fi.notes,
                    },
                }
                s3_put_bytes(debug_key, json.dumps(dbg, ensure_ascii=False, indent=2).encode("utf-8"), "application/json; charset=utf-8")

            return {
                "ok": True,
                "mode": "report_bundle",
                "output_en_key": out_en_key,
                "output_th_key": out_th_key,
                "debug_key": debug_key,
                "first_impression": {
                    "eye_contact_pct": fi.eye_contact_pct,
                    "upright_pct": fi.upright_pct,
                    "stance_stability": fi.stance_stability,
                    "notes": fi.notes,
                },
            }

        raise RuntimeError(f"Unknown mode: {mode}")


def main_loop(poll_seconds: int = 3) -> None:
    logging.info("Worker started. Bucket=%s region=%s", AWS_BUCKET, AWS_REGION)
    logging.info(
        "Recovery cfg: processing_stale_minutes=%s processing_recovery_max_items=%s",
        PROCESSING_STALE_MINUTES,
        PROCESSING_RECOVERY_MAX_ITEMS,
    )
    logging.info(
        "Loop cfg: processing_recovery_interval_seconds=%s idle_heartbeat_seconds=%s poll_interval=%ss",
        PROCESSING_RECOVERY_INTERVAL_SECONDS,
        IDLE_HEARTBEAT_SECONDS,
        poll_seconds,
    )
    logging.info("Retry cfg: max_video_job_retries=%s", MAX_VIDEO_JOB_RETRIES)
    logging.info(
        "Pending scan: WORKER_PENDING_MAX_JSON_READS=%s (legacy WORKER_PENDING_MAX_SCAN; each pending JSON = one S3 read, reports skipped for selection)",
        WORKER_PENDING_MAX_JSON_READS,
    )
    logging.info(
        "Scope: WORKER_HANDLES_MODES=%s (default: dots,skeleton). mode=report stays in %s for the REPORT worker — both should run in parallel.",
        ",".join(sorted(WORKER_HANDLES_MODES)),
        PENDING,
    )
    logging.info(
        "Identity: worker_id=%s force_recover_on_startup=%s orphan_sweep_when_idle=%s",
        WORKER_ID,
        FORCE_RECOVER_ON_STARTUP,
        ORPHAN_SWEEP_WHEN_IDLE,
    )
    if FORCE_RECOVER_ON_STARTUP:
        forced = force_recover_all_processing_on_startup(PROCESSING_RECOVERY_MAX_ITEMS)
        logging.info("[startup_force_recovery] completed reclaimed=%s", forced)
    recovered = recover_stale_processing_jobs(PROCESSING_STALE_MINUTES, PROCESSING_RECOVERY_MAX_ITEMS)
    logging.info("[startup_recovery] completed recovered=%s", recovered)
    last_recovery_at = time.time()
    last_heartbeat_at = 0.0
    last_queue_diag_ts = 0.0
    last_orphan_sweep_at = 0.0
    # We track in-process active claims so an idle orphan sweep only runs when this worker
    # holds nothing — avoids a rare race where we'd reclaim our own freshly-written JSON
    # between status write and heartbeat start.
    active_claim_keys: set = set()

    while True:
        now = time.time()
        if (now - last_recovery_at) >= max(1, PROCESSING_RECOVERY_INTERVAL_SECONDS):
            recovered = recover_stale_processing_jobs(PROCESSING_STALE_MINUTES, PROCESSING_RECOVERY_MAX_ITEMS)
            logging.info("[startup_recovery] completed recovered=%s", recovered)
            last_recovery_at = now
        if (now - last_heartbeat_at) >= max(1, IDLE_HEARTBEAT_SECONDS):
            try:
                pending_count = count_jobs(PENDING)
                processing_count = count_jobs(PROCESSING)
            except Exception:
                pending_count = -1
                processing_count = -1
            logging.info(
                "[heartbeat] worker_alive worker_id=%s active_claims=%s pending_json_all_modes=%s processing=%s poll_interval=%ss "
                "| note: pending count is every jobs/pending/*.json (includes report). "
                "Video worker only claims dots/skeleton; mail is drained by REPORT worker via jobs/email_pending/.",
                WORKER_ID,
                len(active_claim_keys),
                pending_count,
                processing_count,
                poll_seconds,
            )
            last_heartbeat_at = now

            # Idle orphan sweep: if S3 shows processing>0 but this worker is holding no
            # claim, any foreign worker_id on disk belongs to a dead predecessor. This
            # catches the scenario where a predecessor's heartbeat refreshed LastModified
            # just before it died, so the age-based stale check won't reclaim for minutes.
            if (
                ORPHAN_SWEEP_WHEN_IDLE
                and isinstance(processing_count, int)
                and processing_count > 0
                and not active_claim_keys
                and (now - last_orphan_sweep_at) >= max(30.0, float(IDLE_HEARTBEAT_SECONDS))
            ):
                last_orphan_sweep_at = now
                try:
                    reclaimed = recover_orphaned_processing_by_worker_id(
                        WORKER_ID, PROCESSING_RECOVERY_MAX_ITEMS
                    )
                    logging.info(
                        "[orphan_sweep] idle sweep completed reclaimed=%s", reclaimed
                    )
                except Exception as exc:
                    logging.warning("[orphan_sweep] idle sweep errored err=%s", exc)

        # Scan deeper in pending queue so skeleton/dots are not starved
        # by many report jobs.
        keys = list_pending(limit=200)
        if not keys:
            # Throttled: if S3 has pending JSON but none are video jobs, explain why this worker is idle.
            if (now - last_queue_diag_ts) >= 60.0:
                last_queue_diag_ts = now
                try:
                    n_pending = count_jobs(PENDING)
                    if n_pending > 0:
                        logging.warning(
                            "[queue] bucket=%s: %s file(s) under %s but none are dots/skeleton pickable "
                            "(almost certainly all mode=report). This worker will stay idle; run/fix the REPORT worker. "
                            "If you expect skeleton jobs, check Streamlit + this worker share the same AWS_BUCKET.",
                            AWS_BUCKET,
                            n_pending,
                            PENDING,
                        )
                except Exception as exc:
                    logging.warning("[queue] could not count pending prefix=%s err=%s", PENDING, exc)
            time.sleep(poll_seconds)
            continue

        for pending_key in keys:
            # Peek pending job mode first, so this worker does not steal report jobs.
            # Retry on NoSuchKey: another worker may claim between list_pending and peek.
            pending_job = None
            last_peek_err: Exception | None = None
            for attempt in range(4):
                try:
                    pending_job = s3_read_json(pending_key)
                    last_peek_err = None
                    break
                except Exception as e:
                    last_peek_err = e
                    err_s = str(e).lower()
                    transient = "nosuchkey" in err_s or "404" in err_s or "does not exist" in err_s
                    if transient and attempt < 3:
                        time.sleep(0.35)
                        continue
                    pending_job = None
                    break
            if pending_job is None:
                # Key disappeared between list and get (report worker claimed it, another video worker, or delete).
                # Normal under parallel workers — try the next key in this snapshot, not the whole list.
                err_s = str(last_peek_err or "").lower()
                transient = "nosuchkey" in err_s or "404" in err_s or "does not exist" in err_s
                if transient:
                    logging.info(
                        "[peek] key already gone (race OK): %s",
                        pending_key,
                    )
                else:
                    logging.warning(
                        "[peek] failed for %s (%s); skipping this key",
                        pending_key,
                        last_peek_err,
                    )
                continue

            pending_mode = (pending_job.get("mode") or "").strip().lower()
            if pending_mode in REPORT_JOB_MODES:
                logging.info(
                    "Skipping pending report job job_id=%s group_id=%s mode=%s (reserved for report_worker)",
                    pending_job.get("job_id"),
                    pending_job.get("group_id"),
                    pending_mode,
                )
                continue

            if pending_mode == "skeleton" and pending_job.get("require_dots_output") is True:
                gid = str(pending_job.get("group_id") or "").strip()
                dots_key = str(pending_job.get("dots_output_wait_key") or "").strip() or (
                    f"jobs/output/groups/{gid}/dots.mp4" if gid else ""
                )
                if dots_key and not s3_head_exists(dots_key):
                    logging.info(
                        "[skeleton-defer] dots.mp4 not ready yet — leaving skeleton in pending "
                        "job_id=%s group_id=%s dots_key=%s",
                        pending_job.get("job_id"),
                        gid,
                        dots_key,
                    )
                    continue

            processing_key = claim_job(pending_key)
            if not processing_key:
                continue
            active_claim_keys.add(processing_key)

            job: Optional[Dict[str, Any]] = None
            heartbeat: Optional[ProcessingHeartbeat] = None
            try:
                job = s3_read_json(processing_key)
                mode = (job.get("mode") or "").strip()
                group_id = str(job.get("group_id") or "").strip()
                
                # Extra safety: keep report jobs for report_worker if one slips through.
                if mode in ("report", "report_th_en", "report_generator"):
                    logging.info(
                        "Skipping report mode job job_id=%s group_id=%s (handled by report_worker), moving back to pending",
                        job.get("job_id"),
                        group_id,
                    )
                    move_job(processing_key, PENDING)
                    continue
                
                job["status"] = "processing"
                job["worker_id"] = WORKER_ID
                job["claimed_at"] = datetime.now(timezone.utc).isoformat()
                s3_write_json(processing_key, job)

                if mode in ("skeleton", "dots"):
                    hb_interval = max(30.0, float(PROCESSING_STALE_MINUTES) * 60.0 / 3.0)
                    heartbeat = ProcessingHeartbeat(processing_key, job, interval_s=hb_interval)
                    heartbeat.start()

                logging.info("Processing job job_id=%s group_id=%s mode=%s", job.get("job_id"), group_id, mode)
                try:
                    rss_before = _get_peak_rss_mb()
                    result = process_job(job)
                    if mode == "skeleton":
                        rss_after = _get_peak_rss_mb()
                        logging.info(
                            "[skeleton] finished job_id=%s peak_rss_mb=%.1f (before=%.1f)",
                            job.get("job_id"), rss_after, rss_before,
                        )
                finally:
                    if heartbeat is not None:
                        heartbeat.stop()

                email_sent, email_status = send_mode_ready_email(job, result)

                job["status"] = "finished"
                job["result"] = result
                job["notification"] = {
                    "notify_email": str(job.get("notify_email") or "").strip(),
                    "sent": bool(email_sent),
                    "status": email_status,
                }
                s3_write_json(processing_key, job)
                move_job(processing_key, FINISHED)

                logging.info(
                    "Finished job job_id=%s group_id=%s mode=%s email_sent=%s email_status=%s",
                    job.get("job_id"),
                    group_id,
                    mode,
                    email_sent,
                    email_status,
                )
                # One job per poll: re-list pending next loop (avoids stale snapshot + clearer logs).
                gc.collect()
                break

            except Exception as e:
                if heartbeat is not None:
                    heartbeat.stop()
                    heartbeat = None
                _jid = (job.get("job_id") if isinstance(job, dict) else "") or ""
                _gid = (job.get("group_id") if isinstance(job, dict) else "") or ""
                _peak = _get_peak_rss_mb()
                logging.exception(
                    "Job failed job_id=%s group_id=%s peak_rss_mb=%.1f: %s",
                    _jid, _gid, _peak, e,
                )
                try:
                    job = s3_read_json(processing_key)
                    current_mode = str(job.get("mode") or "").strip().lower()
                    retry_count = int(job.get("retry_count") or 0)
                    if current_mode in ("dots", "skeleton") and retry_count < max(0, MAX_VIDEO_JOB_RETRIES):
                        next_retry = retry_count + 1
                        job["status"] = "pending"
                        job["retry_count"] = next_retry
                        job["message"] = str(e)
                        s3_write_json(processing_key, job)
                        move_job(processing_key, PENDING)
                        logging.warning(
                            "Requeued transient video job job_id=%s group_id=%s mode=%s retry=%s/%s",
                            job.get("job_id"),
                            job.get("group_id"),
                            current_mode,
                            next_retry,
                            MAX_VIDEO_JOB_RETRIES,
                        )
                        continue
                    job["status"] = "failed"
                    job["message"] = str(e)
                    s3_write_json(processing_key, job)
                except Exception:
                    pass
                move_job(processing_key, FAILED)
            finally:
                active_claim_keys.discard(processing_key)

        time.sleep(0.2)


if __name__ == "__main__":
    main_loop()
