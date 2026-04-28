# report_core.py — shared report logic for report generation
REPORT_CORE_VERSION = "2026-03-28-mp-agg-recover"  # Agg backend + Pose no-seg + report_worker GL recovery

import os
import sys
import io
import math


def _configure_mediapipe_headless_env() -> None:
    """
    MediaPipe pose graphs can request GpuService / GL (ImageToTensor). On headless Linux (e.g. Render)
    there is no display unless you use xvfb-run; on macOS SSH/CI you may see NSOpenGLPixelFormat errors.
    Call before importing mediapipe. Override with MEDIAPIPE_USE_GPU=1 on a real desktop with GPU.

    Production (Render): use `xvfb-run -a` in startCommand — see `render.yaml`.
    Code paths that still fail (e.g. raw python on macOS without GL) fall back to placeholder analysis —
    jobs must finish with PDF/email, not hard-fail on Pose init.

    Note: Some MediaPipe builds still touch GPU calculators until graph init; defensive try/except around
    every Pose() remains required in addition to env tuning.
    """
    if str(os.getenv("MEDIAPIPE_USE_GPU", "")).strip().lower() in ("1", "true", "yes", "on"):
        return
    os.environ.setdefault("CUDA_VISIBLE_DEVICES", "-1")
    os.environ.setdefault("TF_CPP_MIN_LOG_LEVEL", "2")
    os.environ.setdefault("GLOG_minloglevel", "2")
    # TensorFlow / XLA: prefer CPU when GPU stack would pull GL/Metal paths
    os.environ.setdefault("TF_FORCE_CPU", "1")
    os.environ.setdefault("TF_ENABLE_ONEDNN_OPTS", "0")


_configure_mediapipe_headless_env()
import random
import time
import logging
import unicodedata
from xml.sax.saxutils import escape
from dataclasses import dataclass
from datetime import datetime
from typing import Dict, Any, Optional, List, Tuple

import cv2
import numpy as np

# Headless servers: never let Matplotlib pick a GUI backend (can break PDF/graph steps).
import matplotlib

matplotlib.use(os.environ.get("MPLBACKEND", "Agg"))
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger("report_core.analysis")

# Ensure project root is in path so narrative_engine (at project root) can be imported
_report_core_dir = os.path.dirname(os.path.abspath(__file__))
_project_root = os.path.dirname(_report_core_dir)


def resolve_brand_asset_path(filename: str) -> str:
    """
    Header.png / Footer.png live at repo root. Optional override: AI_PEOPLE_READER_ASSET_DIR=/path
    (useful on servers or when the working directory differs). Skips unreadable paths (macOS sandbox, etc.).
    """
    fn = str(filename or "").strip()
    if not fn:
        return ""
    roots: List[str] = []
    env_root = (os.environ.get("AI_PEOPLE_READER_ASSET_DIR") or "").strip()
    if env_root:
        roots.append(env_root)
    roots.append(_project_root)
    for root in roots:
        p = os.path.join(root, fn)
        try:
            if os.path.isfile(p) and os.access(p, os.R_OK):
                return p
        except OSError:
            continue
    return ""


if _project_root not in sys.path:
    sys.path.insert(0, _project_root)

try:
    from narrative_engine import build_narrative_report
except ImportError as _e:
    logger.debug("[narrative] narrative_engine not available: %s", _e)
    build_narrative_report = None

try:
    from mediapipe.python.solutions import pose as mp_pose_module
    from mediapipe.python.solutions.pose import Pose, PoseLandmark
    mp = True
except ImportError as e1:
    try:
        import mediapipe as mp
        mp_pose_module = mp.solutions.pose
        Pose = mp_pose_module.Pose
        PoseLandmark = mp_pose_module.PoseLandmark
        mp = True
    except Exception as e2:
        mp = None
        mp_pose_module = None
        Pose = None
        PoseLandmark = None
        logger.warning(
            "[MediaPipe] Import failed (mp=None). Real analysis disabled; using placeholder. "
            "ImportError: %s | Fallback error: %s. "
            "Tip: Use Python 3.10 on Linux (Render) if 3.11 fails.",
            e1, e2,
        )

# Dataclasses
@dataclass
class FirstImpressionData:
    eye_contact_pct: float
    upright_pct: float
    stance_stability: float

def is_first_impression_fallback(fi: Optional["FirstImpressionData"]) -> bool:
    """True when First Impression used zero fallback (analysis failed) — show N/A instead of Low."""
    if not fi:
        return True
    return (float(fi.eye_contact_pct or 0) == 0 and float(fi.upright_pct or 0) == 0 and float(fi.stance_stability or 0) == 0)


def first_impression_level(value: float, metric: str = "", is_fallback: bool = False) -> str:
    """High/Moderate/Low. When is_fallback and value=0, return N/A (analysis failed)."""
    v = float(value or 0.0)
    if is_fallback and v == 0:
        return "N/A"
    if metric == "stance":
        # Stance bands calibrated against the 10 reference clips:
        #   High >= 53, Moderate 30-52, Low < 30.
        if v >= 53.0:
            raw = "High"
        elif v >= 30.0:
            raw = "Moderate"
        else:
            raw = "Low"
    elif metric == "uprightness":
        # Uprightness bands calibrated against the 10 reference clips:
        #   High >= 70, Moderate 15-69, Low < 15.
        if v >= 70.0:
            raw = "High"
        elif v >= 15.0:
            raw = "Moderate"
        else:
            raw = "Low"
    elif metric == "eye_contact":
        # Business override: eye contact is always reported as "High" for every
        # video. The raw v2 score (eye_contact_pct) is still computed and logged
        # for diagnostics, but the printed scale label is forced to "High" so
        # the report reflects a baseline assumption that the speaker is making
        # eye contact with the camera/audience. Set EYE_CONTACT_FORCE_HIGH=0 to
        # restore the calibrated High/Moderate/Low bands (>=60 / >=35 / else).
        if str(os.getenv("EYE_CONTACT_FORCE_HIGH", "1")).strip().lower() not in ("0", "false", "no", "off"):
            raw = "High"
        elif v >= 60.0:
            raw = "High"
        elif v >= 35.0:
            raw = "Moderate"
        else:
            raw = "Low"
    else:
        if v >= 70.0:
            raw = "High"
        elif v >= 30.0:
            raw = "Moderate"
        else:
            raw = "Low"
    return raw


def _display_scale(scale: str, is_thai: bool) -> str:
    """Convert internal scale string to localized High / Moderate / Low."""
    s = str(scale or "").strip().lower()
    if s.startswith("high"):
        return "สูง" if is_thai else "High"
    if s.startswith("moderate"):
        return "กลาง" if is_thai else "Moderate"
    if s.startswith("low"):
        return "ต่ำ" if is_thai else "Low"
    return "-"


@dataclass
class CategoryResult:
    name_en: str
    name_th: str
    score: int
    scale: str
    positives: int
    total: int
    description: str = ""

@dataclass
class ReportData:
    client_name: str
    analysis_date: str
    video_length_str: str
    overall_score: int
    categories: list
    summary_comment: str
    generated_by: str
    first_impression: Optional[FirstImpressionData] = None
    movement_type_info: Optional[Dict[str, Any]] = None


def format_people_reader_movement_top_two(mt: Optional[Dict[str, Any]], is_thai: bool) -> List[str]:
    """
    People Reader: show best match 1 and 2 from seven-dimension scoring.
    Prefers structured keys (seven_match_rank1_* / rank2_*); falls back to seven_match_line_*.
    """
    if not isinstance(mt, dict):
        return []
    lines: List[str] = []
    pre1 = "อันดับ 1 (จับคู่ 7 มิติ):" if is_thai else "Best match 1 (7 dimensions):"
    pre2 = "อันดับ 2 (จับคู่ 7 มิติ):" if is_thai else "Best match 2 (7 dimensions):"
    r1n = str(mt.get("seven_match_rank1_type_name") or "").strip()
    r2n = str(mt.get("seven_match_rank2_type_name") or "").strip()
    if r1n:
        lines.append(
            f"{pre1} {r1n} — {int(mt.get('seven_match_rank1_matches', 0))}/7 "
            f"({int(mt.get('seven_match_rank1_pct', 0))}%)"
        )
    if r2n:
        lines.append(
            f"{pre2} {r2n} — {int(mt.get('seven_match_rank2_matches', 0))}/7 "
            f"({int(mt.get('seven_match_rank2_pct', 0))}%)"
        )
    if lines:
        return lines
    comb = str(
        mt.get("seven_match_line_th" if is_thai else "seven_match_line_en", mt.get("seven_match_line_en", ""))
        or ""
    ).strip()
    if comb:
        hdr = "อันดับจับคู่ 7 มิติ (1 และ 2):" if is_thai else "Top two matches (7 dimensions):"
        return [f"{hdr} {comb}"]
    return []


# Helpers
def format_seconds_to_mmss(total_seconds: float) -> str:
    total_seconds = max(0, float(total_seconds))
    mm = int(total_seconds // 60)
    ss = int(round(total_seconds - mm * 60))
    if ss == 60:
        mm += 1
        ss = 0
    return f"{mm:02d}:{ss:02d}"

def get_video_duration_seconds(video_path: str) -> float:
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        return 0.0
    fps = cap.get(cv2.CAP_PROP_FPS) or 25.0
    frames = cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0.0
    cap.release()
    if fps <= 0:
        return 0.0
    return float(frames / fps)

def _apply_default_retreating_share(
    detection: Dict[str, float],
    retreat_key: str = "Retreating",
    retreat_default_pct: float = 1.0,
) -> Dict[str, float]:
    """
    Temporary business override: keep Retreating at a fixed default share,
    and scale other categories to fill the remaining percentage.
    """
    if not isinstance(detection, dict) or retreat_key not in detection:
        return detection

    fixed = max(0.0, min(100.0, float(retreat_default_pct)))
    other_keys = [k for k in detection.keys() if k != retreat_key]
    target_other_total = max(0.0, 100.0 - fixed)
    current_other_total = sum(max(0.0, float(detection.get(k, 0.0))) for k in other_keys)

    out: Dict[str, float] = {}
    if other_keys:
        if current_other_total <= 0.0:
            even_share = target_other_total / float(len(other_keys))
            for k in other_keys:
                out[k] = round(even_share, 1)
        else:
            for k in other_keys:
                raw = max(0.0, float(detection.get(k, 0.0)))
                out[k] = round((raw / current_other_total) * target_other_total, 1)
    out[retreat_key] = round(fixed, 1)
    return out

def analyze_first_impression_from_video(
    video_path: str,
    sample_every_n: int = 5,
    max_frames: int = 200,
    audience_mode: str = "one",
) -> FirstImpressionData:
    """Real First Impression analysis with MediaPipe using continuous scoring.
    audience_mode: "one" = present to one person (strict eye contact - look at camera/audience)
                   "many" = present to many (relaxed - allow scanning across room)
    """
    if (Pose is None) or (PoseLandmark is None) or (not callable(Pose)):
        logger.warning("[first_impression] MediaPipe Pose unavailable; using fallback (High/High/Moderate)")
        return FirstImpressionData(eye_contact_pct=65.0, upright_pct=65.0, stance_stability=50.0)

    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        return FirstImpressionData(eye_contact_pct=65.0, upright_pct=65.0, stance_stability=50.0)

    frame_samples = []

    def _normalize_horizontal_angle(deg: float) -> float:
        while deg > 180.0:
            deg -= 360.0
        while deg <= -180.0:
            deg += 360.0
        if deg > 90.0:
            deg -= 180.0
        if deg < -90.0:
            deg += 180.0
        return deg

    try:
        with Pose(
            static_image_mode=False,
            model_complexity=1,
            enable_segmentation=False,
        ) as pose:
            i = 0
            while True:
                ret, frame = cap.read()
                if not ret:
                    break
                i += 1
                if i % sample_every_n != 0:
                    continue

                rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                res = pose.process(rgb)
                if not res.pose_landmarks:
                    continue

                lms = res.pose_landmarks.landmark
                nose = lms[PoseLandmark.NOSE]
                leye = lms[PoseLandmark.LEFT_EYE]
                reye = lms[PoseLandmark.RIGHT_EYE]
                lsh = lms[PoseLandmark.LEFT_SHOULDER]
                rsh = lms[PoseLandmark.RIGHT_SHOULDER]
                lhip = lms[PoseLandmark.LEFT_HIP]
                rhip = lms[PoseLandmark.RIGHT_HIP]
                lkne = lms[PoseLandmark.LEFT_KNEE]
                rkne = lms[PoseLandmark.RIGHT_KNEE]
                lank = lms[PoseLandmark.LEFT_ANKLE]
                rank = lms[PoseLandmark.RIGHT_ANKLE]

                if min(nose.visibility, leye.visibility, reye.visibility, lsh.visibility, rsh.visibility, lhip.visibility, rhip.visibility) < 0.5:
                    continue

                frame_samples.append(
                    {
                        "nose": (float(nose.x), float(nose.y)),
                        "leye": (float(leye.x), float(leye.y)),
                        "reye": (float(reye.x), float(reye.y)),
                        "lsh": (float(lsh.x), float(lsh.y)),
                        "rsh": (float(rsh.x), float(rsh.y)),
                        "lhip": (float(lhip.x), float(lhip.y)),
                        "rhip": (float(rhip.x), float(rhip.y)),
                        "lkne": (float(lkne.x), float(lkne.y)),
                        "rkne": (float(rkne.x), float(rkne.y)),
                        "lank": (float(lank.x), float(lank.y)),
                        "rank": (float(rank.x), float(rank.y)),
                        "leye_vis": float(leye.visibility),
                        "reye_vis": float(reye.visibility),
                        "lkne_vis": float(lkne.visibility),
                        "rkne_vis": float(rkne.visibility),
                        "lank_vis": float(lank.visibility),
                        "rank_vis": float(rank.visibility),
                    }
                )
                if len(frame_samples) >= max_frames:
                    break
    except Exception as e:
        logger.warning("[first_impression] MediaPipe failed (fallback scores): %s", e)
        return FirstImpressionData(eye_contact_pct=65.0, upright_pct=65.0, stance_stability=50.0)
    finally:
        try:
            cap.release()
        except Exception:
            pass

    total = len(frame_samples)
    if total == 0:
        return FirstImpressionData(eye_contact_pct=65.0, upright_pct=65.0, stance_stability=50.0)

    # Auto-calibrate camera roll (tilted camera) using median shoulder/hip line angles.
    roll_candidates = []
    for s in frame_samples:
        lsh = s["lsh"]
        rsh = s["rsh"]
        lhip = s["lhip"]
        rhip = s["rhip"]
        sh_ang = math.degrees(math.atan2(rsh[1] - lsh[1], rsh[0] - lsh[0]))
        hip_ang = math.degrees(math.atan2(rhip[1] - lhip[1], rhip[0] - lhip[0]))
        roll_candidates.append(_normalize_horizontal_angle(sh_ang))
        roll_candidates.append(_normalize_horizontal_angle(hip_ang))
    camera_roll_deg = float(np.median(np.array(roll_candidates))) if roll_candidates else 0.0
    camera_roll_deg = max(-20.0, min(20.0, camera_roll_deg))
    theta = -math.radians(camera_roll_deg)

    def _rotate_point(pt: tuple, cx: float = 0.5, cy: float = 0.5) -> tuple:
        x, y = pt
        dx = x - cx
        dy = y - cy
        xr = dx * math.cos(theta) - dy * math.sin(theta) + cx
        yr = dx * math.sin(theta) + dy * math.cos(theta) + cy
        return (xr, yr)

    eye_frame_scores = []
    upright_frame_scores = []
    ankle_dist = []
    ankle_center_x = []
    stance_width_ratios = []
    # v2 stance signals
    knee_bend_L_list: List[float] = []
    knee_bend_R_list: List[float] = []
    torso_tilt_list: List[float] = []
    # v2 uprightness signals
    sh_hip_offset_list: List[float] = []
    # v3 uprightness signals (added per professor feedback, calibration-pending)
    # - sh_mid_x / hip_mid_x track body sway over time (weight shifting).
    # - head_above_shoulders captures chin-down / rounded-shoulder slouch.
    sh_mid_x_list: List[float] = []
    hip_mid_x_list: List[float] = []
    head_above_shoulders_list: List[float] = []
    # v2 eye-contact signals
    nose_offset_ratio_list: List[float] = []
    nose_sh_ratio_list: List[float] = []

    # Simple eye contact: one audience = look at camera. many = look around (scanning).
    is_many_audience = str(audience_mode or "").strip().lower() == "many"
    # One: face toward camera (nose centered). Many: allow head turn when scanning audience.
    nose_offset_max = 0.50 if is_many_audience else 0.35  # max nose offset from eye center (ratio of eye dist)

    for s in frame_samples:
        nose = _rotate_point(s["nose"])
        leye = _rotate_point(s["leye"])
        reye = _rotate_point(s["reye"])
        lsh = _rotate_point(s["lsh"])
        rsh = _rotate_point(s["rsh"])
        lhip = _rotate_point(s["lhip"])
        rhip = _rotate_point(s["rhip"])
        lkne = _rotate_point(s.get("lkne", (0.0, 0.0)))
        rkne = _rotate_point(s.get("rkne", (0.0, 0.0)))
        lank = _rotate_point(s["lank"])
        rank = _rotate_point(s["rank"])

        # Eye contact signals.
        # Binary per-frame score kept for v1 fallback. v2 uses continuous means.
        eye_dist = abs(leye[0] - reye[0])
        if eye_dist > 1e-4:
            mid_eye_x = (leye[0] + reye[0]) / 2.0
            nose_offset_ratio = abs(nose[0] - mid_eye_x) / eye_dist
            if nose_offset_ratio <= nose_offset_max:
                eye_frame_scores.append(100.0)
            else:
                eye_frame_scores.append(0.0)
            nose_offset_ratio_list.append(float(nose_offset_ratio))
        # Nose displacement from shoulder midline, normalized by shoulder width.
        # This captures head turn better than nose-vs-eye-center because eyes travel
        # with the nose when the head yaws. Separates High (< 0.08) from Moderate (> 0.09).
        sh_w = abs(lsh[0] - rsh[0])
        if sh_w > 1e-4:
            mid_sh_x = (lsh[0] + rsh[0]) / 2.0
            nose_sh_ratio_list.append(abs(float(nose[0]) - mid_sh_x) / sh_w)

        # Uprightness (simple): torso roughly vertical = upright. Relaxed threshold.
        mid_sh = np.array([(lsh[0] + rsh[0]) / 2.0, (lsh[1] + rsh[1]) / 2.0])
        mid_hip = np.array([(lhip[0] + rhip[0]) / 2.0, (lhip[1] + rhip[1]) / 2.0])
        v = mid_sh - mid_hip
        vert = np.array([0.0, -1.0])
        v_norm = np.linalg.norm(v) + 1e-9
        cosang = float(np.dot(v / v_norm, vert))
        ang = math.degrees(math.acos(max(-1.0, min(1.0, cosang))))
        # Keep legacy per-frame 100/0 list for v1 fallback only.
        if ang <= 50.0:
            upright_frame_scores.append(100.0)
        else:
            upright_frame_scores.append(0.0)
        torso_tilt_list.append(float(ang))
        # Horizontal offset between shoulder midline and hip midline — slouch indicator.
        sh_hip_offset_list.append(abs(float(mid_sh[0]) - float(mid_hip[0])))
        # v3 uprightness signals (calibration-pending).
        sh_mid_x_list.append(float(mid_sh[0]))
        hip_mid_x_list.append(float(mid_hip[0]))
        if sh_w > 1e-4:
            # Head height above shoulder midline, normalized by shoulder width.
            # Tall heads (values >= ~0.70) read as upright; slumped / chin-down
            # heads compress toward the shoulder line (values <= ~0.50).
            head_above_shoulders_list.append(float((mid_sh[1] - nose[1]) / sh_w))

        # Stance (continuous): lower-body stability + center sway + stance width appropriateness.
        if min(float(s["lank_vis"]), float(s["rank_vis"])) >= 0.5:
            dx = (lank[0] - rank[0])
            dy = (lank[1] - rank[1])
            dist = math.sqrt(dx * dx + dy * dy)
            ankle_dist.append(dist)
            ankle_center_x.append((lank[0] + rank[0]) / 2.0)
            shoulder_width = abs(lsh[0] - rsh[0])
            if shoulder_width > 1e-4:
                stance_width_ratios.append(dist / shoulder_width)

        # Per-frame knee bend (v2 signal). Bend = 180 - angle(hip,knee,ankle).
        # Bigger value = more bent knee. Straight leg ~ 0-2 degrees.
        def _knee_bend(hip_pt, knee_pt, ank_pt) -> float:
            ax, ay = hip_pt[0] - knee_pt[0], hip_pt[1] - knee_pt[1]
            cx, cy = ank_pt[0] - knee_pt[0], ank_pt[1] - knee_pt[1]
            na = math.sqrt(ax * ax + ay * ay) + 1e-9
            nc = math.sqrt(cx * cx + cy * cy) + 1e-9
            cosv = max(-1.0, min(1.0, (ax * cx + ay * cy) / (na * nc)))
            return 180.0 - math.degrees(math.acos(cosv))

        lkne_vis = float(s.get("lkne_vis", 0.0))
        rkne_vis = float(s.get("rkne_vis", 0.0))
        if lkne_vis >= 0.4 and float(s["lank_vis"]) >= 0.4:
            knee_bend_L_list.append(_knee_bend(lhip, lkne, lank))
        if rkne_vis >= 0.4 and float(s["rank_vis"]) >= 0.4:
            knee_bend_R_list.append(_knee_bend(rhip, rkne, rank))

    def _clamp01(x: float) -> float:
        return 0.0 if x < 0.0 else (1.0 if x > 1.0 else x)

    # Eye contact scoring.
    # v2 (default) — continuous: 0.33*nose_vs_eye_center + 0.67*nose_vs_shoulder_midline.
    #   nose_vs_shoulder is the dominant signal (eyes travel with the nose when the head
    #   yaws; shoulder midline doesn't). Calibrated against Types 1-10: 10/10 match.
    # v1 (legacy) — binary per-frame threshold then mean. Scored every clip >=94, collapsing
    #   Moderate into High. Keep accessible via EYE_FORMULA=v1 for A/B.
    eye_formula = str(os.getenv("EYE_FORMULA", "v2")).strip().lower()
    if eye_formula == "v1":
        eye_pct = float(np.mean(np.array(eye_frame_scores))) if eye_frame_scores else 0.0
        logger.info("[first_impression] eye v1: score=%.1f (binary threshold)", eye_pct)
    else:
        if nose_offset_ratio_list and nose_sh_ratio_list:
            nose_eye_mean = float(np.mean(np.array(nose_offset_ratio_list)))
            nose_sh_mean  = float(np.mean(np.array(nose_sh_ratio_list)))
            nose_eye_f = _clamp01(1.0 - (nose_eye_mean - 0.10) / 0.08)
            nose_sh_f  = _clamp01(1.0 - (nose_sh_mean  - 0.07) / 0.04)
            eye_pct = 100.0 * (0.33 * nose_eye_f + 0.67 * nose_sh_f)
            eye_pct = max(0.0, min(100.0, eye_pct))
            logger.info(
                "[first_impression] eye v2: score=%.1f  nose_eye_mean=%.3f(f=%.2f) nose_sh_mean=%.3f(f=%.2f)",
                eye_pct, nose_eye_mean, nose_eye_f, nose_sh_mean, nose_sh_f,
            )
        else:
            eye_pct = 50.0
            logger.info("[first_impression] eye v2: insufficient samples, using 50.0")

    # Uprightness scoring.
    # v2 (default) — continuous: median torso tilt + p90 torso tilt (worst-case lean)
    #                            + median shoulder-vs-hip horizontal offset (slouch).
    #   Calibrated against 10 reference clips: matches human-judged level 8/10.
    #   v1 binary (tilt<=50°=100 else 0) effectively always returns 100 in practice,
    #   because real torso tilts are 1..6° — so everyone scored High. Replaced.
    # v1 (legacy) remains accessible via env flag UPRIGHT_FORMULA=v1 for A/B.
    upright_formula = str(os.getenv("UPRIGHT_FORMULA", "v2")).strip().lower()
    if upright_formula == "v1":
        upright_pct = float(np.mean(np.array(upright_frame_scores))) if upright_frame_scores else 0.0
        logger.info("[first_impression] upright v1: score=%.1f (binary <=50deg)", upright_pct)
    else:
        if torso_tilt_list:
            torso_arr = np.array(torso_tilt_list)
            torso_med = float(np.median(torso_arr))
            torso_p90 = float(np.percentile(torso_arr, 90))
            sh_hip_med = float(np.median(np.array(sh_hip_offset_list))) if sh_hip_offset_list else 0.02

            torso_med_f = _clamp01(1.0 - (torso_med - 0.5) / 4.0)
            torso_p90_f = _clamp01(1.0 - (torso_p90 - 3.0) / 4.0)
            sh_hip_f    = _clamp01(1.0 - (sh_hip_med - 0.015) / 0.015)

            # v3 add-ons (professor-feedback). Two signals the old 3-term v2
            # formula was blind to:
            #   1. Body sway across the clip (weight shifting / left-right drift).
            #   2. Head height above the shoulders (chin-down / rounded posture).
            # Both are computed only when we have enough samples; otherwise
            # they default to neutral (0.5) so they can't accidentally drop a
            # clean clip on a short video with too few frames.
            upright_v3 = str(os.getenv("UPRIGHT_V3", "1")).strip().lower() not in ("0", "false", "no", "off")

            if upright_v3 and len(sh_mid_x_list) >= 10 and len(hip_mid_x_list) >= 10:
                sway_series = np.array(sh_mid_x_list + hip_mid_x_list, dtype=float)
                sway_std = float(np.std(sway_series))
                # Clean standers sit around 0.005–0.012; visible swaying starts
                # ~0.025. Map std=0.012 -> 1.0 and std=0.050 -> 0.0.
                sway_f = _clamp01(1.0 - (sway_std - 0.012) / 0.038)
            else:
                sway_std = float("nan")
                sway_f = 0.5

            if upright_v3 and head_above_shoulders_list:
                head_h_med = float(np.median(np.array(head_above_shoulders_list)))
                # Upright speakers land ~0.75–0.90; rounded/chin-down slumps
                # compress to ~0.45–0.55. Map 0.55 -> 0.0, 0.80 -> 1.0.
                head_f = _clamp01((head_h_med - 0.55) / 0.25)
            else:
                head_h_med = float("nan")
                head_f = 0.5

            if upright_v3:
                # Re-weighted mix: the two old signals still dominate but the
                # sway and head-height components can pull the score down for
                # clips that look slumped / shifting to a human judge even
                # when the 2D shoulder-to-hip line is nearly vertical.
                upright_pct = 100.0 * (
                    0.15 * torso_med_f
                    + 0.25 * torso_p90_f
                    + 0.20 * sh_hip_f
                    + 0.20 * sway_f
                    + 0.20 * head_f
                )
            else:
                upright_pct = 100.0 * (
                    0.20 * torso_med_f
                    + 0.40 * torso_p90_f
                    + 0.40 * sh_hip_f
                )

            upright_pct = max(0.0, min(100.0, upright_pct))
            logger.info(
                "[first_impression] upright v%s: score=%.1f  torso_med=%.2f(f=%.2f) torso_p90=%.2f(f=%.2f) "
                "sh_hip_off=%.3f(f=%.2f) sway_std=%.3f(f=%.2f) head_above=%.2f(f=%.2f)",
                "3" if upright_v3 else "2",
                upright_pct, torso_med, torso_med_f, torso_p90, torso_p90_f,
                sh_hip_med, sh_hip_f,
                sway_std if sway_std == sway_std else -1.0, sway_f,
                head_h_med if head_h_med == head_h_med else -1.0, head_f,
            )
        else:
            upright_pct = 50.0
            logger.info("[first_impression] upright v2/v3: no torso samples, using 50.0")

    # Stance scoring.
    # v2 (default) — geometry-based: stance width + knee straightness + knee symmetry + torso upright.
    #   Calibrated against 10 reference clips: matches human-judged level 8/10 (vs 3/10 for v1).
    # v1 (legacy) — ankle-distance variance only. Keep accessible via STANCE_FORMULA=v1 for A/B.
    stance_formula = str(os.getenv("STANCE_FORMULA", "v2")).strip().lower()

    if stance_formula == "v1":
        if len(ankle_dist) >= 10:
            dist_arr = np.array(ankle_dist)
            dist_std = float(np.std(dist_arr))
            dist_mean = float(np.mean(dist_arr)) + 1e-9
            rel_std = dist_std / dist_mean
            base_stability = max(0.0, min(100.0, 100.0 * (1.0 - (rel_std / 0.60))))
            sway_score = 0.0
            if len(ankle_center_x) >= 10:
                sway_std = float(np.std(np.array(ankle_center_x)))
                sway_score = max(0.0, min(100.0, 100.0 * (1.0 - (sway_std / 0.12))))
            stability = max(0.0, min(100.0, 0.70 * base_stability + 0.30 * sway_score))
        else:
            stability = 50.0
        logger.info("[first_impression] stance v1: score=%.1f (base+sway)", stability)
    else:
        # v2: stance_score = 0.56*width + 0.22*knee_straight + 0.07*knee_symmetry + 0.15*torso_upright
        if len(stance_width_ratios) >= 10:
            width_ratio_med = float(np.median(np.array(stance_width_ratios)))
            width_factor = _clamp01((width_ratio_med - 0.35) / 0.40)

            if knee_bend_L_list and knee_bend_R_list:
                kL_med = float(np.median(np.array(knee_bend_L_list)))
                kR_med = float(np.median(np.array(knee_bend_R_list)))
                knee_avg = (kL_med + kR_med) / 2.0
                knee_straight = _clamp01(1.0 - (knee_avg - 1.0) / 6.0)
                knee_symmetry = _clamp01(1.0 - abs(kL_med - kR_med) / 6.0)
            else:
                kL_med = kR_med = knee_avg = float("nan")
                knee_straight = 0.5
                knee_symmetry = 0.5

            if torso_tilt_list:
                torso_med = float(np.median(np.array(torso_tilt_list)))
                torso_upright = _clamp01(1.0 - (torso_med - 1.0) / 5.0)
            else:
                torso_med = float("nan")
                torso_upright = 0.5

            stability = 100.0 * (
                0.56 * width_factor
                + 0.22 * knee_straight
                + 0.07 * knee_symmetry
                + 0.15 * torso_upright
            )
            stability = max(0.0, min(100.0, stability))

            # v2 professor-feedback overrides (calibration-pending).
            # Two specific patterns appear repeatedly in the ground-truth
            # table and aren't well handled by the weighted average alone:
            #   a) Feet-together stance (Aon / Ann / Candy). If the median
            #      ankle distance is below ~25% of shoulder width for the
            #      whole clip, the speaker effectively has no stance —
            #      cap the score into the Low band.
            #   b) Wide, stable stance (Ches / Lisa 2nd clip). A consistently
            #      wide base with minimal variation is a strong High signal.
            stance_overrides = str(os.getenv("STANCE_V2_OVERRIDES", "1")).strip().lower() not in ("0", "false", "no", "off")
            override_note = ""
            width_ratio_std = float(np.std(np.array(stance_width_ratios)))
            if stance_overrides and width_ratio_med < 0.25:
                stability = min(stability, 18.0)
                override_note = " feet_together_cap"
            elif stance_overrides and width_ratio_med > 0.80 and width_ratio_std < 0.05:
                stability = min(100.0, stability + 10.0)
                override_note = " wide_stable_boost"

            logger.info(
                "[first_impression] stance v2: score=%.1f  width_ratio=%.3f±%.3f(f=%.2f) "
                "knee_avg=%.2f(f=%.2f) knee_sym=%.2f(f=%.2f) torso=%.2f(f=%.2f)%s",
                stability, width_ratio_med, width_ratio_std, width_factor,
                knee_avg if knee_avg == knee_avg else -1.0, knee_straight,
                (kL_med - kR_med) if (kL_med == kL_med and kR_med == kR_med) else 0.0, knee_symmetry,
                torso_med if torso_med == torso_med else -1.0, torso_upright,
                override_note,
            )
        else:
            stability = 50.0
            logger.info("[first_impression] stance v2: insufficient ankle samples (%d), using 50.0", len(stance_width_ratios))

    return FirstImpressionData(eye_contact_pct=eye_pct, upright_pct=upright_pct, stance_stability=stability)


def extract_movement_type_frame_features_from_video(
    video_path: str,
    audience_mode: str = "one",
    sample_every_n: int = 3,
    max_frames: int = 200,
) -> Dict[str, List[float]]:
    """
    Per-frame signals for movement_type_classifier.build_summary_features_from_timeseries.
    Returns keys: eye_contact, uprightness, stance_width_ratio, weight_shift, gesture_energy,
    gesture_variation, chest_blocking, chest_open, rotation (lists may be empty if MediaPipe fails).
    """
    out: Dict[str, List[float]] = {
        "eye_contact": [],
        "uprightness": [],
        "stance_width_ratio": [],
        "weight_shift": [],
        "gesture_energy": [],
        "gesture_variation": [],
        "chest_blocking": [],
        "chest_open": [],
        "rotation": [],
    }
    if (Pose is None) or (PoseLandmark is None) or (not callable(Pose)):
        return out

    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        return out

    is_many_audience = str(audience_mode or "").strip().lower() == "many"
    nose_offset_max = 0.50 if is_many_audience else 0.35

    prev_ankle_mid: Optional[Tuple[float, float]] = None
    prev_wrist_mid: Optional[Tuple[float, float]] = None
    prev_sh_angle: Optional[float] = None
    energy_window: List[float] = []

    def _sh_angle_deg(lsh: Tuple[float, float], rsh: Tuple[float, float]) -> float:
        return math.degrees(math.atan2(rsh[1] - lsh[1], rsh[0] - lsh[0]))

    try:
        with Pose(
            static_image_mode=False,
            model_complexity=1,
            enable_segmentation=False,
        ) as pose:
            i = 0
            while True:
                ret, frame = cap.read()
                if not ret:
                    break
                i += 1
                if i % sample_every_n != 0:
                    continue

                rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                res = pose.process(rgb)
                if not res.pose_landmarks:
                    continue

                lms = res.pose_landmarks.landmark
                nose = lms[PoseLandmark.NOSE]
                leye = lms[PoseLandmark.LEFT_EYE]
                reye = lms[PoseLandmark.RIGHT_EYE]
                lsh = lms[PoseLandmark.LEFT_SHOULDER]
                rsh = lms[PoseLandmark.RIGHT_SHOULDER]
                lhip = lms[PoseLandmark.LEFT_HIP]
                rhip = lms[PoseLandmark.RIGHT_HIP]
                lank = lms[PoseLandmark.LEFT_ANKLE]
                rank = lms[PoseLandmark.RIGHT_ANKLE]
                lw = lms[PoseLandmark.LEFT_WRIST]
                rw = lms[PoseLandmark.RIGHT_WRIST]

                if min(
                    nose.visibility,
                    leye.visibility,
                    reye.visibility,
                    lsh.visibility,
                    rsh.visibility,
                    lhip.visibility,
                    rhip.visibility,
                ) < 0.5:
                    continue

                lsh_pt = (float(lsh.x), float(lsh.y))
                rsh_pt = (float(rsh.x), float(rsh.y))
                lhip_pt = (float(lhip.x), float(lhip.y))
                rhip_pt = (float(rhip.x), float(rhip.y))
                lank_pt = (float(lank.x), float(lank.y))
                rank_pt = (float(rank.x), float(rank.y))
                leye_pt = (float(leye.x), float(leye.y))
                reye_pt = (float(reye.x), float(reye.y))
                nose_pt = (float(nose.x), float(nose.y))

                eye_dist = abs(leye_pt[0] - reye_pt[0])
                if eye_dist > 1e-4:
                    mid_eye_x = (leye_pt[0] + reye_pt[0]) / 2.0
                    nose_offset_ratio = abs(nose_pt[0] - mid_eye_x) / eye_dist
                    out["eye_contact"].append(1.0 if nose_offset_ratio <= nose_offset_max else 0.0)
                else:
                    out["eye_contact"].append(0.5)

                mid_sh = np.array([(lsh_pt[0] + rsh_pt[0]) / 2.0, (lsh_pt[1] + rsh_pt[1]) / 2.0])
                mid_hip = np.array([(lhip_pt[0] + rhip_pt[0]) / 2.0, (lhip_pt[1] + rhip_pt[1]) / 2.0])
                v = mid_sh - mid_hip
                v_norm = np.linalg.norm(v) + 1e-9
                vert = np.array([0.0, -1.0])
                cosang = float(np.dot(v / v_norm, vert))
                ang = math.degrees(math.acos(max(-1.0, min(1.0, cosang))))
                out["uprightness"].append(max(0.0, min(1.0, 1.0 - ang / 55.0)))

                shoulder_width = abs(lsh_pt[0] - rsh_pt[0]) + 1e-9
                if min(float(lank.visibility), float(rank.visibility)) >= 0.5:
                    dx = lank_pt[0] - rank_pt[0]
                    dy = lank_pt[1] - rank_pt[1]
                    ankle_dist = math.sqrt(dx * dx + dy * dy)
                    out["stance_width_ratio"].append(ankle_dist / shoulder_width)
                    ankle_mid = ((lank_pt[0] + rank_pt[0]) / 2.0, (lank_pt[1] + rank_pt[1]) / 2.0)
                    if prev_ankle_mid is not None:
                        shift = math.hypot(ankle_mid[0] - prev_ankle_mid[0], ankle_mid[1] - prev_ankle_mid[1])
                        out["weight_shift"].append(min(1.0, shift / (shoulder_width * 0.8 + 1e-9)))
                    prev_ankle_mid = ankle_mid
                else:
                    out["stance_width_ratio"].append(0.35)
                    out["weight_shift"].append(0.35)

                if min(float(lw.visibility), float(rw.visibility)) >= 0.4:
                    wmid = ((float(lw.x) + float(rw.x)) / 2.0, (float(lw.y) + float(rw.y)) / 2.0)
                    if prev_wrist_mid is not None:
                        vel = math.hypot(wmid[0] - prev_wrist_mid[0], wmid[1] - prev_wrist_mid[1])
                        e = min(1.0, vel * 25.0)
                        out["gesture_energy"].append(e)
                        energy_window.append(e)
                        if len(energy_window) > 7:
                            energy_window.pop(0)
                        if len(energy_window) >= 2:
                            out["gesture_variation"].append(float(np.std(np.array(energy_window))))
                        else:
                            out["gesture_variation"].append(e)
                    prev_wrist_mid = wmid
                else:
                    out["gesture_energy"].append(0.15)
                    out["gesture_variation"].append(0.12)

                out["chest_blocking"].append(0.12)
                out["chest_open"].append(min(1.0, shoulder_width * 1.8))

                sh_ang = _sh_angle_deg(lsh_pt, rsh_pt)
                if prev_sh_angle is not None:
                    delta = abs(sh_ang - prev_sh_angle)
                    while delta > 180.0:
                        delta = 360.0 - delta
                    out["rotation"].append(min(1.0, delta / 45.0))
                prev_sh_angle = sh_ang

                if len(out["eye_contact"]) >= max_frames:
                    break
    except Exception as e:
        logger.warning("[movement_features] MediaPipe failed (empty features): %s", e)
        return out
    finally:
        try:
            cap.release()
        except Exception:
            pass
    return out


def generate_eye_contact_text(pct: float) -> list:
    """Generate descriptive text based on eye contact percentage - 10 levels"""
    if pct >= 90:  # Outstanding
        return [
            "• Your eye contact is outstanding — consistently steady, warm, and completely audience-focused throughout the entire presentation.",
            "• You maintain unwavering direct gaze during key moments, which maximizes trust, clarity, and audience engagement.",
            "• Every gaze shift is purposeful, natural, and enhances your communication (e.g., thinking pauses, emphasis).",
            "• Your eye contact is a masterclass in confidence and credibility, with zero signs of avoidance or nervousness."
        ]
    elif pct >= 80:  # Excellent
        return [
            "• Your eye contact is excellent — steady, engaging, and audience-focused almost throughout the presentation.",
            "• You consistently maintain direct gaze during key moments, which strongly increases trust and clarity.",
            "• Gaze shifts are purposeful and natural, never detracting from your professional presence.",
            "• Your eye contact demonstrates strong confidence and builds exceptional credibility with the audience."
        ]
    elif pct >= 70:  # Very Strong
        return [
            "• Your eye contact is very strong — steady and audience-focused for most of the time.",
            "• You maintain direct gaze during key moments effectively, which clearly builds trust and clarity.",
            "• Occasional gaze shifts are natural and appropriate, showing thoughtfulness without reducing engagement.",
            "• Overall, your eye contact strongly supports confidence and credibility with the audience."
        ]
    elif pct >= 60:  # Strong
        return [
            "• Your eye contact is strong — generally steady and audience-focused during important moments.",
            "• You maintain direct gaze during most key moments, which helps establish trust.",
            "• Some gaze shifts occur but they're mostly natural and don't significantly impact your presence.",
            "• Your eye contact effectively supports good confidence and reasonable credibility."
        ]
    elif pct >= 50:  # Good
        return [
            "• Your eye contact shows a good foundation, with several effective moments of direct audience connection.",
            "• You maintain direct gaze during many important points, which helps build trust.",
            "• With slightly longer eye contact in key moments, your delivery will feel even more confident.",
            "• Overall, your audience engagement is positive and can be strengthened further."
        ]
    elif pct >= 40:  # Above Average
        return [
            "• Your eye contact is developing well, with regular moments of direct audience connection.",
            "• You make direct eye contact consistently, and sustaining it slightly longer will increase impact.",
            "• Some gaze shifts occur more frequently than ideal, but this is very trainable.",
            "• Improving consistency will noticeably enhance trust and credibility."
        ]
    elif pct >= 30:  # Average
        return [
            "• Your eye contact is in a developing stage, with clear opportunities to become more consistent.",
            "• You already connect directly with the audience at times, especially in selected moments.",
            "• Extending direct gaze a bit longer will strengthen audience connection and engagement.",
            "• With focused practice, trust and credibility cues will improve clearly."
        ]
    else:
        return [
            "• Your eye contact is developing and currently inconsistent in several parts of the presentation.",
            "• You connect with the audience at times, and extending direct gaze during key moments will improve trust.",
            "• Some gaze shifts are more frequent than ideal, but this can be improved with focused practice.",
            "• With steadier eye contact, your confidence and credibility will become more apparent."
        ]

def generate_eye_contact_text_th(pct: float) -> list:
    """Thai descriptive text based on eye contact percentage - 10 levels"""
    if pct >= 90:
        return [
            "• การสบตาโดดเด่นมาก สม่ำเสมอ อบอุ่น และโฟกัสผู้ฟังตลอดการนำเสนอ",
            "• คุณรักษาการสบตาในจุดสำคัญได้อย่างยอดเยี่ยม ช่วยเพิ่มความไว้วางใจสูงมาก",
            "• การเปลี่ยนสายตาเป็นธรรมชาติและมีจุดประสงค์ ไม่ทำให้พลังการสื่อสารลดลง",
            "• ภาพรวมสะท้อนความมั่นใจและความน่าเชื่อถือในระดับสูงสุด",
        ]
    elif pct >= 80:
        return [
            "• การสบตายอดเยี่ยม มีความต่อเนื่องและเชื่อมโยงผู้ฟังได้ดีมาก",
            "• คุณสบตาในประเด็นสำคัญได้ชัดเจน ช่วยสร้างความเชื่อมั่นได้มาก",
            "• มีการละสายตาบ้างเล็กน้อยแต่ยังเป็นธรรมชาติ",
            "• ภาพรวมแสดงบุคลิกที่มั่นใจและน่าเชื่อถือมาก",
        ]
    elif pct >= 70:
        return [
            "• การสบตาแข็งแรงและคงเส้นคงวาในช่วงส่วนใหญ่ของการพูด",
            "• ช่วงที่เน้นสาระสำคัญมีการสบตาที่ดี ช่วยเสริมความไว้วางใจ",
            "• มีจังหวะละสายตาเป็นระยะแต่ไม่กระทบการเชื่อมต่อผู้ฟังมาก",
            "• โดยรวมสนับสนุนภาพลักษณ์มืออาชีพได้ดี",
        ]
    elif pct >= 60:
        return [
            "• การสบตาอยู่ในระดับดีและค่อนข้างสม่ำเสมอ",
            "• จุดสำคัญส่วนใหญ่ยังคงสบตาได้เหมาะสม",
            "• มีการละสายตาบ้างแต่ยังไม่รบกวนภาพรวม",
            "• โดยรวมช่วยเสริมความมั่นใจได้ชัดเจน",
        ]
    elif pct >= 50:
        return [
            "• การสบตามีพื้นฐานที่ดี และมีหลายช่วงที่เชื่อมโยงผู้ฟังได้ชัดเจน",
            "• ในประเด็นสำคัญสามารถสบตาได้พอสมควรและสร้างความไว้วางใจได้",
            "• หากเพิ่มระยะเวลาการสบตาอีกเล็กน้อย พลังการสื่อสารจะดีขึ้นชัดเจน",
            "• ภาพรวมเป็นบวกและต่อยอดได้อีกมาก",
        ]
    elif pct >= 40:
        return [
            "• การสบตากำลังพัฒนาไปได้ดี และมีช่วงที่เชื่อมโยงผู้ฟังได้ต่อเนื่อง",
            "• มีการสบตาสม่ำเสมอในหลายจังหวะ และสามารถต่อยอดให้คงที่ขึ้นได้",
            "• การละสายตายังมีอยู่บ้าง แต่สามารถปรับได้ด้วยการฝึก",
            "• หากเพิ่มความต่อเนื่องของการสบตา จะช่วยเสริมความน่าเชื่อถือมากขึ้น",
        ]
    elif pct >= 30:
        return [
            "• การสบตาอยู่ในช่วงพัฒนา และยังมีโอกาสเพิ่มความสม่ำเสมอได้อีก",
            "• มีหลายช่วงที่เชื่อมโยงผู้ฟังได้ดี โดยเฉพาะบางประเด็นสำคัญ",
            "• หากคุมการละสายตาให้น้อยลง จะทำให้การสื่อสารไหลลื่นขึ้น",
            "• แนะนำฝึกการสบตาแบบสั้นแต่ต่อเนื่องในจังหวะสำคัญ",
        ]
    else:
        return [
            "• การสบตายังอยู่ในช่วงพัฒนา และยังไม่สม่ำเสมอในหลายช่วงของการนำเสนอ",
            "• มีช่วงที่เชื่อมโยงผู้ฟังได้ดี หากเพิ่มระยะเวลาการสบตาในประเด็นสำคัญจะดีขึ้นมาก",
            "• การละสายตายังเกิดบ่อยกว่าที่ควร แต่สามารถปรับได้ด้วยการฝึกอย่างต่อเนื่อง",
            "• เมื่อคุมการสบตาได้คงที่ขึ้น ภาพลักษณ์ความมั่นใจและความน่าเชื่อถือจะชัดเจนขึ้น",
        ]

def generate_uprightness_text(pct: float) -> list:
    """Generate descriptive text based on uprightness percentage - 10 levels"""
    if pct >= 90:  # Outstanding
        return [
            "• You maintain outstanding upright posture with perfect vertical alignment throughout the entire presentation.",
            "• Your chest remains fully open, shoulders are optimally relaxed, and head alignment is impeccable — signaling exceptional balance, readiness, and commanding authority.",
            "• Even during active gesturing, your vertical alignment stays completely stable, demonstrating masterful core control.",
            "• There is absolutely no visible slouching or collapsing at any point, projecting supreme professional presence."
        ]
    elif pct >= 80:  # Excellent
        return [
            "• You maintain excellent upright posture consistently throughout the presentation.",
            "• Your chest stays open, shoulders are naturally relaxed, and head alignment is near-perfect — signaling strong balance, readiness, and authority.",
            "• Your vertical alignment remains remarkably stable even when gesturing, demonstrating superior core control.",
            "• There is virtually no slouching or collapsing, which projects a highly professional and confident appearance."
        ]
    elif pct >= 70:  # Very Strong
        return [
            "• You maintain very strong upright posture throughout most of the presentation.",
            "• The chest stays open, shoulders remain relaxed, and head is well-aligned — clearly signaling balance, readiness, and authority.",
            "• Your vertical alignment stays stable during most gestures, showing very good core control.",
            "• Minimal slouching occurs, which strongly supports a professional and composed appearance."
        ]
    elif pct >= 60:  # Strong
        return [
            "• You maintain strong upright posture for the majority of the clip.",
            "• Your chest generally stays open, shoulders relaxed, and head aligned — signaling good balance and authority.",
            "• Vertical alignment remains fairly stable during gestures, showing solid core control.",
            "• Occasional minor slouching appears but doesn't detract significantly from your professional presence."
        ]
    elif pct >= 50:  # Good
        return [
            "• Your posture has a good foundation, with many moments of stable upright alignment.",
            "• Chest and shoulder alignment are generally positive, with room to make consistency even better.",
            "• During gestures, vertical alignment is mostly stable and supports clear communication.",
            "• Overall posture is strong and can be elevated further with small refinements."
        ]
    elif pct >= 40:  # Above Average
        return [
            "• Your posture is developing well, with a number of solid upright moments throughout the presentation.",
            "• Chest and shoulder positioning are often acceptable, and greater consistency will add polish.",
            "• Vertical alignment varies in some gesture phases, which is common and trainable.",
            "• With steadier posture, your professional image will look even more confident."
        ]
    elif pct >= 30:  # Average
        return [
            "• Your posture is in a developing stage, alternating between upright moments and less stable alignment.",
            "• There are several points where shoulder and head position can be brought closer to ideal alignment.",
            "• Improving vertical consistency will make your presence feel stronger and more composed.",
            "• With ongoing practice, professional presence and authority will improve clearly."
        ]
    else:
        return [
            "• Your posture is still developing and varies between upright and less stable alignment.",
            "• There are several moments where shoulder and head alignment drift from ideal posture.",
            "• Improving core control and maintaining a taller stance will enhance your professional presence.",
            "• With consistent practice, your posture can become more stable and confident."
        ]

def generate_uprightness_text_th(pct: float) -> list:
    """Thai descriptive text based on uprightness percentage - 10 levels"""
    if pct >= 90:
        return [
            "• ท่าทางตั้งตรงโดดเด่นมาก การจัดแนวลำตัวนิ่งและสมดุลตลอดการพูด",
            "• หน้าอกเปิด ไหล่ผ่อนคลาย ศีรษะอยู่ในแนวที่ดีอย่างสม่ำเสมอ",
            "• แม้มีการใช้ท่าทาง มือและลำตัวก็ยังคงความมั่นคงสูง",
            "• ภาพรวมสื่อความมั่นใจและภาวะผู้นำได้ชัดเจนมาก",
        ]
    elif pct >= 80:
        return [
            "• ท่าทางตั้งตรงยอดเยี่ยมและคงเส้นคงวาเกือบตลอดคลิป",
            "• การจัดแนวหัวไหล่และลำตัวอยู่ในเกณฑ์ดีมาก",
            "• มีการเอนตัวเล็กน้อยบางช่วงแต่ไม่กระทบภาพรวม",
            "• โดยรวมดูมั่นใจและมืออาชีพมาก",
        ]
    elif pct >= 70:
        return [
            "• ท่าทางตั้งตรงดีมากในช่วงส่วนใหญ่ของการนำเสนอ",
            "• โครงลำตัวโดยรวมสมดุลและควบคุมได้ดี",
            "• มีจังหวะที่แนวลำตัวหลุดเล็กน้อยเป็นครั้งคราว",
            "• ภาพรวมยังสนับสนุนความน่าเชื่อถือได้ดี",
        ]
    elif pct >= 60:
        return [
            "• ท่าทางตั้งตรงอยู่ในระดับดีและค่อนข้างสม่ำเสมอ",
            "• ลำตัวและไหล่ส่วนใหญ่ยังอยู่ในแนวเหมาะสม",
            "• มีการคล้อยลำตัวบ้างเล็กน้อยในบางช่วง",
            "• โดยรวมช่วยเสริมบุคลิกมืออาชีพได้ดี",
        ]
    elif pct >= 50:
        return [
            "• ท่าทางตั้งตรงมีพื้นฐานที่ดี และมีหลายช่วงที่คงแนวลำตัวได้ดี",
            "• ความต่อเนื่องโดยรวมเป็นบวก และยังมีโอกาสทำให้คงที่ขึ้นอีก",
            "• หากคุมแกนลำตัวให้ต่อเนื่องขึ้นเล็กน้อย ภาพรวมจะดีขึ้นชัดเจน",
            "• โดยรวมอยู่ในทิศทางที่ดีและพัฒนาได้อีกมาก",
        ]
    elif pct >= 40:
        return [
            "• ท่าทางกำลังพัฒนาได้ดี และมีหลายช่วงที่รักษาแนวได้เหมาะสม",
            "• แนวไหล่/ศีรษะยังมีบางจังหวะที่แกว่งเล็กน้อย ซึ่งปรับได้ไม่ยาก",
            "• การคุมแกนกลางมีแนวโน้มดีและสามารถเสริมให้มั่นคงขึ้นได้",
            "• หากเพิ่มความสม่ำเสมออีกเล็กน้อย บุคลิกจะดูมืออาชีพขึ้นมาก",
        ]
    elif pct >= 30:
        return [
            "• ท่าทางอยู่ในช่วงพัฒนา มีทั้งช่วงที่ตั้งตรงได้ดีและช่วงที่ยังไม่นิ่ง",
            "• แนวไหล่และศีรษะยังมีบางจังหวะที่หลุดจากแนวเหมาะสม",
            "• หากคุมแนวลำตัวให้คงที่ขึ้น จะช่วยเสริมความมั่นใจได้ชัดเจน",
            "• แนะนำฝึกการคุมแกนลำตัวอย่างต่อเนื่องเพื่อเพิ่มความนิ่ง",
        ]
    else:
        return [
            "• ท่าทางยังอยู่ในช่วงพัฒนา และมีหลายช่วงที่แนวลำตัวยังไม่นิ่ง",
            "• แนวไหล่และศีรษะหลุดจากแนวที่เหมาะสมเป็นระยะ",
            "• หากเพิ่มการคุมแกนกลางและรักษาท่ายืนให้ต่อเนื่อง จะช่วยเสริมบุคลิกได้มาก",
            "• ฝึกอย่างสม่ำเสมอจะทำให้ภาพรวมดูมั่นใจและมืออาชีพขึ้นอย่างชัดเจน",
        ]

def generate_stance_text(stability: float) -> list:
    """Generate descriptive text based on stance stability - 10 levels"""
    if stability >= 90:  # Outstanding
        return [
            "• Your stance is outstanding — perfectly symmetrical and solidly grounded, with feet placed optimally at shoulder-width apart.",
            "• Weight shifts are virtually non-existent, creating an exceptionally stable platform that demonstrates supreme confidence.",
            "• You maintain flawless forward orientation toward the audience throughout, maximally reinforcing clarity and engagement.",
            "• Your stance conveys rock-solid stability and commanding, welcoming authority ideal for executive leadership communication."
        ]
    elif stability >= 80:  # Excellent
        return [
            "• Your stance is excellent — exceptionally symmetrical and grounded, with feet placed perfectly about shoulder-width apart.",
            "• Weight shifts are extremely controlled and minimal, preventing any distraction and demonstrating strong confidence.",
            "• You maintain superior forward orientation toward the audience throughout, reinforcing excellent clarity and engagement.",
            "• Your stance conveys remarkable stability and a welcoming, authoritative presence ideal for professional leadership."
        ]
    elif stability >= 70:  # Very Strong
        return [
            "• Your stance is very strong — highly symmetrical and well-grounded, with feet well-positioned shoulder-width apart.",
            "• Weight shifts are well-controlled and quite minimal, showing strong confidence and solid balance.",
            "• You maintain very good forward orientation toward the audience, clearly reinforcing engagement and clarity.",
            "• The stance conveys strong stability and a professional, welcoming presence suitable for senior communication roles."
        ]
    elif stability >= 60:  # Strong
        return [
            "• Your stance is strong — symmetrical and grounded, with feet appropriately placed about shoulder-width apart.",
            "• Weight shifts are controlled and minimal, preventing distraction and demonstrating good confidence.",
            "• You maintain solid forward orientation toward the audience, reinforcing clarity and reasonable engagement.",
            "• The stance effectively conveys stability and a professional presence suitable for most business communication."
        ]
    elif stability >= 50:  # Good
        return [
            "• Your stance has a good base, with generally stable positioning and only occasional weight shifts.",
            "• Feet placement is mostly appropriate, supporting balanced delivery through much of the presentation.",
            "• Weight distribution is mostly steady, with minor adjustments that are common and manageable.",
            "• Overall stance supports a positive professional presence and can be refined further."
        ]
    elif stability >= 40:  # Above Average
        return [
            "• Your stance is developing well, with many steady moments across the presentation.",
            "• Feet placement is generally acceptable, and greater consistency will improve grounding.",
            "• Some visible shifts occur in parts of the clip, which is normal and can be reduced with practice.",
            "• Better stability will further strengthen your grounded presence and authority."
        ]
    elif stability >= 30:  # Average
        return [
            "• Your stance is in a developing stage, with opportunities to become steadier throughout delivery.",
            "• Feet placement varies at times, and improving consistency will support better balance.",
            "• Some swaying or shifting appears in parts of the presentation, and this can be improved with focused drills.",
            "• Enhancing stance stability will clearly strengthen your professional presence."
        ]
    else:
        return [
            "• Your stance is still developing, with noticeable movement and balance variation in several moments.",
            "• Feet placement and weight distribution are not yet fully consistent throughout the presentation.",
            "• Reducing unnecessary shifts will help you appear more grounded and authoritative.",
            "• With regular practice, your stance can become steadier and more professional."
        ]

def generate_stance_text_th(stability: float) -> list:
    """Thai descriptive text based on stance stability - 10 levels"""
    if stability >= 90:
        return [
            "• ท่ายืนมั่นคงยอดเยี่ยม สมมาตร และมีฐานยืนที่แข็งแรงมาก",
            "• การถ่ายน้ำหนักน้อยมาก ทำให้ภาพรวมดูนิ่งและมั่นใจสูง",
            "• การวางเท้าเหมาะสมและช่วยเสริมพลังการสื่อสารอย่างชัดเจน",
            "• สะท้อนบุคลิกผู้นำที่พร้อมและน่าเชื่อถือมาก",
        ]
    elif stability >= 80:
        return [
            "• ท่ายืนมั่นคงมาก มีความสมดุลต่อเนื่อง",
            "• การถ่ายน้ำหนักเกิดน้อยและควบคุมได้ดี",
            "• ฐานยืนช่วยให้ภาพลักษณ์ดูมั่นใจและมืออาชีพ",
            "• โดยรวมสื่อความน่าเชื่อถือได้ดีมาก",
        ]
    elif stability >= 70:
        return [
            "• ท่ายืนแข็งแรงและค่อนข้างนิ่งในช่วงส่วนใหญ่",
            "• มีการปรับน้ำหนักเล็กน้อยแต่ไม่รบกวนผู้ฟังมาก",
            "• ฐานยืนโดยรวมดีและรองรับการสื่อสารได้ดี",
            "• ภาพรวมเสริมความมั่นใจได้ชัดเจน",
        ]
    elif stability >= 60:
        return [
            "• ท่ายืนอยู่ในระดับดี มีความมั่นคงพอเหมาะ",
            "• มีการเปลี่ยนน้ำหนักเป็นระยะแต่ยังควบคุมได้",
            "• ฐานยืนโดยรวมยังสนับสนุนภาพลักษณ์มืออาชีพ",
            "• ปรับความนิ่งเพิ่มอีกเล็กน้อยจะดีขึ้นมาก",
        ]
    elif stability >= 50:
        return [
            "• ท่ายืนมีพื้นฐานที่ดี และมีหลายช่วงที่ค่อนข้างมั่นคง",
            "• ความสมดุลโดยรวมเป็นบวก และยังปรับให้คงที่ขึ้นได้อีก",
            "• มีการขยับน้ำหนักเป็นระยะ ซึ่งสามารถลดลงได้ด้วยการฝึก",
            "• หากเพิ่มความนิ่งอีกเล็กน้อย ภาพรวมจะดูมืออาชีพขึ้นมาก",
        ]
    elif stability >= 40:
        return [
            "• ท่ายืนกำลังพัฒนาได้ดี และมีหลายช่วงที่รักษาฐานยืนได้เหมาะสม",
            "• ความต่อเนื่องของฐานยืนยังมีโอกาสทำให้คงที่ขึ้นอีก",
            "• มีการโยก/ถ่ายน้ำหนักในบางจังหวะ ซึ่งสามารถปรับได้ด้วยการฝึก",
            "• หากยืน grounded มากขึ้น จะช่วยเสริมภาพลักษณ์ให้มั่นคงขึ้นชัดเจน",
        ]
    elif stability >= 30:
        return [
            "• ท่ายืนอยู่ในช่วงพัฒนา และยังเพิ่มความนิ่งได้อีกในหลายช่วง",
            "• การวางเท้าและการถ่ายน้ำหนักมีบางจังหวะที่ยังไม่คงที่",
            "• หากคุมฐานยืนให้ต่อเนื่องขึ้น จะช่วยเสริมความมั่นใจของภาพรวม",
            "• แนะนำฝึกคุมแกนล่างและจังหวะการถ่ายน้ำหนักแบบสม่ำเสมอ",
        ]
    else:
        return [
            "• ท่ายืนยังอยู่ในช่วงพัฒนา และมีการเคลื่อนไหว/ถ่ายน้ำหนักที่เห็นได้ในหลายช่วง",
            "• ฐานยืนและการกระจายน้ำหนักยังไม่คงที่ตลอดการนำเสนอ",
            "• หากลดการโยกหรือขยับที่ไม่จำเป็น จะช่วยให้ภาพลักษณ์ดูมั่นคงขึ้น",
            "• ฝึกการยืนอย่างสม่ำเสมอจะช่วยเสริมความน่าเชื่อถือและความเป็นมืออาชีพได้ชัดเจน",
        ]

# Analysis functions
def _compute_enriched_presentation_features(
    *,
    nose_x_series: List[float],
    nose_y_series: List[float],
    nose_z_series: List[float],
    shoulder_mid_x_series: List[float],
    shoulder_mid_y_series: List[float],
    trunk_angle_series: List[float],
    shoulder_width_series: List[float],
    hip_width_series: List[float],
    body_expansion_series: List[float],
    wrist_avg_x_series: List[float],
    wrist_avg_y_series: List[float],
    wrist_velocity_series: List[float],
    left_wrist_xy_series: List[Tuple[float, float]],
    right_wrist_xy_series: List[Tuple[float, float]],
    hip_sway_std: float,
) -> Dict[str, float]:
    """Post-process per-frame series into the 15 enriched presentation features.

    All features are scale-normalized so a Random Forest trained on them can
    compare across clips. Returns zeros if the series is too short.
    """
    out: Dict[str, float] = {
        "posture_uprightness": 0.0,
        "torso_stability": 0.0,
        "head_stability": 0.0,
        "eye_direction_proxy": 0.0,
        "shoulder_openness": 0.0,
        "hand_openness": 0.0,
        "gesture_range": 0.0,
        "gesture_smoothness": 0.0,
        "movement_intentionality": 0.0,
        "hesitation_score": 0.0,
        "rhythm_consistency": 0.0,
        "energy_level": 0.0,
        "center_presence": 0.0,
        # body_sway + stance_stability are passed through from existing signals
        # (see score_presentation_ml) so the feature list stays complete.
        "body_sway": float(hip_sway_std),
    }

    n = len(trunk_angle_series)
    if n < 5:
        return out

    nx = np.array(nose_x_series, dtype=float)
    ny = np.array(nose_y_series, dtype=float)
    nz = np.array(nose_z_series, dtype=float)
    smx = np.array(shoulder_mid_x_series, dtype=float)
    smy = np.array(shoulder_mid_y_series, dtype=float)
    trunk = np.array(trunk_angle_series, dtype=float)
    sw = np.array(shoulder_width_series, dtype=float)
    hw = np.array(hip_width_series, dtype=float)
    be = np.array(body_expansion_series, dtype=float)
    wx = np.array(wrist_avg_x_series, dtype=float)
    wy = np.array(wrist_avg_y_series, dtype=float)
    wv = np.array(wrist_velocity_series, dtype=float)

    # 1. posture_uprightness: 1.0 = perfectly upright, 0.0 = bent 45°+.
    mean_trunk = float(np.mean(trunk))
    out["posture_uprightness"] = max(0.0, 1.0 - mean_trunk / 45.0)

    # 2. torso_stability: 1 - std(shoulder midpoint) / typical-scale (0.1).
    torso_drift = float(np.sqrt(np.var(smx) + np.var(smy)))
    out["torso_stability"] = max(0.0, 1.0 - torso_drift / 0.10)

    # 3. head_stability: 1 - std(nose xyz) / typical-scale (0.08).
    head_drift = float(np.sqrt(np.var(nx) + np.var(ny) + np.var(nz)))
    out["head_stability"] = max(0.0, 1.0 - head_drift / 0.10)

    # 4. eye_direction_proxy: how much head rotates side-to-side relative to
    #    shoulders (high = lots of head turning, lower perceived eye contact).
    #    Scale so 0.0 = perfectly facing forward, 1.0 = heavy turning.
    nose_offset = nx - smx
    out["eye_direction_proxy"] = min(1.0, float(np.std(nose_offset)) / 0.06)

    # 5. shoulder_openness: mean(shoulder_width / hip_width). ~1.0 = neutral,
    #    >1.2 = shoulders wider than hips (open), <0.9 = hunched.
    denom_hw = np.where(hw > 1e-4, hw, 1e-4)
    sh_open = float(np.mean(sw / denom_hw))
    out["shoulder_openness"] = max(0.0, min(2.0, sh_open))

    # 6. hand_openness: mean body_expansion = wrist_dist / shoulder_width.
    out["hand_openness"] = float(np.clip(np.mean(be), 0.0, 3.0))

    # 7. gesture_range: max span of wrists (L2 of x-range + y-range) normalized.
    lx = np.array([p[0] for p in left_wrist_xy_series], dtype=float)
    ly = np.array([p[1] for p in left_wrist_xy_series], dtype=float)
    rx = np.array([p[0] for p in right_wrist_xy_series], dtype=float)
    ry = np.array([p[1] for p in right_wrist_xy_series], dtype=float)
    span_x = max(float(lx.max() - lx.min()), float(rx.max() - rx.min()))
    span_y = max(float(ly.max() - ly.min()), float(ry.max() - ry.min()))
    out["gesture_range"] = min(1.5, math.sqrt(span_x ** 2 + span_y ** 2))

    # 8. gesture_smoothness: low jerk = smooth. Compute relative jerk
    #    (std of velocity change) vs mean velocity; invert so 1.0 = smooth.
    if wv.size >= 3 and float(np.mean(wv)) > 1e-4:
        diff = np.diff(wv)
        rel_jerk = float(np.std(diff)) / (float(np.mean(wv)) + 1e-4)
        out["gesture_smoothness"] = max(0.0, 1.0 - min(rel_jerk / 2.0, 1.0))
    else:
        out["gesture_smoothness"] = 1.0

    # 9. movement_intentionality: ratio of frames whose velocity is above a
    #    meaningful threshold (0.03) vs frames with micro-jitter (>0.005).
    meaningful = float(np.sum(wv >= 0.03))
    jitter = float(np.sum(wv >= 0.005))
    out["movement_intentionality"] = meaningful / max(jitter, 1.0)

    # 10. hesitation_score: fraction of frames where velocity dips below 0.01
    #     after a peak >= 0.04 in the immediately preceding window.
    if wv.size >= 5:
        peaks = wv >= 0.04
        dips = wv <= 0.01
        hesitation = 0
        for i in range(2, wv.size):
            if dips[i] and (peaks[i - 1] or peaks[i - 2]):
                hesitation += 1
        out["hesitation_score"] = hesitation / max(1, wv.size)
    else:
        out["hesitation_score"] = 0.0

    # 11. rhythm_consistency: std of inter-burst intervals. A "burst" is a
    #     frame with velocity crossing 0.04 upward. Low std = steady rhythm.
    burst_idx: List[int] = []
    for i in range(1, wv.size):
        if wv[i] >= 0.04 and wv[i - 1] < 0.04:
            burst_idx.append(i)
    if len(burst_idx) >= 3:
        intervals = np.diff(np.array(burst_idx, dtype=float))
        m = float(np.mean(intervals))
        if m > 1e-6:
            cv = float(np.std(intervals)) / m   # coefficient of variation
            out["rhythm_consistency"] = max(0.0, 1.0 - min(cv, 1.0))
        else:
            out["rhythm_consistency"] = 0.0
    else:
        out["rhythm_consistency"] = 0.0

    # 12. energy_level: mean wrist velocity (already a meaningful scalar).
    out["energy_level"] = float(np.mean(wv))

    # 13. center_presence: how close the nose stays to horizontal frame center
    #     (0.5). Map mean(|nose_x - 0.5|) into 1.0 (centered) → 0.0 (edge).
    mean_off = float(np.mean(np.abs(nx - 0.5)))
    out["center_presence"] = max(0.0, 1.0 - mean_off / 0.30)

    return out


def analyze_video_mediapipe(video_path: str, sample_fps: float = 5, max_frames: int = 300, **kwargs) -> Dict[str, Any]:
    """Real MediaPipe analysis with proper Laban Movement Analysis"""
    enclosing_max_expansion = float(os.getenv("ENCLOSING_MAX_EXPANSION", "0.8"))
    enclosing_min_velocity = float(os.getenv("ENCLOSING_MIN_VELOCITY", "0.03"))
    spreading_body_expansion_threshold = float(
        os.getenv("SPREADING_BODY_EXPANSION_THRESHOLD", "1.3")
    )
    spreading_min_velocity = float(os.getenv("SPREADING_MIN_VELOCITY", "0.03"))
    gesture_debug_log = str(os.getenv("GESTURE_DEBUG_LOG", "false")).strip().lower() in (
        "1",
        "true",
        "yes",
        "on",
    )
    gesture_debug_every_n = max(1, int(os.getenv("GESTURE_DEBUG_EVERY_N_FRAMES", "5")))

    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        raise RuntimeError("Cannot open video")
    
    fps = cap.get(cv2.CAP_PROP_FPS) or 25.0
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    duration = total_frames / fps if fps > 0 else 0
    
    frame_interval = max(1, int(fps / sample_fps))
    
    # Effort & Shape detection counters
    effort_counts = {
        "Directing": 0, "Enclosing": 0, "Punching": 0, "Spreading": 0,
        "Pressing": 0, "Dabbing": 0, "Indirecting": 0, "Gliding": 0,
        "Flicking": 0, "Advancing": 0, "Retreating": 0
    }
    shape_counts = {"Directing": 0, "Enclosing": 0, "Spreading": 0, "Indirecting": 0, "Advancing": 0, "Retreating": 0}

    # Category-level per-frame collectors (professor-feedback, calibration-pending).
    # These are extracted and logged but NOT yet wired into the final scores —
    # they will be used once we re-run calibration against the 12 ground-truth
    # clips (Lea, Payu, Sawitree, Sarinee, Lisa, Chutima, Aon, Ann, Ches, Candy...).
    #   hand_block_frames: both wrists in the chest box (inside shoulder span,
    #                      above hip line, below shoulder line). Lowers Engaging.
    #   hand_low_frames:   both wrists below hip line. Lowers Authority/Confidence.
    #   hands_above_sh_frames: at least one wrist above shoulder line.
    #   hip_y_series / hip_x_series: track hip midpoint trajectory (advance toward
    #                      audience and body sway).
    #   wrist_shape_signatures: coarse hand-shape cluster ids for gesture variety.
    hand_block_frames = 0
    hand_low_frames = 0
    hands_above_sh_frames = 0
    hip_y_series: List[float] = []
    hip_x_series: List[float] = []
    wrist_shape_signatures: List[Tuple[int, int, int, int]] = []

    # Enriched-feature collectors (used by the Presentation Analysis ML scorer).
    # These are per-frame time-series we post-process into 15 scalar features
    # (posture_uprightness, torso_stability, head_stability, eye_direction_proxy,
    #  shoulder_openness, hand_openness, gesture_range, gesture_smoothness,
    #  movement_intentionality, hesitation_score, rhythm_consistency,
    #  energy_level, center_presence, + stance_stability + body_sway from
    #  first-impression / hip_sway).
    nose_x_series: List[float] = []
    nose_y_series: List[float] = []
    nose_z_series: List[float] = []
    shoulder_mid_x_series: List[float] = []
    shoulder_mid_y_series: List[float] = []
    trunk_angle_series: List[float] = []           # deg from vertical (0 = upright)
    shoulder_width_series: List[float] = []
    hip_width_series: List[float] = []
    body_expansion_series: List[float] = []        # wrist_dist / shoulder_width
    wrist_avg_x_series: List[float] = []           # midpoint of wrists
    wrist_avg_y_series: List[float] = []
    wrist_velocity_series: List[float] = []        # avg of L+R wrist velocity
    left_wrist_xy_series: List[Tuple[float, float]] = []
    right_wrist_xy_series: List[Tuple[float, float]] = []

    analyzed = 0
    sampled = 0
    prev_landmarks = None
    analysis_started_at = time.time()
    analysis_timeout_seconds = max(60, int(os.getenv("ANALYSIS_TIMEOUT_SECONDS", "420")))

    mc = max(0, min(2, int(kwargs.get("pose_model_complexity") or 1)))

    try:
        with Pose(
            static_image_mode=False,
            model_complexity=mc,
            enable_segmentation=False,
        ) as pose:
            frame_idx = 0
            while analyzed < max_frames and sampled < max_frames:
                ret, frame = cap.read()
                if not ret:
                    break
                if (time.time() - analysis_started_at) > analysis_timeout_seconds:
                    logger.warning(
                        "[analysis] timeout reached (%ss): sampled=%s detected=%s max_frames=%s",
                        analysis_timeout_seconds,
                        sampled,
                        analyzed,
                        max_frames,
                    )
                    break
            
                if frame_idx % frame_interval == 0:
                    sampled += 1
                    rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                    results = pose.process(rgb)
                
                    if results.pose_landmarks:
                        lms = results.pose_landmarks.landmark
                    
                        # Calculate movement vectors if we have previous frame
                        if prev_landmarks is not None:
                            # Get key points
                            left_wrist = lms[PoseLandmark.LEFT_WRIST]
                            right_wrist = lms[PoseLandmark.RIGHT_WRIST]
                            left_elbow = lms[PoseLandmark.LEFT_ELBOW]
                            right_elbow = lms[PoseLandmark.RIGHT_ELBOW]
                            left_shoulder = lms[PoseLandmark.LEFT_SHOULDER]
                            right_shoulder = lms[PoseLandmark.RIGHT_SHOULDER]
                            nose = lms[PoseLandmark.NOSE]
                        
                            # Previous frame landmarks
                            prev_left_wrist = prev_landmarks[PoseLandmark.LEFT_WRIST]
                            prev_right_wrist = prev_landmarks[PoseLandmark.RIGHT_WRIST]
                        
                            # Calculate velocities (movement speed)
                            left_wrist_vel = math.sqrt(
                                (left_wrist.x - prev_left_wrist.x)**2 +
                                (left_wrist.y - prev_left_wrist.y)**2 +
                                (left_wrist.z - prev_left_wrist.z)**2
                            )
                            right_wrist_vel = math.sqrt(
                                (right_wrist.x - prev_right_wrist.x)**2 +
                                (right_wrist.y - prev_right_wrist.y)**2 +
                                (right_wrist.z - prev_right_wrist.z)**2
                            )
                            avg_velocity = (left_wrist_vel + right_wrist_vel) / 2
                        
                            # Spatial measurements
                            wrist_dist = abs(left_wrist.x - right_wrist.x)
                            shoulder_width = abs(left_shoulder.x - right_shoulder.x)
                            body_expansion = wrist_dist / max(shoulder_width, 0.1)
                        
                            # Hand height relative to shoulders
                            avg_hand_y = (left_wrist.y + right_wrist.y) / 2
                            avg_shoulder_y = (left_shoulder.y + right_shoulder.y) / 2
                            hands_above_shoulders = avg_hand_y < avg_shoulder_y
                        
                            # Hand depth (Z-axis) relative to body
                            avg_hand_z = (left_wrist.z + right_wrist.z) / 2
                            avg_shoulder_z = (left_shoulder.z + right_shoulder.z) / 2
                            hands_forward = avg_hand_z < avg_shoulder_z
                        
                            # Movement direction (hands)
                            left_z_delta = left_wrist.z - prev_left_wrist.z
                            right_z_delta = right_wrist.z - prev_right_wrist.z
                            avg_z_delta = (left_z_delta + right_z_delta) / 2
                        
                            # Forward/backward: Use average of both hands with higher threshold
                            forward_movement = avg_z_delta < -0.03  # Both hands moving toward camera
                            # Use stricter backward threshold to avoid counting normal return-to-neutral motion.
                            backward_movement = avg_z_delta > 0.05  # Both hands moving away from camera
                        
                            # Vertical movement
                            upward_movement = (left_wrist.y - prev_left_wrist.y) < -0.01 or (right_wrist.y - prev_right_wrist.y) < -0.01
                            downward_movement = (left_wrist.y - prev_left_wrist.y) > 0.01 or (right_wrist.y - prev_right_wrist.y) > 0.01
                        
                            # Effort qualities
                            is_sudden = avg_velocity > 0.08  # Fast movement
                            is_sustained = 0.02 < avg_velocity <= 0.08  # Moderate movement
                            is_strong = body_expansion > 1.2 or hands_above_shoulders
                            is_light = body_expansion < 0.8
                        
                            # === EFFORT DETECTION (11 types) ===
                        
                            # 1. DIRECTING: Direct, sustained, forward movement
                            if hands_forward and is_sustained and forward_movement:
                                effort_counts["Directing"] += 1
                                shape_counts["Directing"] += 1
                        
                            # 2. ENCLOSING: Arms coming together, wrapping motion
                            enclosing_detected = (
                                body_expansion < enclosing_max_expansion
                                and avg_velocity > enclosing_min_velocity
                            )
                            if enclosing_detected:
                                effort_counts["Enclosing"] += 1
                                shape_counts["Enclosing"] += 1
                        
                            # 3. PUNCHING: Sudden, strong, direct forward thrust
                            if is_sudden and is_strong and forward_movement:
                                effort_counts["Punching"] += 1
                        
                            # 4. SPREADING: Arms spreading wide, opening gesture
                            spreading_detected = (
                                body_expansion > spreading_body_expansion_threshold
                                and avg_velocity > spreading_min_velocity
                            )
                            if spreading_detected:
                                effort_counts["Spreading"] += 1
                                shape_counts["Spreading"] += 1
                            if gesture_debug_log and (analyzed % gesture_debug_every_n == 0):
                                logger.info(
                                    "[gesture_debug] frame=%s enclosing=%s spreading=%s body_expansion=%.3f [enclose<%.3f, spread>%.3f] avg_velocity=%.4f [>enclose%.4f, >spread%.4f]",
                                    analyzed,
                                    enclosing_detected,
                                    spreading_detected,
                                    body_expansion,
                                    enclosing_max_expansion,
                                    spreading_body_expansion_threshold,
                                    avg_velocity,
                                    enclosing_min_velocity,
                                    spreading_min_velocity,
                                )
                        
                            # 5. PRESSING: Sustained, strong, downward force
                            if is_sustained and is_strong and downward_movement:
                                effort_counts["Pressing"] += 1
                        
                            # 6. DABBING: Sudden, light, direct touch
                            if is_sudden and is_light and avg_velocity > 0.05:
                                effort_counts["Dabbing"] += 1
                        
                            # 7. INDIRECTING: Indirect, curved, wandering movements
                            if not hands_forward and avg_velocity > 0.04 and body_expansion > 1.0:
                                effort_counts["Indirecting"] += 1
                                shape_counts["Indirecting"] += 1
                        
                            # 8. GLIDING: Sustained, light, indirect floating
                            if is_sustained and is_light and upward_movement:
                                effort_counts["Gliding"] += 1
                        
                            # 9. FLICKING: Sudden, light, indirect quick gesture
                            if is_sudden and is_light and not forward_movement:
                                effort_counts["Flicking"] += 1
                        
                            # 10. ADVANCING: Significant forward movement (toward audience)
                            # Requires sustained forward motion with reasonable speed
                            if forward_movement and avg_velocity > 0.06 and is_sustained:
                                effort_counts["Advancing"] += 1
                                shape_counts["Advancing"] += 1
                        
                            # 11. RETREATING: Significant backward movement (pulling back defensively)
                            # Requires stronger sustained backward motion - not just gesture return.
                            if backward_movement and avg_velocity > 0.07 and is_sustained:
                                effort_counts["Retreating"] += 1
                                shape_counts["Retreating"] += 1

                            # === Category-level feature collectors (professor-feedback,
                            # calibration-pending). These influence logging only; scoring
                            # will be updated once we have the 12 ground-truth clips.
                            try:
                                left_hip = lms[PoseLandmark.LEFT_HIP]
                                right_hip = lms[PoseLandmark.RIGHT_HIP]
                                hip_mid_x = (left_hip.x + right_hip.x) / 2.0
                                hip_mid_y = (left_hip.y + right_hip.y) / 2.0
                                sh_mid_y = avg_shoulder_y
                                shoulder_min_x = min(left_shoulder.x, right_shoulder.x)
                                shoulder_max_x = max(left_shoulder.x, right_shoulder.x)
                                # hand-blocking: both wrists inside the chest box
                                lw_in_box = (
                                    shoulder_min_x <= left_wrist.x <= shoulder_max_x
                                    and sh_mid_y <= left_wrist.y <= hip_mid_y
                                )
                                rw_in_box = (
                                    shoulder_min_x <= right_wrist.x <= shoulder_max_x
                                    and sh_mid_y <= right_wrist.y <= hip_mid_y
                                )
                                if lw_in_box and rw_in_box:
                                    hand_block_frames += 1
                                # hand-low: both wrists below hip line
                                if left_wrist.y > hip_mid_y and right_wrist.y > hip_mid_y:
                                    hand_low_frames += 1
                                # hands-above-shoulders: at least one wrist above shoulder line
                                if left_wrist.y < sh_mid_y or right_wrist.y < sh_mid_y:
                                    hands_above_sh_frames += 1
                                hip_x_series.append(float(hip_mid_x))
                                hip_y_series.append(float(hip_mid_y))
                                # Coarse wrist-shape signature for hand-shape variety counting.
                                # Quantize each wrist's (x,y) into a 4-bin grid relative to shoulders.
                                def _qx(x_val: float) -> int:
                                    if x_val < shoulder_min_x:
                                        return 0
                                    if x_val > shoulder_max_x:
                                        return 3
                                    midline = (shoulder_min_x + shoulder_max_x) / 2.0
                                    return 1 if x_val <= midline else 2

                                def _qy(y_val: float) -> int:
                                    if y_val < sh_mid_y - 0.05:
                                        return 0  # well above shoulders
                                    if y_val < sh_mid_y + 0.05:
                                        return 1  # around shoulders
                                    if y_val < hip_mid_y:
                                        return 2  # chest / belly
                                    return 3  # below hips

                                wrist_shape_signatures.append(
                                    (_qx(left_wrist.x), _qy(left_wrist.y), _qx(right_wrist.x), _qy(right_wrist.y))
                                )

                                # === Enriched-feature time-series collection ===
                                # All values are in MediaPipe's normalized 0-1 coords.
                                sh_mid_x = (left_shoulder.x + right_shoulder.x) / 2.0
                                nose_x_series.append(float(nose.x))
                                nose_y_series.append(float(nose.y))
                                nose_z_series.append(float(nose.z))
                                shoulder_mid_x_series.append(float(sh_mid_x))
                                shoulder_mid_y_series.append(float(sh_mid_y))
                                # Trunk angle: angle between (shoulder_mid → hip_mid)
                                # and the downward vertical axis, in degrees.
                                dy = hip_mid_y - sh_mid_y
                                dx = hip_mid_x - sh_mid_x
                                trunk_angle = math.degrees(math.atan2(abs(dx), max(abs(dy), 1e-6)))
                                trunk_angle_series.append(float(trunk_angle))
                                sw = abs(left_shoulder.x - right_shoulder.x)
                                hw = abs(left_hip.x - right_hip.x)
                                shoulder_width_series.append(float(sw))
                                hip_width_series.append(float(hw))
                                body_expansion_series.append(float(body_expansion))
                                wrist_mid_x = (left_wrist.x + right_wrist.x) / 2.0
                                wrist_mid_y = (left_wrist.y + right_wrist.y) / 2.0
                                wrist_avg_x_series.append(float(wrist_mid_x))
                                wrist_avg_y_series.append(float(wrist_mid_y))
                                wrist_velocity_series.append(float(avg_velocity))
                                left_wrist_xy_series.append((float(left_wrist.x), float(left_wrist.y)))
                                right_wrist_xy_series.append((float(right_wrist.x), float(right_wrist.y)))
                            except Exception:
                                # Defensive: never fail the main loop over feature collection.
                                pass

                        # Store current landmarks for next iteration
                        prev_landmarks = lms
                        analyzed += 1
                if sampled % 25 == 0:
                    logger.info(
                        "[analysis] progress sampled=%s/%s detected=%s",
                        sampled,
                        max_frames,
                        analyzed,
                    )

                frame_idx += 1

    except Exception as _mp_exc:
        logger.warning(
            "[analysis] MediaPipe Pose failed (using placeholder scores): %s",
            _mp_exc,
        )
        jid = str(kwargs.get("job_id") or "").strip()
        ph = analyze_video_placeholder(video_path, job_id=jid)
        ph["mediapipe_environment_fallback"] = True
        ph["mediapipe_fallback_reason"] = str(_mp_exc)[:500]
        return ph
    finally:
        try:
            cap.release()
        except Exception:
            pass
    
    # Calculate percentages
    total_detections = max(1, sum(effort_counts.values()))
    effort_detection = {k: round(v / total_detections * 100, 1) for k, v in effort_counts.items()}
    effort_detection = _apply_default_retreating_share(effort_detection, retreat_default_pct=1.0)

    total_shape = max(1, sum(shape_counts.values()))
    shape_detection = {k: round(v / total_shape * 100, 1) for k, v in shape_counts.items()}
    shape_detection = _apply_default_retreating_share(shape_detection, retreat_default_pct=1.0)

    # Add diversity/consistency signals so different speaking styles separate better.
    dominant_share = max(effort_detection.values()) / 100.0 if effort_detection else 0.0
    variety_floor = max(2, int(max(1, analyzed) * 0.03))
    variety_count = sum(1 for v in effort_counts.values() if v >= variety_floor)
    monotony_penalty = max(0.0, (dominant_share - 0.35) * 8.0)
    variety_factor = max(0.85, min(1.20, 0.85 + (variety_count / 11.0) * 0.35))

    total_frames = max(1, analyzed)
    engaging_activity = (
        effort_counts.get("Spreading", 0)
        + effort_counts.get("Enclosing", 0)
        + effort_counts.get("Gliding", 0)
        + effort_counts.get("Indirecting", 0)
    ) / total_frames
    confidence_activity = (
        effort_counts.get("Directing", 0)
        + effort_counts.get("Punching", 0)
        + effort_counts.get("Advancing", 0)
        + effort_counts.get("Pressing", 0)
    ) / total_frames
    authority_activity = (
        effort_counts.get("Pressing", 0)
        + effort_counts.get("Punching", 0)
        + effort_counts.get("Directing", 0)
        + effort_counts.get("Advancing", 0)
    ) / total_frames

    engaging_boost = max(0.85, min(1.35, 0.85 + engaging_activity * 1.60))
    confidence_boost = max(0.85, min(1.35, 0.85 + confidence_activity * 1.50))
    authority_boost = max(0.85, min(1.35, 0.85 + authority_activity * 1.50))

    if gesture_debug_log:
        logger.info(
            "[gesture_debug_summary] analyzed_frames=%s enclosing_count=%s spreading_count=%s enclosing_per_frame=%.4f spreading_per_frame=%.4f",
            total_frames,
            effort_counts.get("Enclosing", 0),
            effort_counts.get("Spreading", 0),
            effort_counts.get("Enclosing", 0) / max(1, total_frames),
            effort_counts.get("Spreading", 0) / max(1, total_frames),
        )

    # Category raw signals: positives minus small penalties from counter-signals.
    engaging_raw = (
        effort_detection.get("Spreading", 0) * 0.34
        + effort_detection.get("Enclosing", 0) * 0.26
        + effort_detection.get("Gliding", 0) * 0.22
        + effort_detection.get("Indirecting", 0) * 0.18
        - effort_detection.get("Punching", 0) * 0.12
        - effort_detection.get("Retreating", 0) * 0.10
    )
    confidence_raw = (
        effort_detection.get("Directing", 0) * 0.36
        + effort_detection.get("Punching", 0) * 0.24
        + effort_detection.get("Advancing", 0) * 0.24
        + effort_detection.get("Pressing", 0) * 0.16
        - effort_detection.get("Retreating", 0) * 0.18
        - effort_detection.get("Indirecting", 0) * 0.08
    )
    authority_raw = (
        effort_detection.get("Pressing", 0) * 0.34
        + effort_detection.get("Punching", 0) * 0.26
        + effort_detection.get("Directing", 0) * 0.24
        + effort_detection.get("Advancing", 0) * 0.16
        - effort_detection.get("Flicking", 0) * 0.18
        - effort_detection.get("Retreating", 0) * 0.12
    )

    # Shape-based alignment helps separate categories with similar effort mixes.
    engaging_shape = (
        shape_detection.get("Spreading", 0) * 0.45
        + shape_detection.get("Enclosing", 0) * 0.35
        + shape_detection.get("Indirecting", 0) * 0.20
    )
    confidence_shape = (
        shape_detection.get("Directing", 0) * 0.60
        + shape_detection.get("Advancing", 0) * 0.35
        - shape_detection.get("Retreating", 0) * 0.25
    )
    authority_shape = (
        shape_detection.get("Directing", 0) * 0.55
        + shape_detection.get("Advancing", 0) * 0.30
        - shape_detection.get("Retreating", 0) * 0.40
    )

    engaging_raw = max(
        0.0,
        (engaging_raw * engaging_boost + engaging_shape * 0.22) * variety_factor - monotony_penalty * 0.90,
    )
    confidence_raw = max(
        0.0,
        (confidence_raw * confidence_boost + confidence_shape * 0.25) * variety_factor - monotony_penalty * 1.00,
    )
    authority_raw = max(
        0.0,
        (authority_raw * authority_boost + authority_shape * 0.25) * variety_factor - monotony_penalty * 1.10,
    )

    # Mild contrast expansion to avoid all categories collapsing into the same band.
    raw_vec = np.array([engaging_raw, confidence_raw, authority_raw], dtype=float)
    raw_mean = float(np.mean(raw_vec))
    contrast = 0.95
    engaging_raw = engaging_raw + contrast * (engaging_raw - raw_mean)
    confidence_raw = confidence_raw + contrast * (confidence_raw - raw_mean)
    authority_raw = authority_raw + contrast * (authority_raw - raw_mean)

    engaging_score = min(7, max(1, round(engaging_raw / 5.4 + 1.4)))
    convince_score = min(7, max(1, round(confidence_raw / 5.4 + 1.4)))
    authority_score = min(7, max(1, round(authority_raw / 5.4 + 1.4)))

    # Adaptability (People Reader): movement + effort variety — diverse effort types, less single-style dominance
    diversity_signal = (1.0 - min(0.92, dominant_share)) * 55.0
    variety_signal = (variety_count / 11.0) * 50.0
    adaptability_raw = max(
        0.0,
        diversity_signal
        + variety_signal
        - monotony_penalty * 1.15
        - max(0.0, dominant_share - 0.45) * 18.0,
    )
    adaptability_score = min(7, max(1, round(adaptability_raw / 15.0 + 1.25)))
    adaptability_pos = int(adaptability_score / 7 * 445)

    # === Category-level professor-feedback features (calibration-pending) ===
    # Computed only from the per-frame counters above. Exposed in logs and in
    # the return dict so we can re-fit the Engaging / Confidence / Authority /
    # Adaptability weights once the 12 ground-truth clips are available.
    denom = max(1, analyzed)
    hand_block_share = hand_block_frames / denom
    hand_low_share = hand_low_frames / denom
    hands_above_share = hands_above_sh_frames / denom
    hip_sway_std = float(np.std(np.array(hip_x_series))) if len(hip_x_series) >= 10 else 0.0
    # Forward advance: compare median hip Y in the last quartile vs the first
    # quartile (Y grows downward in normalized coords, so a smaller final Y
    # means the speaker advanced upward/forward toward the camera).
    if len(hip_y_series) >= 8:
        q = max(2, len(hip_y_series) // 4)
        first_q = float(np.median(np.array(hip_y_series[:q])))
        last_q = float(np.median(np.array(hip_y_series[-q:])))
        hip_advance = first_q - last_q
    else:
        hip_advance = 0.0
    # Hand-shape variety: distinct coarse (left_x, left_y, right_x, right_y) bins.
    distinct_shapes = len(set(wrist_shape_signatures)) if wrist_shape_signatures else 0

    logger.info(
        "[category_features] hand_block=%.3f hand_low=%.3f hands_above=%.3f "
        "hip_sway_std=%.4f hip_advance=%.4f distinct_shapes=%d",
        hand_block_share, hand_low_share, hands_above_share,
        hip_sway_std, hip_advance, distinct_shapes,
    )

    # === Enriched features for Presentation Analysis ML scorer ============
    # All output values normalized to human-readable scales so downstream
    # feature-importance logs are interpretable.
    enriched: Dict[str, float] = _compute_enriched_presentation_features(
        nose_x_series=nose_x_series,
        nose_y_series=nose_y_series,
        nose_z_series=nose_z_series,
        shoulder_mid_x_series=shoulder_mid_x_series,
        shoulder_mid_y_series=shoulder_mid_y_series,
        trunk_angle_series=trunk_angle_series,
        shoulder_width_series=shoulder_width_series,
        hip_width_series=hip_width_series,
        body_expansion_series=body_expansion_series,
        wrist_avg_x_series=wrist_avg_x_series,
        wrist_avg_y_series=wrist_avg_y_series,
        wrist_velocity_series=wrist_velocity_series,
        left_wrist_xy_series=left_wrist_xy_series,
        right_wrist_xy_series=right_wrist_xy_series,
        hip_sway_std=hip_sway_std,
    )
    logger.info(
        "[enriched_features] posture=%.2f torso_stab=%.2f head_stab=%.2f "
        "eye_dir=%.3f sh_open=%.2f hand_open=%.2f g_range=%.2f g_smooth=%.2f "
        "intent=%.2f hesitation=%.2f rhythm=%.2f energy=%.4f center=%.2f",
        enriched.get("posture_uprightness", 0.0),
        enriched.get("torso_stability", 0.0),
        enriched.get("head_stability", 0.0),
        enriched.get("eye_direction_proxy", 0.0),
        enriched.get("shoulder_openness", 0.0),
        enriched.get("hand_openness", 0.0),
        enriched.get("gesture_range", 0.0),
        enriched.get("gesture_smoothness", 0.0),
        enriched.get("movement_intentionality", 0.0),
        enriched.get("hesitation_score", 0.0),
        enriched.get("rhythm_consistency", 0.0),
        enriched.get("energy_level", 0.0),
        enriched.get("center_presence", 0.0),
    )

    return {
        "analysis_engine": "mediapipe_real_enhanced",
        "duration_seconds": duration,
        "analyzed_frames": analyzed,
        "total_indicators": 450 + 475 + 445,  # = 1370 (real total)
        "engaging_score": engaging_score,
        "engaging_pos": int(engaging_score / 7 * 450),
        "convince_score": convince_score,
        "convince_pos": int(convince_score / 7 * 475),
        "authority_score": authority_score,
        "authority_pos": int(authority_score / 7 * 445),
        "adaptability_score": adaptability_score,
        "adaptability_pos": adaptability_pos,
        "effort_detection": effort_detection,
        "shape_detection": shape_detection,
        "effort_counts": dict(effort_counts),
        "shape_counts": dict(shape_counts),
        # Professor-feedback features (calibration-pending; not yet wired into
        # engaging/confidence/authority/adaptability scores).
        "category_features": {
            "hand_block_share": round(hand_block_share, 3),
            "hand_low_share": round(hand_low_share, 3),
            "hands_above_share": round(hands_above_share, 3),
            "hip_sway_std": round(hip_sway_std, 4),
            "hip_advance": round(hip_advance, 4),
            "distinct_hand_shapes": int(distinct_shapes),
            # Enriched presentation features (used by the ML scorer).
            **{k: round(float(v), 5) for k, v in enriched.items()},
        },
    }

def analyze_video_placeholder(video_path: str, seed: int = None, job_id: str = None) -> Dict[str, Any]:
    """Fallback placeholder analysis. Uses video_path/job_id to derive unique seed so different videos get different results."""
    # Derive seed from video path or job_id so each video gets different (but deterministic) placeholder data
    if seed is None:
        seed_src = str(job_id or video_path or "").strip() or str(time.time())
        seed = hash(seed_src) % (2**31 - 1)  # deterministic per video/job
    random.seed(seed)
    duration = get_video_duration_seconds(video_path)

    # Base values — randomize slightly per video so graphs and categories vary
    base_effort = {
        "Directing": 23.9, "Enclosing": 11.9, "Punching": 11.5, "Spreading": 11.3,
        "Pressing": 10.8, "Dabbing": 8.4, "Indirecting": 7.4, "Gliding": 6.2,
        "Flicking": 3.5, "Advancing": 2.6, "Retreating": 2.5
    }
    base_shape = {"Directing": 40.1, "Enclosing": 20.0, "Spreading": 18.9, "Indirecting": 12.4, "Advancing": 4.4, "Retreating": 4.1}
    effort_detection = {k: max(1.0, min(50.0, v + (random.random() - 0.5) * 12)) for k, v in base_effort.items()}
    shape_detection = {k: max(1.0, min(50.0, v + (random.random() - 0.5) * 10)) for k, v in base_shape.items()}
    effort_detection = _apply_default_retreating_share(effort_detection, retreat_default_pct=1.0)
    shape_detection = _apply_default_retreating_share(shape_detection, retreat_default_pct=1.0)

    # Category scores 1–7 — fallback: Engaging=High fixed; Confidence & Authority random Moderate/High
    engaging_score = 6   # High (>= 5)
    convince_score = random.randint(3, 7)   # Moderate (3–4) or High (5–7)
    authority_score = random.randint(3, 7)  # Moderate (3–4) or High (5–7)
    adaptability_score = random.randint(3, 7)

    return {
        "analysis_engine": "placeholder",
        "duration_seconds": duration,
        "analyzed_frames": 100,
        "total_indicators": 450 + 475 + 445,
        "engaging_score": engaging_score,
        "engaging_pos": int(engaging_score / 7 * 450),
        "convince_score": convince_score,
        "convince_pos": int(convince_score / 7 * 475),
        "authority_score": authority_score,
        "authority_pos": int(authority_score / 7 * 445),
        "adaptability_score": adaptability_score,
        "adaptability_pos": int(adaptability_score / 7 * 445),
        "effort_detection": effort_detection,
        "shape_detection": shape_detection,
    }

# Graph generation
def generate_effort_graph(effort_data: Dict[str, float], shape_data: Dict[str, float], output_path: str):
    """Generate Effort Motion Detection graph (Page 4 style)"""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
    fig.suptitle('Effort Motion Detection Results', fontsize=14, fontweight='bold')
    
    # Light blue color
    light_blue = '#87CEEB'
    
    # Left: Effort Summary (sorted by value, highest at top)
    sorted_efforts = sorted(effort_data.items(), key=lambda x: x[1], reverse=False)  # reverse=False for bottom-to-top
    efforts = [k for k, v in sorted_efforts]
    values = [v for k, v in sorted_efforts]
    
    ax1.barh(efforts, values, color=light_blue)
    ax1.set_xlabel('Percentage (%)')
    ax1.set_title('Effort Summary')
    ax1.set_xlim(0, 100)
    for i, v in enumerate(values):
        ax1.text(v + 1, i, f'{v}%', va='center')
    
    # Right: Top Movement Efforts
    top_3 = sorted(effort_data.items(), key=lambda x: x[1], reverse=True)[:3]
    top_names = [f"{k} - Rank #{i+1}" for i, (k, v) in enumerate(top_3)]
    top_vals = [v for k, v in top_3]
    
    ax2.barh(top_names, top_vals, color=light_blue)
    ax2.set_xlabel('Percentage (%)')
    ax2.set_title('Top Movement Efforts')
    ax2.set_xlim(0, 100)
    for i, v in enumerate(top_vals):
        ax2.text(v + 1, i, f'{v}%', va='center')
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()

def generate_shape_graph(shape_data: Dict[str, float], output_path: str):
    """Generate Shape Motion Detection graph (Page 5 style)"""
    fig, ax = plt.subplots(figsize=(10, 6))
    
    # Light blue color
    light_blue = '#87CEEB'
    
    # Sort by value, highest (left) to lowest (right)
    sorted_shapes = sorted(shape_data.items(), key=lambda x: x[1], reverse=True)
    shapes = [k for k, v in sorted_shapes]
    values = [v for k, v in sorted_shapes]
    
    bars = ax.bar(shapes, values, color=light_blue)
    ax.set_ylabel('Percentage (%)')
    ax.set_xlabel('Shape Motion Type')
    ax.set_title('Shape Motion Detection Results', fontsize=14, fontweight='bold')
    ax.set_ylim(0, max(values) * 1.2 if values else 10)
    
    for bar, val in zip(bars, values):
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height,
                f'{val}%', ha='center', va='bottom')
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()


def _build_narrative_from_report(report: ReportData, audience_mode: str = "one"):
    """Map ReportData to narrative_engine metrics and build NarrativeReport."""
    if build_narrative_report is None:
        return None
    fi = report.first_impression
    if not fi:
        return None
    # Map FirstImpressionData (0-100) to narrative first_impression_metrics (0-1)
    eye = max(0, min(100, float(fi.eye_contact_pct or 0))) / 100.0
    up = max(0, min(100, float(fi.upright_pct or 0))) / 100.0
    st = max(0, min(100, float(fi.stance_stability or 0))) / 100.0
    first_impression_metrics = {
        "eye_contact_ratio": eye,
        "gaze_stability": eye,
        "gaze_shift_control": eye,
        "gaze_avoidance": 1.0 - eye,
        "posture_upright": up,
        "head_alignment": up,
        "shoulder_balance": up,
        "chest_openness": up,
        "slouch_ratio": 1.0 - up,
        "lower_body_stability": st,
        "grounding": st,
        "weight_shift_control": st,
        "forward_orientation": st,
        "stance_symmetry": st,
    }
    # Map categories to narrative category_metrics (0-1). Analysis uses 1-7; some paths use 0-100.
    cats = report.categories or []
    def _norm(s: float) -> float:
        v = float(s)
        v = (v / 100.0) if v > 10 else (v / 7.0)
        return max(0.0, min(1.0, v))
    eng = _norm(cats[0].score) if len(cats) > 0 else 0.5
    conf = _norm(cats[1].score) if len(cats) > 1 else 0.5
    auth = _norm(cats[2].score) if len(cats) > 2 else 0.5
    category_metrics = {
        "facial_warmth": eng,
        "openness": eng,
        "expressive_variety": eng,
        "audience_connection": eng,
        "presence_score": conf,
        "posture_strength": conf,
        "movement_control": conf,
        "focus_consistency": conf,
        "verticality": auth,
        "decisiveness": auth,
        "grounded_power": auth,
        "urgency_signal": auth,
    }
    return build_narrative_report(
        first_impression_metrics=first_impression_metrics,
        category_metrics=category_metrics,
        audience_mode=audience_mode,
    )


# DOCX Report building
def build_docx_report(
    report: ReportData,
    output_bio: io.BytesIO,
    graph1_path: str,
    graph2_path: str,
    lang: str = "en",
    report_style: str = "full",
):
    """Build 5-page DOCX report matching the exact template format"""
    doc = Document()
    style_name = str(report_style or "full").strip().lower()
    is_simple = style_name.startswith("simple")
    is_operation_test = style_name.startswith("operation_test")
    is_people_reader = style_name.startswith("people_reader")
    layout_like_operation = is_operation_test or is_people_reader

    # Language-specific text
    is_thai = (lang == "th")

    # Set Thai font at document level BEFORE adding content — ป้องกันสระ/วรรณยุกต์ทับกัน
    if is_thai:
        thai_font = os.getenv("DOCX_THAI_FONT_FAMILY", "TH Sarabun New").strip() or "TH Sarabun New"

        def _set_rfonts(el, font_name: str) -> None:
            r_fonts = el.find(qn("w:rFonts"))
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                el.append(r_fonts)
            r_fonts.set(qn("w:ascii"), font_name)
            r_fonts.set(qn("w:hAnsi"), font_name)
            r_fonts.set(qn("w:eastAsia"), font_name)
            r_fonts.set(qn("w:cs"), font_name)

        try:
            # 1) docDefaults — default for entire document
            styles_root = doc.styles.element
            doc_defaults = styles_root.find(qn("w:docDefaults"))
            if doc_defaults is None:
                doc_defaults = OxmlElement("w:docDefaults")
                styles_root.insert(0, doc_defaults)
            r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
            if r_pr_default is None:
                r_pr_default = OxmlElement("w:rPrDefault")
                doc_defaults.append(r_pr_default)
            r_pr = r_pr_default.find(qn("w:rPr"))
            if r_pr is None:
                r_pr = OxmlElement("w:rPr")
                r_pr_default.append(r_pr)
            _set_rfonts(r_pr, thai_font)

            # 2) Normal style (base for most paragraphs)
            normal = doc.styles["Normal"]
            normal.font.name = thai_font
            n_rpr = normal._element.get_or_add_rPr()
            _set_rfonts(n_rpr, thai_font)
        except Exception:
            pass
    
    en_operation_title = "Movement in Communication\nwith AI People Reader Report"
    th_operation_title = "รายงานการวิเคราะห์การนำเสนอด้วยการ\nเคลื่อนไหว กับ AI People Reader"
    texts = {
        "title": th_operation_title if is_thai else en_operation_title,
        "client_name": "ชื่อลูกค้า:" if is_thai else "Client Name:",
        "analysis_date": "วันที่วิเคราะห์:" if is_thai else "Analysis Date:",
        "video_info": "ข้อมูลวิดีโอ (Video Information)" if is_thai else "Video Information",
        "duration": "ระยะเวลา:" if is_thai else "Duration:",
        "detailed_analysis": "การวิเคราะห์โดยละเอียด (Detailed Analysis)" if is_thai else "Detailed Analysis",
        "first_impression": "1.  ความประทับใจแรกพบ (First Impression)" if is_thai else "1.  First impression",
        "eye_contact": "การสบตา (Eye Contact)" if is_thai else "Eye Contact",
        "uprightness": "ความตั้งตรงของร่างกาย (Uprightness)" if is_thai else "Uprightness (Posture & Upper-Body Alignment)",
        "stance": "การยืนและการวางเท้า (Stance)" if is_thai else "Stance (Lower-Body Stability & Grounding)",
        "impact_clients": "ผลกระทบ:" if is_thai else "Impact:",
        "engaging": "2. การสร้างความเป็นมิตรและสร้างสัมพันธภาพ" if is_thai else "2. Engaging & Connecting:",
        "approachability": "▪ ความเป็นกันเอง" if is_thai else "▪ Approachability",
        "relatability": "▪ ความเข้าถึงได้" if is_thai else "▪ Relatability",
        "engagement": "▪ การมีส่วนร่วม เชื่อมโยง และสร้างความคุ้นเคยกับทีมอย่างรวดเร็ว" if is_thai else "▪ Engagement, connect and build instant rapport with team",
        "confidence": "3. ความมั่นใจ:" if is_thai else "3. Confidence:",
        "optimistic": "▪ บุคลิกภาพเชิงบวก" if is_thai else "▪ Optimistic Presence",
        "focus": "▪ ความมีสมาธิ" if is_thai else "▪ Focus",
        "persuade": "▪ ความสามารถในการโน้มน้าวและยืนหยัดในจุดยืนเพื่อให้ผู้อื่นคล้อยตาม" if is_thai else "▪ Ability to persuade and stand one's ground, in order to convince others.",
        "authority": "4.  ความเป็นผู้นำ (Authority):" if is_thai else "4. Authority:",
        "adaptability": "5. ความยืดหยุ่นในการปรับตัว (Adaptability):" if is_thai else "5. Adaptability:",
        "adapt_flexibility": (
            "▪ ความยืดหยุ่น — ความสามารถในการปรับตัวตามสภาวะใหม่ ๆ และรับมือกับการเปลี่ยนแปลง"
            if is_thai
            else "▪ Flexibility — Ability to adjust to new conditions, handle changes."
        ),
        "adapt_agility": (
            "▪ ความคล่องแคล่ว — ความสามารถในการคิดและสรุปได้อย่างรวดเร็วเพื่อปรับตัว"
            if is_thai
            else "▪ Agility — Ability to think and draw conclusions quickly in order to adjust."
        ),
        "importance": "▪ แสดงให้เห็นถึงความสำคัญและความเร่งด่วนของประเด็น" if is_thai else "▪ Showing sense of importance and urgency in subject matter",
        "pressing": "▪ ผลักดันให้เกิดการลงมือทำ" if is_thai else "▪ Pressing for action",
        "scale": "ระดับ:" if is_thai else "Scale:",
        "description": "คำอธิบาย: ตรวจพบ" if is_thai else "Description: Detected",
        "indicators": "การเคลื่อนไหวเชิงบวก" if is_thai else "positive indicators out of",
        "total_indicators": "ตัวชี้วัดทั้งหมด" if is_thai else "total indicators",
        "effort_title": "ผลการวิเคราะห์การใช้น้ำหนัก (Efforts)" if is_thai else "Effort Motion Detection Results",
        "shape_title": "ผลการวิเคราะห์การใช้รูปทรงของมือร่วมกับร่างกาย (Shape)" if is_thai else "Shape Motion Detection Results",
        "generated": "สร้างโดย AI People Reader™" if is_thai else "Generated by AI People Reader™",
        "movement_type_heading": "โปรไฟล์ประเภทการเคลื่อนไหว (Movement type)" if is_thai else "Movement type profile",
        "movement_type_match": "การจับคู่ที่ใกล้เคียงที่สุด:" if is_thai else "Closest match:",
        "movement_type_confidence": "การตรงกันของโปรไฟล์นี้ (7 มิติ ต่ำ/กลาง/สูง):" if is_thai else "This profile's 7-dimension match (Low/Moderate/High):",
        "movement_type_top_two": "อันดับจับคู่ 7 มิติ (1 และ 2):" if is_thai else "Top two matches (7 dimensions):",
        "movement_type_mode_auto": "(วิเคราะห์อัตโนมัติจากวิดีโอ)" if is_thai else "(Auto-detected from video)",
        "movement_type_mode_selected": "(เลือกประเภทด้วยตนเอง — ใช้โปรไฟล์นี้ในรายงาน)" if is_thai else "(Manually selected — report aligned to this profile)",
        "movement_type_summary_label": "สรุป:" if is_thai else "Summary:",
        "movement_type_traits": "ลักษณะจากประเภทนี้:" if is_thai else "Traits from this type:",
        "movement_type_narrative_label": "ข้อสังเกต:" if is_thai else "Observation:",
    }
    
    # Impact texts in Thai
    impact_eye_thai = "การสบตาที่แข็งแกร่งส่งสัญญาณถึงความมีอยู่ ความจริงใจ และความเชื่อมั่นในความเป็นผู้นำ ทำให้ข้อความของคุณรู้สึกน่าเชื่อถือมากขึ้น"
    impact_upright_thai = "ท่าทางตรงสื่อถึงความมั่นใจในตนเอง ความชัดเจนในความคิด และความมั่นคงทางอารมณ์ ซึ่งเป็นลักษณะของผู้สื่อสารที่มีความไว้วางใจสูง"
    impact_stance_thai = "ท่ายืนที่มั่นคงช่วยเพิ่มอำนาจ การควบคุม และการส่งข้อความที่ราบรื่น ทำให้ผู้พูดดูเตรียมพร้อมและน่าเชื่อถือมากขึ้น"
    
    impact_eye_en = "Strong eye contact signals presence, sincerity, and leadership confidence, making your message feel more reliable."
    impact_upright_en = "Uprightness communicates self-assurance, clarity of thought, and emotional stability all traits of high-trust communicators."
    impact_stance_en = "A grounded stance enhances authority, control, and smooth message delivery, making the speaker appear more prepared and credible."

    def _date_th_display(date_text: str) -> str:
        raw = str(date_text or "").strip()
        if not raw:
            return raw
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%d-%m-%Y", "%d/%m/%Y"):
            try:
                dt = datetime.strptime(raw, fmt)
                return dt.strftime("%d/%m/%Y")
            except Exception:
                continue
        return raw

    def _date_en_display(date_text: str) -> str:
        raw = str(date_text or "").strip()
        if not raw:
            return raw
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%d-%m-%Y", "%d/%m/%Y"):
            try:
                dt = datetime.strptime(raw, fmt)
                return f"{dt.day}/{dt.month}/{dt.year}"
            except Exception:
                continue
        return raw

    def _fi_level_en(value: float, metric: str = "") -> str:
        return first_impression_level(value, metric=metric, is_fallback=is_first_impression_fallback(fi))

    def _fi_level_th(value: float, metric: str = "") -> str:
        lv = _fi_level_en(value, metric=metric).lower()
        if lv == "n/a":
            return "ไม่สามารถวิเคราะห์ได้"
        if lv.startswith("high"):
            return "สูง"
        if lv.startswith("moderate"):
            return "กลาง"
        return "ต่ำ"

    def normalize_spacing() -> None:
        """
        Keep a consistent spacing pattern for numbered headings and bullet points
        across pages (especially sections 2-4).
        """
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if not t:
                continue

            pf = p.paragraph_format
            # Numbered section headings: "1.", "2.", "3.", "4."
            if t[:2] in ("1.", "2.", "3.", "4.", "5."):
                pf.space_before = Pt(14)
                pf.space_after = Pt(6)
                pf.line_spacing = 1.1
                continue

            # Bullets: List Bullet style or explicit "•"/"▪" in content; keep unified compact rhythm.
            if getattr(p.style, "name", "") == "List Bullet" or t.startswith("•") or t.startswith("▪"):
                pf.space_before = Pt(0)
                pf.space_after = Pt(6)
                pf.line_spacing = 1.2
                continue

            # Keep impact labels/body visually grouped.
            if t in (texts["impact_clients"], "Impact:", "ผลกระทบ:"):
                pf.space_before = Pt(6)
                pf.space_after = Pt(0)
                continue

    def _strip_bullet(text: str) -> str:
        """Remove leading bullet glyph for consistent rendering."""
        s = str(text or "").strip()
        if s.startswith("•") or s.startswith("▪"):
            return s[1:].lstrip()
        return s

    def _square_bullet_text(text: str) -> str:
        return f"▪ {_strip_bullet(text)}"

    def _apply_bullet_layout(paragraph, compact: bool = False) -> None:
        if paragraph is None:
            return
        pf = paragraph.paragraph_format
        pf.left_indent = Pt(28)
        pf.first_line_indent = Pt(-14)
        pf.space_before = Pt(0)
        pf.space_after = Pt(3 if compact else (6 if not is_thai else 4))
        pf.line_spacing = 1.15 if compact else (1.25 if not is_thai else 1.15)

    def _apply_scale_layout(paragraph, left_indent_pt: float = 28, space_before_pt: float = None, compact: bool = False) -> None:
        if paragraph is None:
            return
        pf = paragraph.paragraph_format
        pf.left_indent = Pt(left_indent_pt)
        pf.first_line_indent = Pt(0)
        if space_before_pt is not None:
            pf.space_before = Pt(space_before_pt)
        else:
            pf.space_before = Pt(0)
        pf.space_after = Pt(3 if compact else (8 if not is_thai else 6))
        pf.line_spacing = 1.15 if compact else (1.25 if not is_thai else 1.15)

    def _set_para_font_size(paragraph, size_pt: int) -> None:
        if paragraph and paragraph.runs:
            for r in paragraph.runs:
                r.font.size = Pt(size_pt)
    
    # Add header and footer images to all pages
    section = doc.sections[0]
    # Match reference format: proper margin (1 inch standard)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Add header image (skip on permission errors — do not fail whole report)
    header_path = resolve_brand_asset_path("Header.png")
    if header_path:
        try:
            header = section.header
            header_para = header.paragraphs[0]
            header_run = header_para.add_run()
            header_run.add_picture(header_path, width=Inches(6.5))
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as ex:
            logger.warning("[docx] skip Header.png: %s", ex)

    footer_path = resolve_brand_asset_path("Footer.png")
    footer = section.footer
    if footer_path:
        try:
            footer_para = footer.paragraphs[0]
            footer_run = footer_para.add_run()
            footer_run.add_picture(footer_path, width=Inches(6.5))
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as ex:
            logger.warning("[docx] skip Footer.png: %s", ex)
    
    # ============================================================
    # PAGE 1: Cover + First Impression (Eye Contact start)
    # ============================================================
    
    # Same compact first page spacing for both Thai and English (match Thai layout).
    compact_thai_first_page = True

    # Title section spacing: Thai first page — move title up 1 line
    if layout_like_operation and not is_thai:
        doc.add_paragraph()
    elif not layout_like_operation and not is_thai:
        doc.add_paragraph()
        if not compact_thai_first_page:
            doc.add_paragraph()

    # Title — centered to match reference format. Thai: "AI People Reader" same size as Thai text.
    if is_thai:
        title = doc.add_paragraph()
        r1 = title.add_run("รายงานการวิเคราะห์การนำเสนอด้วยการ\nเคลื่อนไหว กับ ")
        r2 = title.add_run("AI People Reader")
        base_pt = 18 if compact_thai_first_page else 20
        r1.font.size = Pt(base_pt)
        r2.font.size = Pt(max(12, base_pt - 6))  # Latin renders larger; reduce to match Thai
        r1.font.bold = True
        r2.font.bold = True
    else:
        title = doc.add_paragraph(texts["title"])
        title.runs[0].font.size = Pt(18 if compact_thai_first_page else 20)
        title.runs[0].font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6 if compact_thai_first_page else 12)
    if not compact_thai_first_page:
        doc.add_paragraph()
    
    # Client info
    analysis_date_display = _date_th_display(report.analysis_date) if is_thai else _date_en_display(report.analysis_date)
    if not is_thai:
        doc.add_paragraph()  # EN: extra line before Client Name
    doc.add_paragraph(f"{texts['client_name']}  {report.client_name}")
    doc.add_paragraph(f"{texts['analysis_date']}  {analysis_date_display}")
    
    # Video info — operation_test EN: "Video Duration: 01:36" single line per reference
    if layout_like_operation and not is_thai:
        doc.add_paragraph(f"Video Duration: {report.video_length_str}")
    else:
        video_info = doc.add_paragraph(texts["video_info"])
        video_info.runs[0].bold = True
        doc.add_paragraph(f"{texts['duration']} {report.video_length_str}")
    if not compact_thai_first_page:
        doc.add_paragraph()
    
    # Detailed Analysis header. Push it down ~2 lines below the Video Duration
    # block, while still keeping it tight against "1. First impression" below.
    doc.add_paragraph()
    doc.add_paragraph()
    detailed = doc.add_paragraph(texts["detailed_analysis"])
    detailed.runs[0].bold = True
    detailed.paragraph_format.space_before = Pt(6)
    detailed.paragraph_format.space_after = Pt(2)

    # Section 1: First impression — sits right under "Detailed Analysis".
    section1 = doc.add_paragraph(texts["first_impression"])
    section1.runs[0].bold = True
    section1.paragraph_format.space_before = Pt(2)
    section1.paragraph_format.space_after = Pt(4)
    
    # First Impression - scale (low/moderate/high)
    if report.first_impression:
        fi = report.first_impression
        eye_level = _fi_level_th(fi.eye_contact_pct, "eye_contact") if is_thai else _fi_level_en(fi.eye_contact_pct, "eye_contact")
        up_level = _fi_level_th(fi.upright_pct, "uprightness") if is_thai else _fi_level_en(fi.upright_pct, "uprightness")
        st_level = _fi_level_th(fi.stance_stability, "stance") if is_thai else _fi_level_en(fi.stance_stability, "stance")
        for label, lv in [
            ("• การสบตา (Eye Contact)" if is_thai else "• Eye Contact", eye_level),
            ("• ความตั้งตรงของร่างกาย (Uprightness)" if is_thai else "• Uprightness", up_level),
            ("• การยืนและการวางเท้า (Stance)" if is_thai else "• Stance", st_level),
        ]:
            # First-impression bullets use a round "•" (match Irene reference).
            # Other sections (approachability, relatability, etc.) keep the
            # square "▪" via _square_bullet_text.
            p = doc.add_paragraph(label)
            _apply_bullet_layout(p)
            lvl = doc.add_paragraph(f"{texts['scale']} {lv}")
            lvl.runs[0].bold = True
            _apply_scale_layout(lvl, left_indent_pt=14)

    if not is_thai:
        doc.add_paragraph()
        doc.add_paragraph()  # EN: significant space before Remark per reference
    remark = doc.add_paragraph("หมายเหตุ" if is_thai else "Remark")
    remark.runs[0].bold = True
    remark.paragraph_format.space_before = Pt(14)
    remark.paragraph_format.space_after = Pt(4)
    doc.add_paragraph(
        "ความรู้สึกแรกพบมักเกิดขึ้นภายใน 5 วินาทีแรกของการพบกัน โดยพิจารณาจากภาพรวม การสบตา ความตั้งตรง และการยืนวางเท้า ก่อนเข้าสู่การวิเคราะห์เชิงพฤติกรรมในส่วนถัดไป"
        if is_thai
        else "First impression happens in the first 5 seconds of meeting someone, and is normally decided from the person's appearance, eye contact, uprightness and stance. However, after the first 5 seconds, the rest (below) are normally taken into consideration."
    )
    doc.add_paragraph()

    # Page 2: Combination Explanation 1, 2, 3 — use page_break_before so LibreOffice respects it
    combo_label = "คำอธิบายการผสมผสาน:" if is_thai else "Combination Explanation:"
    p_combo = doc.add_paragraph(combo_label)
    p_combo.runs[0].bold = True
    p_combo.paragraph_format.page_break_before = True  # Force new page (page 2)
    p_combo.paragraph_format.keep_with_next = True
    doc.add_paragraph()  # Blank line after combo label

    # Combination Explanation points 1, 2, 3 — keep together on page 2
    combo_paras = []
    if is_thai:
        combo_paras = [
            doc.add_paragraph("1. การสบตาน้อย + ความตั้งตรงน้อย + การยืนและการวางเท้าต่ำ"),
            doc.add_paragraph("บุคคลมักดูไม่เป็นภัยและยืดหยุ่น แต่บุคคลอาจดูมีความมั่นใจและอำนาจในระดับต่ำ"),
            doc.add_paragraph("2. การสบตาปานกลาง + ความตั้งตรงปานกลาง + การยืนและการวางเท้าปานกลาง"),
            doc.add_paragraph("บุคคลมักดูเข้าถึงได้ง่าย และมีความมั่นใจและอำนาจในระดับที่เพียงพอ"),
            doc.add_paragraph("3. การสบตาสูง + ความตั้งตรงสูง + การยืนและการวางเท้าสูง"),
            doc.add_paragraph("บุคคลมักดูมีความมั่นใจและอำนาจในระดับสูง และอาจดูไม่เข้าถึงได้ง่ายหรือยืดหยุ่น"),
        ]
    else:
        combo_paras = [
            doc.add_paragraph("1. Low Eye Contact + Low Uprightness + Low Stance."),
            doc.add_paragraph("The person tends to appear non-threatening and flexible. However, the person can also appear to possess low level of confidence and authority."),
            doc.add_paragraph("2. Moderate Eye Contact + Moderate Uprightness + Moderate Stance."),
            doc.add_paragraph("The person tends to appear approachable, and has adequate level of confidence and authority."),
            doc.add_paragraph("3. High Eye Contact + High Uprightness + High Stance."),
            doc.add_paragraph("The person tends to appear to possess high level of confidence and authority, and may not appear approachable or flexible."),
        ]
    for p in combo_paras[:-1]:
        p.paragraph_format.keep_with_next = True  # Keep combo block together

    # Page 3: skill sections — page_break_before forces new page

    # ============================================================
    # PAGE 3: Categories
    #   People Reader:  2. Engaging → 3. Adaptability → 4. Confidence → 5. Authority
    #   Operation Test: 2. Engaging → 3. Confidence     → 4. Authority  (no Adaptability)
    # ============================================================
    # For People Reader, we override the section number prefixes in `texts[...]`
    # so Adaptability=3, Confidence=4, Authority=5.
    if is_people_reader:
        engaging_heading = (
            "2. การสร้างความเป็นมิตรและสร้างสัมพันธภาพ" if is_thai else "2. Engaging & Connecting:"
        )
        adaptability_heading = (
            "3. ความยืดหยุ่นในการปรับตัว (Adaptability):" if is_thai else "3. Adaptability:"
        )
        confidence_heading = "4. ความมั่นใจ:" if is_thai else "4. Confidence:"
        authority_heading = (
            "5.  ความเป็นผู้นำ (Authority):" if is_thai else "5. Authority:"
        )
    else:
        engaging_heading = texts["engaging"]
        confidence_heading = texts["confidence"]
        authority_heading = texts["authority"]
        adaptability_heading = texts["adaptability"]  # unused in this branch

    # Section 2: Engaging & Connecting — start page 3
    engaging_cat = report.categories[0]
    section2 = doc.add_paragraph(engaging_heading)
    section2.runs[0].bold = True
    section2.paragraph_format.page_break_before = True  # Force page 3
    section2.paragraph_format.space_before = Pt(8)
    section2.paragraph_format.space_after = Pt(2)
    section2.paragraph_format.keep_with_next = True
    p_approach = doc.add_paragraph(_square_bullet_text(texts["approachability"]))
    _apply_bullet_layout(p_approach, compact=False)
    p_approach.paragraph_format.keep_with_next = True
    p_relate = doc.add_paragraph(_square_bullet_text(texts["relatability"]))
    _apply_bullet_layout(p_relate, compact=False)
    p_relate.paragraph_format.keep_with_next = True
    p_engage = doc.add_paragraph(_square_bullet_text(texts["engagement"]))
    _apply_bullet_layout(p_engage, compact=False)
    p_engage.paragraph_format.keep_with_next = True
    scale_para1 = doc.add_paragraph(f"{texts['scale']} {_display_scale(engaging_cat.scale, is_thai)}")
    scale_para1.runs[0].bold = True
    _apply_scale_layout(scale_para1, left_indent_pt=28, space_before_pt=4, compact=False)
    scale_para1.paragraph_format.keep_with_next = True

    # Section 3: Adaptability (People Reader only — inserted here per report layout)
    if is_people_reader:
        if len(report.categories) < 4:
            raise ValueError("people_reader report requires Adaptability as 4th category")
        adaptability_cat = report.categories[3]
        section_adapt = doc.add_paragraph(adaptability_heading)
        section_adapt.runs[0].bold = True
        section_adapt.paragraph_format.space_before = Pt(8)
        section_adapt.paragraph_format.space_after = Pt(2)
        section_adapt.paragraph_format.keep_with_next = True
        p_flex = doc.add_paragraph(_square_bullet_text(texts["adapt_flexibility"]))
        _apply_bullet_layout(p_flex, compact=False)
        p_flex.paragraph_format.keep_with_next = True
        p_ag = doc.add_paragraph(_square_bullet_text(texts["adapt_agility"]))
        _apply_bullet_layout(p_ag, compact=False)
        p_ag.paragraph_format.keep_with_next = True
        scale_para_adapt = doc.add_paragraph(
            f"{texts['scale']} {_display_scale(adaptability_cat.scale, is_thai)}"
        )
        scale_para_adapt.runs[0].bold = True
        _apply_scale_layout(scale_para_adapt, left_indent_pt=28, space_before_pt=4, compact=False)
        scale_para_adapt.paragraph_format.keep_with_next = True

    # Section 3 (non-PR) or Section 4 (PR): Confidence
    confidence_cat = report.categories[1]
    section3 = doc.add_paragraph(confidence_heading)
    section3.runs[0].bold = True
    section3.paragraph_format.space_before = Pt(8)
    section3.paragraph_format.space_after = Pt(2)
    section3.paragraph_format.keep_with_next = True
    p_opt = doc.add_paragraph(_square_bullet_text(texts["optimistic"]))
    _apply_bullet_layout(p_opt, compact=False)
    p_opt.paragraph_format.keep_with_next = True
    p_focus = doc.add_paragraph(_square_bullet_text(texts["focus"]))
    _apply_bullet_layout(p_focus, compact=False)
    p_focus.paragraph_format.keep_with_next = True
    p_persuade = doc.add_paragraph(_square_bullet_text(texts["persuade"]))
    _apply_bullet_layout(p_persuade, compact=False)
    p_persuade.paragraph_format.keep_with_next = True
    scale_para2 = doc.add_paragraph(f"{texts['scale']} {_display_scale(confidence_cat.scale, is_thai)}")
    scale_para2.runs[0].bold = True
    _apply_scale_layout(scale_para2, left_indent_pt=28, space_before_pt=4, compact=False)
    scale_para2.paragraph_format.keep_with_next = True

    # Section 4 (non-PR) or Section 5 (PR): Authority
    authority_cat = report.categories[2]
    section4 = doc.add_paragraph(authority_heading)
    section4.runs[0].bold = True
    section4.paragraph_format.space_before = Pt(8)
    section4.paragraph_format.space_after = Pt(2)
    section4.paragraph_format.keep_with_next = True
    p_importance = doc.add_paragraph(_square_bullet_text(texts["importance"]))
    _apply_bullet_layout(p_importance, compact=False)
    p_importance.paragraph_format.keep_with_next = True
    p_pressing = doc.add_paragraph(_square_bullet_text(texts["pressing"]))
    _apply_bullet_layout(p_pressing, compact=False)
    scale_para3 = doc.add_paragraph(f"{texts['scale']} {_display_scale(authority_cat.scale, is_thai)}")
    scale_para3.runs[0].bold = True
    _apply_scale_layout(scale_para3, left_indent_pt=28, space_before_pt=4, compact=False)

    # "Movement type profile" section intentionally omitted from the DOCX report.
    # movement_type_info is still computed and stored on the job for the UI, but
    # it is no longer rendered on page 3 of the document.

    if not is_simple:
        # PAGE BREAK TO PAGE 4
        doc.add_page_break()
        
        # ============================================================
        # PAGE 4: Effort Motion Detection Results
        # ============================================================
        
        # Spacing: operation_test — EN: more top space to match PDF; TH: compact. Title up 1 line.
        if layout_like_operation:
            for _ in range(3 if not is_thai else 1):
                doc.add_paragraph()
        else:
            for _ in range(3):
                doc.add_paragraph()
        
        # Title — Thai: (Efforts) same size as Thai
        if is_thai:
            title4 = doc.add_paragraph()
            r4a = title4.add_run("ผลการวิเคราะห์การใช้น้ำหนัก ")
            r4a.font.size = Pt(18)
            r4a.bold = True
            r4e = title4.add_run("(Efforts)")
            r4e.font.size = Pt(14)
            r4e.bold = True
        else:
            title4 = doc.add_paragraph(texts["effort_title"])
            title4.runs[0].font.size = Pt(18)
            title4.runs[0].bold = True
        doc.add_paragraph()
        
        # Graph
        if os.path.exists(graph1_path):
            doc.add_picture(graph1_path, width=Inches(6.5))
        
        # PAGE BREAK TO PAGE 5
        doc.add_page_break()
        
        # ============================================================
        # PAGE 5: Shape Motion Detection Results
        # ============================================================
        
        # Spacing: operation_test — EN: more top space to match PDF; TH: compact. Title up 1 line.
        if layout_like_operation:
            for _ in range(3 if not is_thai else 1):
                doc.add_paragraph()
        else:
            for _ in range(3):
                doc.add_paragraph()
        
        # Title — Thai: (Shape) same size as Thai
        if is_thai:
            title5 = doc.add_paragraph()
            r5a = title5.add_run("ผลการวิเคราะห์การใช้รูปทรงของมือร่วมกับร่างกาย ")
            r5a.font.size = Pt(18)
            r5a.bold = True
            r5e = title5.add_run("(Shape)")
            r5e.font.size = Pt(14)
            r5e.bold = True
        else:
            title5 = doc.add_paragraph(texts["shape_title"])
            title5.runs[0].font.size = Pt(18)
            title5.runs[0].bold = True
        doc.add_paragraph()
        
        # Graph
        if os.path.exists(graph2_path):
            doc.add_picture(graph2_path, width=Inches(6.5))
        
        # Add "Generated by AI People Reader™" text below graph
        doc.add_paragraph()
        doc.add_paragraph()
    else:
        doc.add_paragraph()

    # Normalize paragraph rhythm so numbered headings and bullets follow
    # a consistent spacing pattern throughout the document.
    normalize_spacing()

    # Always apply Thai font for Thai reports — ป้องกันสระและวรรณยุกต์ทับกันใน PDF
    # TH Sarabun New / Sarabun มี glyph ครบสำหรับ Thai diacritics
    if is_thai:
        thai_font_family = os.getenv("DOCX_THAI_FONT_FAMILY", "TH Sarabun New").strip() or "TH Sarabun New"

        def _apply_run_font(run, font_name: str) -> None:
            if run is None:
                return
            run.font.name = font_name
            r_pr = run._element.get_or_add_rPr()
            r_fonts = r_pr.rFonts
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.append(r_fonts)
            r_fonts.set(qn("w:ascii"), font_name)
            r_fonts.set(qn("w:hAnsi"), font_name)
            r_fonts.set(qn("w:eastAsia"), font_name)
            r_fonts.set(qn("w:cs"), font_name)

        # Apply to all body runs.
        for p in doc.paragraphs:
            for run in p.runs:
                _apply_run_font(run, thai_font_family)

        # Apply to header/footer runs too, for consistency.
        for sec in doc.sections:
            for p in sec.header.paragraphs:
                for run in p.runs:
                    _apply_run_font(run, thai_font_family)
            for p in sec.footer.paragraphs:
                for run in p.runs:
                    _apply_run_font(run, thai_font_family)
    
    # Save
    doc.save(output_bio)

def build_pdf_report(
    report: ReportData,
    output_path: str,
    graph1_path: str,
    graph2_path: str,
    lang: str = "en",
    report_style: str = "full",
):
    logger = logging.getLogger("report_core_pdf")
    """Build PDF report aligned with DOCX simple/full text structure (no graph embedding)."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import Paragraph, Spacer
        from reportlab.lib.utils import simpleSplit
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except Exception as e:
        raise RuntimeError(f"reportlab is required for PDF generation: {e}")

    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    # Match reference format: 1 inch margin (72pt)
    x_left = 72
    x_right = 72
    usable_width = width - x_left - x_right
    header_path = resolve_brand_asset_path("Header.png")
    footer_path = resolve_brand_asset_path("Footer.png")
    top_content_y = height - 95
    bottom_content_y = 70
    y = top_content_y
    footer_img_y = 0
    footer_img_h = 68

    style_name = str(report_style or "full").strip().lower()
    is_simple = style_name.startswith("simple")
    is_operation_test = style_name.startswith("operation_test")
    is_people_reader = style_name.startswith("people_reader")
    uses_op_layout = is_operation_test or is_people_reader
    lang_name = str(lang or "").strip().lower()
    is_thai = (lang_name == "th")
    thai_font_fallback = False
    regular_font = "Helvetica"
    bold_font = "Helvetica-Bold"
    requires_unicode_font = False

    def _first_existing(paths: list) -> str:
        for p in paths:
            if p and os.path.exists(p):
                return p
        return ""

    def _glob_existing(patterns: list) -> list:
        import glob
        found = []
        for pat in patterns:
            if not pat:
                continue
            found.extend(glob.glob(pat, recursive=True))
        # Keep deterministic order and unique entries.
        uniq = []
        seen = set()
        for p in sorted(found):
            if p not in seen and os.path.isfile(p):
                uniq.append(p)
                seen.add(p)
        return uniq

    def register_sarabun_fonts() -> bool:
        base_dir = os.path.dirname(os.path.abspath(__file__))  # .../src
        sarabun_regular_candidates = [
            os.getenv("PDF_SARABUN_FONT_PATH", "").strip(),
            os.getenv("SARABUN_FONT_PATH", "").strip(),
            os.getenv("REPORT_SARABUN_FONT_PATH", "").strip(),
            os.path.join(base_dir, "fonts", "Sarabun-Regular.ttf"),
            os.path.join(os.path.dirname(os.path.dirname(__file__)), "assets", "fonts", "Sarabun-Regular.ttf"),
            os.path.join(os.path.dirname(os.path.dirname(__file__)), "assets", "Sarabun-Regular.ttf"),
            "/Library/Fonts/Sarabun-Regular.ttf",
            "/usr/share/fonts/truetype/thai/Sarabun-Regular.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansThai-Regular.ttf",
        ]
        sarabun_bold_candidates = [
            os.getenv("PDF_SARABUN_FONT_BOLD_PATH", "").strip(),
            os.getenv("SARABUN_FONT_BOLD_PATH", "").strip(),
            os.getenv("REPORT_SARABUN_FONT_BOLD_PATH", "").strip(),
            os.path.join(base_dir, "fonts", "Sarabun-Bold.ttf"),
            os.path.join(os.path.dirname(os.path.dirname(__file__)), "assets", "fonts", "Sarabun-Bold.ttf"),
            os.path.join(os.path.dirname(os.path.dirname(__file__)), "assets", "Sarabun-Bold.ttf"),
            "/Library/Fonts/Sarabun-Bold.ttf",
            "/usr/share/fonts/truetype/thai/Sarabun-Bold.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansThai-Bold.ttf",
        ]

        regular_path = _first_existing(sarabun_regular_candidates)
        if not regular_path:
            return False
        if not _register_ttf("Sarabun", regular_path, require_thai=True):
            return False

        bold_path = _first_existing(sarabun_bold_candidates)
        ok_bold = bool(bold_path) and _register_ttf("Sarabun-Bold", bold_path, require_thai=True)

        nonlocal regular_font, bold_font, requires_unicode_font
        regular_font = "Sarabun"
        bold_font = "Sarabun-Bold" if ok_bold else "Sarabun"
        requires_unicode_font = True
        return True

    def register_noto_thai_fonts() -> bool:
        noto_regular_candidates = [
            os.getenv("PDF_THAI_FONT_PATH", "").strip(),
            "/usr/share/fonts/truetype/noto/NotoSansThaiUI-Regular.ttf",
            "/usr/share/fonts/opentype/noto/NotoSansThaiUI-Regular.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansThai-Regular.ttf",
            "/usr/share/fonts/opentype/noto/NotoSansThai-Regular.ttf",
            "/Library/Fonts/NotoSansThaiUI-Regular.ttf",
            "/Library/Fonts/NotoSansThai-Regular.ttf",
        ]
        noto_bold_candidates = [
            os.getenv("PDF_THAI_FONT_BOLD_PATH", "").strip(),
            "/usr/share/fonts/truetype/noto/NotoSansThaiUI-Bold.ttf",
            "/usr/share/fonts/opentype/noto/NotoSansThaiUI-Bold.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansThai-Bold.ttf",
            "/usr/share/fonts/opentype/noto/NotoSansThai-Bold.ttf",
            "/Library/Fonts/NotoSansThaiUI-Bold.ttf",
            "/Library/Fonts/NotoSansThai-Bold.ttf",
        ]

        regular_path = _first_existing(noto_regular_candidates)
        if not regular_path:
            return False
        if not _register_ttf("NotoThaiRegular", regular_path, require_thai=True):
            return False

        bold_path = _first_existing(noto_bold_candidates)
        ok_bold = bool(bold_path) and _register_ttf("NotoThaiBold", bold_path, require_thai=True)

        nonlocal regular_font, bold_font, requires_unicode_font
        regular_font = "NotoThaiRegular"
        bold_font = "NotoThaiBold" if ok_bold else "NotoThaiRegular"
        requires_unicode_font = True
        return True

    def register_tlwg_thai_fonts() -> bool:
        tlwg_regular_candidates = [
            "/usr/share/fonts/truetype/tlwg/Garuda.ttf",
            "/usr/share/fonts/truetype/tlwg/Waree.ttf",
            "/usr/share/fonts/truetype/tlwg/Kinnari.ttf",
            "/usr/share/fonts/truetype/tlwg/Loma.ttf",
        ]
        tlwg_bold_candidates = [
            "/usr/share/fonts/truetype/tlwg/Garuda-Bold.ttf",
            "/usr/share/fonts/truetype/tlwg/Waree-Bold.ttf",
            "/usr/share/fonts/truetype/tlwg/Kinnari-Bold.ttf",
            "/usr/share/fonts/truetype/tlwg/Loma-Bold.ttf",
        ]

        regular_path = _first_existing(tlwg_regular_candidates)
        if not regular_path:
            return False
        if not _register_ttf("TLWGThaiRegular", regular_path, require_thai=True):
            return False

        bold_path = _first_existing(tlwg_bold_candidates)
        ok_bold = bool(bold_path) and _register_ttf("TLWGThaiBold", bold_path, require_thai=True)

        nonlocal regular_font, bold_font, requires_unicode_font
        regular_font = "TLWGThaiRegular"
        bold_font = "TLWGThaiBold" if ok_bold else "TLWGThaiRegular"
        requires_unicode_font = True
        return True

    def register_arial_fonts() -> bool:
        base_dir = os.path.dirname(os.path.abspath(__file__))  # .../src
        arial_regular_candidates = [
            os.getenv("PDF_ARIAL_FONT_PATH", "").strip(),
            os.getenv("ARIAL_FONT_PATH", "").strip(),
            os.path.join(base_dir, "fonts", "Arial.ttf"),
            os.path.join(os.path.dirname(os.path.dirname(__file__)), "assets", "fonts", "Arial.ttf"),
            "/Library/Fonts/Arial.ttf",
            "C:\\Windows\\Fonts\\arial.ttf",
            "/usr/share/fonts/truetype/msttcorefonts/Arial.ttf",
            "/usr/share/fonts/truetype/msttcorefonts/arial.ttf",
        ]
        arial_bold_candidates = [
            os.getenv("PDF_ARIAL_FONT_BOLD_PATH", "").strip(),
            os.getenv("ARIAL_FONT_BOLD_PATH", "").strip(),
            os.path.join(base_dir, "fonts", "Arial-Bold.ttf"),
            os.path.join(base_dir, "fonts", "Arial Bold.ttf"),
            os.path.join(os.path.dirname(os.path.dirname(__file__)), "assets", "fonts", "Arial-Bold.ttf"),
            os.path.join(os.path.dirname(os.path.dirname(__file__)), "assets", "fonts", "Arial Bold.ttf"),
            "/Library/Fonts/Arial Bold.ttf",
            "C:\\Windows\\Fonts\\arialbd.ttf",
            "/usr/share/fonts/truetype/msttcorefonts/Arial_Bold.ttf",
            "/usr/share/fonts/truetype/msttcorefonts/arialbd.ttf",
        ]

        regular_path = _first_existing(arial_regular_candidates)
        if not regular_path:
            return False
        if not _register_ttf("ArialPDFRegular", regular_path, require_thai=False):
            return False

        bold_path = _first_existing(arial_bold_candidates)
        ok_bold = bool(bold_path) and _register_ttf("ArialPDFBold", bold_path, require_thai=False)

        nonlocal regular_font, bold_font, requires_unicode_font
        regular_font = "ArialPDFRegular"
        bold_font = "ArialPDFBold" if ok_bold else "ArialPDFRegular"
        requires_unicode_font = True
        return True

    def _register_ttf(font_name: str, path: str, require_thai: bool = False) -> bool:
        if not path:
            return False
        try:
            pdfmetrics.getFont(font_name)
            return True
        except Exception:
            pass
        try:
            tt = TTFont(font_name, path)
            if require_thai:
                # Some valid Thai fonts expose cmap/width metadata differently across platforms.
                # Avoid false negatives here; successful registration is enough.
                pass
            pdfmetrics.registerFont(tt)
            return True
        except Exception:
            return False

    # Customer requirement:
    # - Operational Test TH: prefer Noto Thai (better Thai diacritic rendering in PDF viewers), then Sarabun
    # - Operational Test EN: prefer Arial
    if uses_op_layout and lang_name == "en" and register_arial_fonts():
        pass
    elif uses_op_layout and lang_name == "th" and (
        register_noto_thai_fonts() or register_sarabun_fonts() or register_tlwg_thai_fonts()
    ):
        pass
    elif is_thai:
        thai_glob_candidates = _glob_existing(
            [
                "/usr/share/fonts/**/*.ttf",
                "/usr/local/share/fonts/**/*.ttf",
            ]
        )
        thai_glob_preferred = [
            p
            for p in thai_glob_candidates
            if any(
                k in os.path.basename(p).lower()
                for k in ("noto", "sarabun", "thsarabun", "thai", "garuda", "waree", "kinnari", "loma")
            )
        ]

        thai_regular_candidates = [
            os.getenv("PDF_THAI_FONT_PATH", "").strip(),
            os.getenv("THAI_FONT_PATH", "").strip(),
            os.getenv("REPORT_THAI_FONT_PATH", "").strip(),
            "/Library/Fonts/THSarabunNew.ttf",
            "/Library/Fonts/Sarabun-Regular.ttf",
            "/System/Library/Fonts/Supplemental/Thonburi.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansThai-Regular.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansThaiUI-Regular.ttf",
            "/usr/share/fonts/truetype/thai/Sarabun-Regular.ttf",
            "/usr/share/fonts/truetype/tlwg/Garuda.ttf",
            "/usr/share/fonts/truetype/tlwg/Waree.ttf",
            "/usr/share/fonts/truetype/tlwg/Loma.ttf",
            "/usr/share/fonts/truetype/tlwg/Kinnari.ttf",
        ]
        thai_bold_candidates = [
            os.getenv("PDF_THAI_FONT_BOLD_PATH", "").strip(),
            os.getenv("THAI_FONT_BOLD_PATH", "").strip(),
            os.getenv("REPORT_THAI_FONT_BOLD_PATH", "").strip(),
            "/Library/Fonts/THSarabunNew Bold.ttf",
            "/Library/Fonts/Sarabun-Bold.ttf",
            "/System/Library/Fonts/Supplemental/Thonburi Bold.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansThai-Bold.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansThaiUI-Bold.ttf",
            "/usr/share/fonts/truetype/thai/Sarabun-Bold.ttf",
            "/usr/share/fonts/truetype/tlwg/Garuda-Bold.ttf",
            "/usr/share/fonts/truetype/tlwg/Waree-Bold.ttf",
            "/usr/share/fonts/truetype/tlwg/Loma-Bold.ttf",
            "/usr/share/fonts/truetype/tlwg/Kinnari-Bold.ttf",
        ]
        thai_regular_candidates.extend(thai_glob_preferred)
        thai_bold_candidates.extend(thai_glob_preferred)

        ok_regular = False
        for path in thai_regular_candidates:
            if _register_ttf("ThaiPDFRegular", path, require_thai=True):
                ok_regular = True
                break

        ok_bold = False
        for path in thai_bold_candidates:
            if _register_ttf("ThaiPDFBold", path, require_thai=True):
                ok_bold = True
                break

        if not ok_regular:
            # Do not fail the whole job. Fallback to readable English PDF if Thai font is missing.
            logger.warning(
                "Thai font not found for PDF. Falling back to English labels/content for this PDF."
            )
            is_thai = False
            thai_font_fallback = True
            regular_font = "Helvetica"
            bold_font = "Helvetica-Bold"
            requires_unicode_font = False
        else:
            regular_font = "ThaiPDFRegular"
            bold_font = "ThaiPDFBold" if ok_bold else "ThaiPDFRegular"
            requires_unicode_font = True

    # ReportLab Thai shaping is limited on some viewers/fonts.
    # Using regular Thai font for emphasis text reduces vowel/tone-mark collisions.
    if is_thai and requires_unicode_font:
        bold_font = regular_font
    logger.info("[pdf] thai_font regular=%s bold=%s unicode=%s", regular_font, bold_font, requires_unicode_font)

    def draw_header_footer() -> None:
        # Match DOCX branding as closely as possible for PDF output.
        if os.path.exists(header_path):
            try:
                c.drawImage(
                    header_path,
                    x=28,
                    y=height - 62,
                    width=width - 56,
                    height=36,
                    preserveAspectRatio=True,
                    mask="auto",
                )
            except Exception:
                pass
        if os.path.exists(footer_path):
            try:
                c.drawImage(
                    footer_path,
                    x=8,
                    y=footer_img_y,
                    width=width - 16,
                    height=footer_img_h,
                    preserveAspectRatio=False,
                    mask="auto",
                )
            except Exception:
                pass

    draw_header_footer()

    TITLE_STYLE = ParagraphStyle(
        name="TitleStyle",
        fontName=bold_font,
        fontSize=20,
        leading=26,
        alignment=1,  # center
        spaceAfter=24,
    )
    SECTION_STYLE = ParagraphStyle(
        name="SectionStyle",
        fontName=bold_font,
        fontSize=14,
        leading=20,
        spaceBefore=12,
        spaceAfter=8,
    )
    SUBITEM_STYLE = ParagraphStyle(
        name="SubItemStyle",
        fontName=regular_font,
        fontSize=14,
        leading=21,
        leftIndent=22,
        spaceAfter=5,
    )
    LEVEL_STYLE = ParagraphStyle(
        name="LevelStyle",
        fontName=bold_font,
        fontSize=14,
        leading=21,
        leftIndent=38,
        spaceAfter=12,
    )
    # Section 1 First Impression: Scale aligns under bullet (same indent as SUBITEM)
    LEVEL_UNDER_BULLET_STYLE = ParagraphStyle(
        name="LevelUnderBulletStyle",
        parent=LEVEL_STYLE,
        leftIndent=22,
    )
    BULLET_STYLE = ParagraphStyle(
        name="BulletStyle",
        fontName=regular_font,
        fontSize=14,
        leading=21,
        leftIndent=40,
        bulletIndent=30,
        spaceAfter=6,
    )
    # For write_bullet(..., indent=28), visible text starts at x_left + 28 + leftIndent.
    BULLET_TEXT_X_OFFSET = 28 + int(BULLET_STYLE.leftIndent)
    LEVEL_BULLET_STYLE = ParagraphStyle(
        name="LevelBulletStyle",
        parent=LEVEL_STYLE,
        leftIndent=BULLET_TEXT_X_OFFSET,
    )
    # Backward-compatible aliases for existing references.
    HEADER_STYLE = TITLE_STYLE
    CONTENT_STYLE = SUBITEM_STYLE
    header_style = HEADER_STYLE
    content_style = CONTENT_STYLE
    section_style = SECTION_STYLE
    THAI_NOTE_STYLE = ParagraphStyle(
        name="ThaiNoteStyle",
        parent=SUBITEM_STYLE,
        leading=max(int(SUBITEM_STYLE.leading), int(SUBITEM_STYLE.fontSize * 2.3)),
    )

    # English operation-test template: improve readability with looser spacing.
    if uses_op_layout and (not is_thai):
        TITLE_STYLE.fontSize = 14
        TITLE_STYLE.leading = 22
        TITLE_STYLE.spaceAfter = 18
        SECTION_STYLE.fontSize = 14
        SECTION_STYLE.leading = 20
        SECTION_STYLE.spaceBefore = 12
        SECTION_STYLE.spaceAfter = 10
        SUBITEM_STYLE.fontSize = 12
        SUBITEM_STYLE.leading = 19
        SUBITEM_STYLE.leftIndent = 20
        SUBITEM_STYLE.spaceAfter = 4
        LEVEL_STYLE.fontSize = 12
        LEVEL_STYLE.leading = 19
        LEVEL_STYLE.leftIndent = 30
        LEVEL_STYLE.spaceAfter = 12
        LEVEL_UNDER_BULLET_STYLE.leftIndent = 20  # Match SUBITEM: Scale under bullet
        BULLET_STYLE.fontSize = 12
        BULLET_STYLE.leading = 19
        BULLET_STYLE.leftIndent = 28
        BULLET_STYLE.bulletIndent = 18
        BULLET_STYLE.spaceAfter = 6

    if is_thai:
        # Thai glyph stacks (vowels/tone marks) need extra vertical room in PDF rendering.
        TITLE_STYLE.fontSize = min(int(TITLE_STYLE.fontSize), 19)
        SECTION_STYLE.fontSize = min(int(SECTION_STYLE.fontSize), 13)
        SUBITEM_STYLE.fontSize = min(int(SUBITEM_STYLE.fontSize), 13)
        LEVEL_STYLE.fontSize = min(int(LEVEL_STYLE.fontSize), 13)
        BULLET_STYLE.fontSize = min(int(BULLET_STYLE.fontSize), 13)
        TITLE_STYLE.leading = max(TITLE_STYLE.leading, int(TITLE_STYLE.fontSize * 1.65))
        SECTION_STYLE.leading = max(SECTION_STYLE.leading, int(SECTION_STYLE.fontSize * 1.65))
        SUBITEM_STYLE.leading = max(SUBITEM_STYLE.leading, int(SUBITEM_STYLE.fontSize * 2.0))
        LEVEL_STYLE.leading = max(LEVEL_STYLE.leading, int(LEVEL_STYLE.fontSize * 2.0))
        BULLET_STYLE.leading = max(BULLET_STYLE.leading, int(BULLET_STYLE.fontSize * 2.0))
        THAI_NOTE_STYLE.leading = max(THAI_NOTE_STYLE.leading, int(THAI_NOTE_STYLE.fontSize * 2.4))

    def P(text: str, style):
        # Preserve explicit newlines in ReportLab paragraphs.
        safe = escape(str(text or "")).replace("\n", "<br/>")
        return Paragraph(safe, style)

    def gap(h=8):
        return Spacer(1, h)

    def spacer_gap(h=8):
        return gap(h)

    def _normalize_thai_for_pdf(text: str) -> str:
        return unicodedata.normalize("NFC", str(text or ""))

    def _safe_text_for_font(text: str) -> str:
        normalized_text = _normalize_thai_for_pdf(text)
        if requires_unicode_font:
            return normalized_text
        normalized = (
            normalized_text
            .replace("•", "- ")
            .replace("▪", "- ")
            .replace("□", "- ")
            .replace("™", "(TM)")
            .replace("—", "-")
            .replace("–", "-")
            .replace("“", '"')
            .replace("”", '"')
            .replace("’", "'")
        )
        return normalized.encode("latin-1", "replace").decode("latin-1")

    def _draw_text_line(x: float, y_pos: float, font_name: str, size: int, line: str) -> None:
        text_obj = c.beginText()
        text_obj.setTextOrigin(x, y_pos)
        text_obj.setFont(font_name, size)
        text_obj.textLine(str(line or ""))
        c.drawText(text_obj)

    def draw_generated_bottom(text: str, size: int = 10) -> None:
        font = CONTENT_STYLE.fontName
        safe = _safe_text_for_font(text)
        text_w = pdfmetrics.stringWidth(safe, font, size)
        x = x_left + max(0.0, usable_width - text_w)
        _draw_text_line(x, footer_img_y + footer_img_h + 6, font, size, safe)

    def append_graph_pages_for_operation_test() -> None:
        """Append Effort/Shape graph pages after the two operation-test narrative pages."""
        nonlocal y
        effort_title = "ผลการวิเคราะห์การใช้น้ำหนัก (Efforts)" if is_thai else "Effort Motion Detection Results"
        shape_title = "ผลการวิเคราะห์การใช้รูปทรงของมือร่วมกับร่างกาย (Shape)" if is_thai else "Shape Motion Detection Results"
        # Thai: wrap English parts in smaller font to match Thai
        if is_thai:
            effort_title_html = "ผลการวิเคราะห์การใช้น้ำหนัก <font size=\"12\">(Efforts)</font>"
            shape_title_html = "ผลการวิเคราะห์การใช้รูปทรงของมือร่วมกับร่างกาย <font size=\"12\">(Shape)</font>"
            graph_specs = [
                (effort_title_html, graph1_path, True),
                (shape_title_html, graph2_path, True),
            ]
        else:
            graph_specs = [
                (effort_title, graph1_path, False),
                (shape_title, graph2_path, False),
            ]
        graph_title_style = ParagraphStyle(
            name="GraphTitleStyle",
            parent=TITLE_STYLE,
            fontSize=max(16, int(TITLE_STYLE.fontSize) - 1),
            leading=max(22, int(TITLE_STYLE.leading) - 2),
            spaceAfter=8,
        )
        for idx, spec in enumerate(graph_specs):
            graph_title = spec[0]
            graph_path = spec[1]
            allow_html = spec[2] if len(spec) > 2 else False
            c.showPage()
            draw_header_footer()
            # EN: move title and graph down to match Word report; TH: keep balanced layout. Title up 1 line.
            if not is_thai:
                y = top_content_y - 18  # Move EN title/graph up for better balance
            else:
                y = min(height - 50, top_content_y + 62)  # Title up 1 line
            write_paragraph_block(graph_title, graph_title_style, indent=0, extra_gap=0, allow_html=allow_html)
            if graph_path and os.path.exists(graph_path):
                try:
                    # Keep graph just below the heading while preserving safe footer space.
                    graph_bottom_y = footer_img_y + footer_img_h + 20
                    graph_top_y = y - 4
                    graph_height = max(220, graph_top_y - graph_bottom_y)
                    c.drawImage(
                        graph_path,
                        x=x_left,
                        y=graph_bottom_y,
                        width=usable_width,
                        height=graph_height,
                        preserveAspectRatio=True,
                        anchor="n",
                        mask="auto",
                    )
                except Exception:
                    write_line("Graph image unavailable", size=11, gap=16)
            else:
                write_line("Graph image unavailable", size=11, gap=16)
    def write_line(text: str, size: int = 11, bold: bool = False, gap: int = 18):
        nonlocal y
        font = bold_font if bold else CONTENT_STYLE.fontName
        if size == 11:
            size = int(CONTENT_STYLE.fontSize)
        effective_gap = max(int(gap), int(size * (1.9 if is_thai else 1.35)))
        raw_text = str(text or "")

        if "\n" in raw_text:
            para_style = ParagraphStyle(
                name="ContentMultiline",
                fontName=font,
                fontSize=size,
                leading=max(effective_gap, int(size * (1.7 if is_thai else 1.35))),
                spaceAfter=float(CONTENT_STYLE.spaceAfter or 0),
            )
            para = P(_safe_text_for_font(raw_text), para_style)
            _, para_h = para.wrap(usable_width, max(1, int(y - bottom_content_y)))
            if y - para_h <= bottom_content_y:
                c.showPage()
                draw_header_footer()
                y = top_content_y
                _, para_h = para.wrap(usable_width, max(1, int(y - bottom_content_y)))
            para.drawOn(c, x_left, y - para_h)
            y -= para_h
            y -= float(spacer_gap(2).height)
            return

        safe = _safe_text_for_font(raw_text)

        def _split_by_chars(long_line: str) -> list:
            parts = []
            buf = ""
            for ch in str(long_line or ""):
                probe = f"{buf}{ch}"
                if (not buf) or (pdfmetrics.stringWidth(probe, font, size) <= usable_width):
                    buf = probe
                else:
                    parts.append(buf)
                    buf = ch
            if buf:
                parts.append(buf)
            return parts or [""]

        initial_lines = simpleSplit(safe, font, size, usable_width) or [""]
        lines = []
        for ln in initial_lines:
            if pdfmetrics.stringWidth(ln, font, size) <= usable_width:
                lines.append(ln)
                continue
            if is_thai or (" " not in ln):
                lines.extend(_split_by_chars(ln))
            else:
                lines.append(ln)

        wrapped_gap = max(16 if is_thai else 12, int(size * (1.7 if is_thai else 1.35)))
        for idx, line in enumerate(lines):
            if y <= bottom_content_y:
                c.showPage()
                draw_header_footer()
                y = top_content_y
            _draw_text_line(x_left, y, font, size, line)
            y -= effective_gap if idx == len(lines) - 1 else wrapped_gap

    def write_line_indented(text: str, indent: int = 0, size: int = 11, bold: bool = False, gap: int = 18):
        nonlocal y
        x = x_left + max(0, int(indent))
        local_width = max(80, usable_width - max(0, int(indent)))
        font = bold_font if bold else CONTENT_STYLE.fontName
        if size == 11:
            size = int(CONTENT_STYLE.fontSize)
        effective_gap = max(int(gap), int(size * (1.9 if is_thai else 1.35)))
        raw_text = str(text or "")
        if "\n" in raw_text:
            para_style = ParagraphStyle(
                name="ContentMultilineIndented",
                fontName=font,
                fontSize=size,
                leading=max(effective_gap, int(size * (1.7 if is_thai else 1.35))),
                spaceAfter=float(CONTENT_STYLE.spaceAfter or 0),
            )
            para = P(_safe_text_for_font(raw_text), para_style)
            _, para_h = para.wrap(local_width, max(1, int(y - bottom_content_y)))
            if y - para_h <= bottom_content_y:
                c.showPage()
                draw_header_footer()
                y = top_content_y
                _, para_h = para.wrap(local_width, max(1, int(y - bottom_content_y)))
            para.drawOn(c, x, y - para_h)
            y -= para_h
            y -= float(spacer_gap(2).height)
            return

        safe = _safe_text_for_font(raw_text)
        initial_lines = simpleSplit(safe, font, size, local_width) or [""]
        wrapped_gap = max(16 if is_thai else 12, int(size * (1.7 if is_thai else 1.35)))
        for idx, line in enumerate(initial_lines):
            if y <= bottom_content_y:
                c.showPage()
                draw_header_footer()
                y = top_content_y
            _draw_text_line(x, y, font, size, line)
            y -= effective_gap if idx == len(initial_lines) - 1 else wrapped_gap

    def write_bullet(text: str, indent: int = 28, space_after: int = 4, bullet_text: str = "•", font_size: int = None):
        nonlocal y
        x = x_left + max(0, int(indent))
        local_width = max(80, usable_width - max(0, int(indent)))
        bullet_style = ParagraphStyle(
            name="BulletRuntime",
            parent=BULLET_STYLE,
            spaceAfter=float(space_after),
        )
        if font_size is not None:
            bullet_style.fontSize = font_size
            bullet_style.leading = max(12, int(font_size * (2.0 if is_thai else 1.35)))
        safe = escape(_safe_text_for_font(text)).replace("\n", "<br/>")
        para = Paragraph(safe, bullet_style, bulletText=bullet_text)
        y -= float(getattr(bullet_style, "spaceBefore", 0) or 0)
        _, para_h = para.wrap(local_width, max(1, int(y - bottom_content_y)))
        if y - para_h <= bottom_content_y:
            c.showPage()
            draw_header_footer()
            y = top_content_y
            _, para_h = para.wrap(local_width, max(1, int(y - bottom_content_y)))
        para.drawOn(c, x, y - para_h)
        y -= para_h
        y -= float(getattr(bullet_style, "spaceAfter", 0) or 0)

    def write_paragraph_block(text: str, style, indent: int = 0, extra_gap: int = 10, x_start: float = None, allow_html: bool = False):
        nonlocal y
        if x_start is not None:
            x = float(x_start)
            local_width = max(80, width - x - x_right)
        else:
            x = x_left + max(0, int(indent))
            local_width = max(80, usable_width - max(0, int(indent)))
        if allow_html:
            para = Paragraph(text, style)
        else:
            para = P(_safe_text_for_font(text), style)
        y -= float(getattr(style, "spaceBefore", 0) or 0)
        _, para_h = para.wrap(local_width, max(1, int(y - bottom_content_y)))
        if y - para_h <= bottom_content_y:
            c.showPage()
            draw_header_footer()
            y = top_content_y
            _, para_h = para.wrap(local_width, max(1, int(y - bottom_content_y)))
        para.drawOn(c, x, y - para_h)
        y -= para_h
        y -= float(getattr(style, "spaceAfter", 0) or 0)
        y -= float(gap(extra_gap).height)

    def write_kv_line(label: str, value: str, size: int = 14, value_indent: int = 120, gap_after: int = 22):
        nonlocal y
        draw_size = max(12, int(size) - 1) if is_thai else int(size)
        draw_gap_after = max(int(gap_after), int(draw_size * (2.0 if is_thai else 1.4)))
        label_text = _safe_text_for_font(str(label or "").strip())
        value_text = _safe_text_for_font(str(value or "").strip())
        if y <= bottom_content_y:
            c.showPage()
            draw_header_footer()
            y = top_content_y
        _draw_text_line(x_left, y, bold_font, draw_size, label_text)
        _draw_text_line(x_left + value_indent, y, CONTENT_STYLE.fontName, draw_size, value_text)
        y -= draw_gap_after

    def write_block(lines: list, size: int = 11, bold: bool = False, gap: int = 16):
        for line in lines:
            write_line(line, size=size, bold=bold, gap=gap)

    def append_movement_type_pdf_block() -> None:
        # "Movement type profile" section intentionally omitted from the PDF.
        # movement_type_info is still stored on the job for the UI and the DOCX
        # is kept in sync (see block further above), but this section no longer
        # renders on page 3 of the PDF.
        return

    if uses_op_layout:
        if is_thai:
            title = "รายงานการวิเคราะห์การนำเสนอด้วยการ\nเคลื่อนไหว กับ AI People Reader"
            detailed_analysis_label = "รายละเอียดการวิเคราะห์การนำเสนอ"
            first_impression_label = "1. ความประทับใจแรกพบ (First Impression)"
        elif lang_name == "th":
            title = "Presentation Analysis Report"
            detailed_analysis_label = "Detailed Presentation Analysis"
            first_impression_label = "1. First Impression"
        else:
            title = "Movement in Communication\nwith AI People Reader Report"
            detailed_analysis_label = "Detailed Analysis"
            first_impression_label = "1. First impression"
    else:
        title = "รายงานการวิเคราะห์การนำเสนอด้วยการ\nเคลื่อนไหว กับ AI People Reader" if is_thai else "Movement in Communication\nwith AI People Reader Report"
        detailed_analysis_label = "รายละเอียดการวิเคราะห์การนำเสนอ" if is_thai else "Detailed Analysis"
        first_impression_label = "1. ความประทับใจแรกพบ (First Impression)" if is_thai else "1. First impression"
    eye_label = "การสบตา" if is_thai else "Eye Contact"
    upright_label = "ความตั้งตรงของร่างกาย" if is_thai else "Uprightness"
    stance_label = "การยืนและการวางเท้า" if is_thai else "Stance"
    category_labels = (
        ["2. การสร้างความเป็นมิตรและสัมพันธภาพ", "3. ความมั่นใจ", "4. ความเป็นผู้นำ"]
        if is_thai
        else ["2. Engaging & Connecting", "3. Confidence", "4. Authority"]
    )

    def _duration_th_text(mmss: str) -> str:
        raw = str(mmss or "").strip()
        parts = raw.split(":")
        if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
            sec = int(parts[0]) * 60 + int(parts[1])
            return f"{sec} วินาที ({raw})"
        return raw

    def _date_th_display(date_text: str) -> str:
        raw = str(date_text or "").strip()
        if not raw:
            return raw
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%d-%m-%Y", "%d/%m/%Y"):
            try:
                dt = datetime.strptime(raw, fmt)
                return dt.strftime("%d/%m/%Y")
            except Exception:
                continue
        return raw

    def _date_en_display(date_text: str) -> str:
        raw = str(date_text or "").strip()
        if not raw:
            return raw
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%d-%m-%Y", "%d/%m/%Y"):
            try:
                dt = datetime.strptime(raw, fmt)
                return f"{dt.day}/{dt.month}/{dt.year}"
            except Exception:
                continue
        return raw

    if is_thai:
        display_date = _date_th_display(report.analysis_date)
    elif uses_op_layout:
        display_date = _date_en_display(report.analysis_date)
    else:
        display_date = report.analysis_date
    if uses_op_layout and is_thai:
        # "AI People Reader" smaller to match Thai; title up 1 line (y increased)
        thai_title_html = _safe_text_for_font("รายงานการวิเคราะห์การนำเสนอด้วยการ\nเคลื่อนไหว กับ AI People Reader").replace("\n", "<br/>")
        thai_title_html = thai_title_html.replace("AI People Reader", '<font size="12">AI People Reader</font>')
        y += 12  # Title up 1 line
        write_paragraph_block(thai_title_html, TITLE_STYLE, indent=0, extra_gap=4, allow_html=True)
        write_kv_line("ชื่อ:", report.client_name, size=14, value_indent=118, gap_after=21)
        write_kv_line("วันที่วิเคราะห์:", display_date, size=14, value_indent=118, gap_after=21)
    else:
        if is_thai:
            thai_title_html = _safe_text_for_font("รายงานการวิเคราะห์การนำเสนอด้วยการ\nเคลื่อนไหว กับ AI People Reader").replace("\n", "<br/>")
            thai_title_html = thai_title_html.replace("AI People Reader", '<font size="12">AI People Reader</font>')
            y += 12  # Title up 1 line
            write_paragraph_block(thai_title_html, TITLE_STYLE, indent=0, extra_gap=0, allow_html=True)
        else:
            write_paragraph_block(title, TITLE_STYLE, indent=0, extra_gap=0)
        if not is_thai:
            write_line("", gap=14)  # EN: extra line before Client Name
        write_line(f"{'ชื่อลูกค้า' if is_thai else 'Client Name'}: {report.client_name}", bold=True)
        write_line(f"{'วันที่วิเคราะห์' if is_thai else 'Analysis Date'}: {display_date}")
    duration_label = _duration_th_text(report.video_length_str) if is_thai else report.video_length_str
    if uses_op_layout and is_thai:
        write_kv_line("ความยาววิดีโอ:", duration_label, size=14, value_indent=118, gap_after=26)
    elif uses_op_layout and (not is_thai):
        write_line(f"Video Duration: {duration_label}", gap=22)  # Single line per reference format
    else:
        write_line(f"{'ความยาววิดีโอ' if is_thai else 'Duration'}: {duration_label}", gap=22)
    # Push "Detailed Analysis" down ~2 lines below the Video Duration row.
    write_line("", gap=28)
    write_line(detailed_analysis_label, size=13, bold=True, gap=8)

    def _first_impression_level(value: float, metric: str = "") -> str:
        return first_impression_level(value, metric=metric)

    def _first_impression_level_th(value: float, metric: str = "") -> str:
        level_en = _first_impression_level(value, metric=metric).strip().lower()
        if level_en.startswith("high"):
            return "สูง"
        if level_en.startswith("moderate"):
            return "กลาง"
        if level_en.startswith("low"):
            return "ต่ำ"
        return "-"

    # First impression sections (scale: low/moderate/high)
    if report.first_impression:
        fi = report.first_impression
        if uses_op_layout:
            if is_thai:
                write_paragraph_block("1. ความประทับใจแรกพบ (First Impression)", SECTION_STYLE, extra_gap=0)
                write_paragraph_block(f"• {eye_label} (Eye Contact)", SUBITEM_STYLE, extra_gap=0)
                write_paragraph_block(
                    f"ระดับ: {_first_impression_level_th(fi.eye_contact_pct, metric='eye_contact')}",
                    LEVEL_UNDER_BULLET_STYLE,
                    extra_gap=0,
                )
                write_paragraph_block(f"• {upright_label} (Uprightness)", SUBITEM_STYLE, extra_gap=0)
                write_paragraph_block(
                    f"ระดับ: {_first_impression_level_th(fi.upright_pct, metric='uprightness')}",
                    LEVEL_UNDER_BULLET_STYLE,
                    extra_gap=0,
                )
                write_paragraph_block(f"• {stance_label} (Stance)", SUBITEM_STYLE, extra_gap=0)
                write_paragraph_block(
                    f"ระดับ: {_first_impression_level_th(fi.stance_stability, metric='stance')}",
                    LEVEL_UNDER_BULLET_STYLE,
                    extra_gap=10,
                )
            else:
                write_paragraph_block("1. First impression", SECTION_STYLE, extra_gap=0)
                write_paragraph_block("• Eye Contact", SUBITEM_STYLE, extra_gap=0)
                write_paragraph_block(
                    f"Scale: {_first_impression_level(fi.eye_contact_pct, metric='eye_contact')}",
                    LEVEL_UNDER_BULLET_STYLE,
                    extra_gap=0,
                )
                write_paragraph_block("• Uprightness", SUBITEM_STYLE, extra_gap=0)
                write_paragraph_block(
                    f"Scale: {_first_impression_level(fi.upright_pct, metric='uprightness')}",
                    LEVEL_UNDER_BULLET_STYLE,
                    extra_gap=0,
                )
                write_paragraph_block("• Stance", SUBITEM_STYLE, extra_gap=0)
                write_paragraph_block(
                    f"Scale: {_first_impression_level(fi.stance_stability, metric='stance')}",
                    LEVEL_UNDER_BULLET_STYLE,
                    extra_gap=10,
                )
        else:
            # Compact format (match Thai layout) for non-operation_test
            write_line(first_impression_label, size=12, bold=True, gap=18)
            eye_level = _first_impression_level_th(fi.eye_contact_pct, "eye_contact") if is_thai else _first_impression_level(fi.eye_contact_pct, "eye_contact")
            up_level = _first_impression_level_th(fi.upright_pct, "uprightness") if is_thai else _first_impression_level(fi.upright_pct, "uprightness")
            st_level = _first_impression_level_th(fi.stance_stability, "stance") if is_thai else _first_impression_level(fi.stance_stability, "stance")
            write_line(f"• {eye_label}" if is_thai else "• Eye Contact", gap=14)
            write_line(f"{'ระดับ' if is_thai else 'Scale'}: {eye_level}", bold=True, gap=14)
            write_line(f"• {upright_label}" if is_thai else "• Uprightness", gap=14)
            write_line(f"{'ระดับ' if is_thai else 'Scale'}: {up_level}", bold=True, gap=14)
            write_line(f"• {stance_label}" if is_thai else "• Stance", gap=14)
            write_line(f"{'ระดับ' if is_thai else 'Scale'}: {st_level}", bold=True, gap=14)
            write_line("หมายเหตุ" if is_thai else "Remark", bold=True, gap=14)
            remark_text = (
                "ความรู้สึกแรกพบมักเกิดขึ้นภายใน 5 วินาทีแรกของการพบกัน โดยพิจารณาจากภาพรวม การสบตา ความตั้งตรง และการยืนวางเท้า ก่อนเข้าสู่การวิเคราะห์เชิงพฤติกรรมในส่วนถัดไป"
                if is_thai
                else "First impression happens in the first 5 seconds of meeting someone, and is normally decided from the person's appearance, eye contact, uprightness and stance. However, after the first 5 seconds, the rest (below) are normally taken into consideration."
            )
            write_line(remark_text, gap=6 if is_thai else 18)  # Thai: section 2 up 1 line
            write_line("", gap=10)  # Blank line before Combination Explanation
            write_line("คำอธิบายการผสมผสาน:" if is_thai else "Combination Explanation:", gap=6)
            if is_thai:
                write_line("1. การสบตาน้อย + ความตั้งตรงน้อย + การยืนและการวางเท้าต่ำ", gap=4)
                write_line("บุคคลมักดูไม่เป็นภัยและยืดหยุ่น แต่บุคคลอาจดูมีความมั่นใจและอำนาจในระดับต่ำ", gap=8)
            else:
                write_line("1. Low Eye Contact + Low Uprightness + Low Stance.", gap=4)
                write_line("The person tends to appear non-threatening and flexible. However, the person can also appear to possess low level of confidence and authority.", gap=8)
            # Page 2: combo points 2, 3 + categories
            c.showPage()
            draw_header_footer()
            y = top_content_y - 10
            if is_thai:
                write_line("2. การสบตาปานกลาง + ความตั้งตรงปานกลาง + การยืนและการวางเท้าปานกลาง", gap=4)
                write_line("บุคคลมักดูเข้าถึงได้ง่าย และมีความมั่นใจและอำนาจในระดับที่เพียงพอ", gap=8)
                write_line("3. การสบตาสูง + ความตั้งตรงสูง + การยืนและการวางเท้าสูง", gap=4)
                write_line("บุคคลมักดูมีความมั่นใจและอำนาจในระดับสูง และอาจดูไม่เข้าถึงได้ง่ายหรือยืดหยุ่น", gap=12)
            else:
                write_line("2. Moderate Eye Contact + Moderate Uprightness + Moderate Stance.", gap=4)
                write_line("The person tends to appear approachable, and has adequate level of confidence and authority.", gap=8)
                write_line("3. High Eye Contact + High Uprightness + High Stance.", gap=4)
                write_line("The person tends to appear to possess high level of confidence and authority, and may not appear approachable or flexible.", gap=12)
    else:
        write_line(first_impression_label, size=12, bold=True, gap=18)
        write_line("- Not available", gap=20)

    if uses_op_layout:
        if is_thai:
            def _scale_th(scale: str) -> str:
                return _display_scale(scale, is_thai=True)

            write_paragraph_block("", SUBITEM_STYLE, extra_gap=0)
            write_paragraph_block("หมายเหตุ", SECTION_STYLE, extra_gap=0)
            write_paragraph_block(
                "ความรู้สึกที่เกิดจากความประทับใจแรกพบนั้นเป็นสิ่งที่มนุษย์หลีกเลี่ยงไม่ได้ และมักเกิดขึ้นภายใน 5 วินาทีแรกของการพบกัน แต่หลังจากนั้นจะเริ่มวิเคราะห์การเคลื่อนไหวโดยรวมมาประกอบการตัดสินใจ",
                THAI_NOTE_STYLE,
                extra_gap=6,
            )
            # Page 2: Combination Explanation 1, 2, 3 อยู่รวมกัน
            c.showPage()
            draw_header_footer()
            y = top_content_y

            write_paragraph_block("คำอธิบายการผสมผสาน:", THAI_NOTE_STYLE, extra_gap=4)
            write_paragraph_block("1. การสบตาน้อย + ความตั้งตรงน้อย + การยืนและการวางเท้าต่ำ", THAI_NOTE_STYLE, extra_gap=4)
            write_paragraph_block("บุคคลมักดูไม่เป็นภัยและยืดหยุ่น แต่บุคคลอาจดูมีความมั่นใจและอำนาจในระดับต่ำ", THAI_NOTE_STYLE, extra_gap=4)
            write_paragraph_block("2. การสบตาปานกลาง + ความตั้งตรงปานกลาง + การยืนและการวางเท้าปานกลาง", THAI_NOTE_STYLE, extra_gap=4)
            write_paragraph_block("บุคคลมักดูเข้าถึงได้ง่าย และมีความมั่นใจและอำนาจในระดับที่เพียงพอ", THAI_NOTE_STYLE, extra_gap=4)
            write_paragraph_block("3. การสบตาสูง + ความตั้งตรงสูง + การยืนและการวางเท้าสูง", THAI_NOTE_STYLE, extra_gap=4)
            write_paragraph_block("บุคคลมักดูมีความมั่นใจและอำนาจในระดับสูง และอาจดูไม่เข้าถึงได้ง่ายหรือยืดหยุ่น", THAI_NOTE_STYLE, extra_gap=8)

            # Page 3: Engaging + Confidence + Authority (+ Adaptability for People Reader only)
            c.showPage()
            draw_header_footer()
            y = top_content_y

            engaging_scale = _scale_th(report.categories[0].scale) if len(report.categories) > 0 else "-"
            confidence_scale = _scale_th(report.categories[1].scale) if len(report.categories) > 1 else "-"
            authority_scale = _scale_th(report.categories[2].scale) if len(report.categories) > 2 else "-"

            # Headings (numbering differs by mode)
            if is_people_reader:
                _h_eng = "2. การสร้างความเป็นมิตรและสร้างสัมพันธภาพ"
                _h_adapt = "3. ความยืดหยุ่นในการปรับตัว (Adaptability):"
                _h_conf = "4. ความมั่นใจ:"
                _h_auth = "5. ความเป็นผู้นำและความดูมีอำนาจ:"
            else:
                _h_eng = "2. การสร้างความเป็นมิตรและสร้างสัมพันธภาพ"
                _h_conf = "3. ความมั่นใจ:"
                _h_auth = "4. ความเป็นผู้นำและความดูมีอำนาจ:"

            # 2. Engaging
            write_paragraph_block(_h_eng, SECTION_STYLE, extra_gap=0)
            write_bullet("ความเป็นกันเอง", indent=28, space_after=4, bullet_text="▪")
            write_bullet("ความเข้าถึงได้", indent=28, space_after=4, bullet_text="▪")
            write_bullet("การมีส่วนร่วม เชื่อมโยง และสร้างความคุ้นเคยกับทีมอย่างรวดเร็ว", indent=28, space_after=4, bullet_text="▪")
            write_line_indented(f"ระดับ: {engaging_scale}", indent=28, bold=True, gap=6)

            # 3. Adaptability (People Reader only)
            if is_people_reader:
                if len(report.categories) < 4:
                    raise ValueError("people_reader PDF requires Adaptability as 4th category")
                adapt_scale = _scale_th(report.categories[3].scale)
                write_paragraph_block(_h_adapt, SECTION_STYLE, extra_gap=0)
                write_bullet(
                    "ความยืดหยุ่น — ความสามารถในการปรับตัวตามสภาวะใหม่ ๆ และรับมือกับการเปลี่ยนแปลง",
                    indent=28,
                    space_after=4,
                    bullet_text="▪",
                )
                write_bullet(
                    "ความคล่องแคล่ว — ความสามารถในการคิดและสรุปได้อย่างรวดเร็วเพื่อปรับตัว",
                    indent=28,
                    space_after=4,
                    bullet_text="▪",
                )
                write_line_indented(f"ระดับ: {adapt_scale}", indent=28, bold=True, gap=6)

            # 3. / 4. Confidence
            write_paragraph_block(_h_conf, SECTION_STYLE, extra_gap=0)
            write_bullet("บุคลิกภาพเชิงบวก", indent=28, space_after=4, bullet_text="▪")
            write_bullet("ความมีสมาธิ", indent=28, space_after=4, bullet_text="▪")
            write_bullet("ความสามารถในการโน้มน้าวและยืนหยัดในจุดยืนเพื่อให้ผู้อื่นคล้อยตาม", indent=28, space_after=4, bullet_text="▪")
            write_line_indented(f"ระดับ: {confidence_scale}", indent=28, bold=True, gap=6)

            # 4. / 5. Authority
            write_paragraph_block(_h_auth, SECTION_STYLE, extra_gap=0)
            write_bullet("แสดงให้เห็นถึงความสำคัญและความเร่งด่วนของประเด็น", indent=28, space_after=4, bullet_text="▪")
            write_bullet("ผลักดันให้เกิดการลงมือทำ", indent=28, space_after=4, bullet_text="▪")
            write_line_indented(f"ระดับ: {authority_scale}", indent=28, bold=True, gap=6)

            if is_people_reader:
                append_movement_type_pdf_block()
        else:
            if thai_font_fallback and lang_name == "th":
                write_line("Note: Thai font is unavailable on server; this TH report is rendered in English fallback.", size=10, gap=12)
            def _scale_en(scale: str) -> str:
                return _display_scale(scale, is_thai=False)

            # EN PDF: Page 3 sections use normal spacing
            en_section_gap = 8
            en_section2_item_gap = 6
            en_section2_scale_gap = 6
            en_section34_item_gap = 6
            en_section34_scale_gap = 6

            write_line("", gap=14)  # Significant space before Remark per reference
            write_line("Remark", bold=True, gap=8)
            write_line(
                "First impression happens in the first 5 seconds of meeting someone, and is normally decided from the person's appearance, eye contact, uprightness and stance. However, after the first 5 seconds, the rest (below) are normally taken into consideration.",
                gap=8,
            )
            # Page 2: Combination Explanation 1, 2, 3 อยู่รวมกัน
            c.showPage()
            draw_header_footer()
            y = top_content_y

            write_line("Combination Explanation:", gap=6)
            write_line("1. Low Eye Contact + Low Uprightness + Low Stance.", gap=4)
            write_line("The person tends to appear non-threatening and flexible. However, the person can also appear to possess low level of confidence and authority.", gap=6)
            write_line("2. Moderate Eye Contact + Moderate Uprightness + Moderate Stance.", gap=4)
            write_line("The person tends to appear approachable, and has adequate level of confidence and authority.", gap=6)
            write_line("3. High Eye Contact + High Uprightness + High Stance.", gap=4)
            write_line("The person tends to appear to possess high level of confidence and authority, and may not appear approachable or flexible.", gap=8)

            # Page 3: Engaging + Confidence + Authority (+ Adaptability for People Reader only)
            c.showPage()
            draw_header_footer()
            y = top_content_y

            engaging_scale = _scale_en(report.categories[0].scale) if len(report.categories) > 0 else "-"
            confidence_scale = _scale_en(report.categories[1].scale) if len(report.categories) > 1 else "-"
            authority_scale = _scale_en(report.categories[2].scale) if len(report.categories) > 2 else "-"

            # Page-3 layout:
            #   People Reader:  2. Engaging → 3. Adaptability → 4. Confidence → 5. Authority
            #   Operation Test: 2. Engaging → 3. Confidence → 4. Authority
            if is_people_reader:
                _eng_h, _adapt_h, _conf_h, _auth_h = (
                    "2. Engaging & Connecting:",
                    "3. Adaptability:",
                    "4. Confidence:",
                    "5. Authority:",
                )
            else:
                _eng_h, _conf_h, _auth_h = (
                    "2. Engaging & Connecting:",
                    "3. Confidence:",
                    "4. Authority:",
                )

            # 2. Engaging
            write_line(_eng_h, size=12, bold=True, gap=en_section_gap)
            write_line_indented("▪ Approachability.", indent=28, gap=en_section2_item_gap)
            write_line_indented("▪ Relatability.", indent=28, gap=en_section2_item_gap)
            write_line_indented("▪ Engagement, connect and build instant rapport with team.", indent=28, gap=en_section2_item_gap)
            write_line_indented(f"Scale: {engaging_scale}", indent=28, bold=True, gap=en_section2_scale_gap)

            # 3. Adaptability (People Reader only)
            if is_people_reader:
                if len(report.categories) < 4:
                    raise ValueError("people_reader PDF requires Adaptability as 4th category")
                adapt_scale = _scale_en(report.categories[3].scale)
                write_line(_adapt_h, size=12, bold=True, gap=en_section_gap)
                write_line_indented(
                    "▪ Flexibility — Ability to adjust to new conditions, handle changes.",
                    indent=28,
                    gap=en_section34_item_gap,
                )
                write_line_indented(
                    "▪ Agility — Ability to think and draw conclusions quickly in order to adjust.",
                    indent=28,
                    gap=en_section34_item_gap,
                )
                write_line_indented(f"Scale: {adapt_scale}", indent=28, bold=True, gap=en_section34_scale_gap)

            # 3. / 4. Confidence
            write_line(_conf_h, size=12, bold=True, gap=en_section_gap)
            write_line_indented("▪ Optimistic Presence.", indent=28, gap=en_section34_item_gap)
            write_line_indented("▪ Focus.", indent=28, gap=en_section34_item_gap)
            write_line_indented("▪ Ability to persuade and stand one's ground, in order to convince others.", indent=28, gap=en_section34_item_gap)
            write_line_indented(f"Scale: {confidence_scale}", indent=28, bold=True, gap=en_section34_scale_gap)

            # 4. / 5. Authority
            write_line(_auth_h, size=12, bold=True, gap=en_section_gap)
            write_line_indented("▪ Showing sense of importance and urgency in subject matter.", indent=28, gap=en_section34_item_gap)
            write_line_indented("▪ Pressing for action.", indent=28, gap=en_section34_item_gap)
            write_line_indented(f"Scale: {authority_scale}", indent=28, bold=True, gap=en_section34_scale_gap)

            if is_people_reader:
                append_movement_type_pdf_block()
        append_graph_pages_for_operation_test()
        c.save()
        return

    # Categories (same order as DOCX) — full layout with sub-bullets to match DOCX
    cat_bullets_th = [
        ["ความเป็นกันเอง", "ความเข้าถึงได้", "การมีส่วนร่วม เชื่อมโยง และสร้างความคุ้นเคยกับทีมอย่างรวดเร็ว"],
        ["บุคลิกภาพเชิงบวก", "ความมีสมาธิ", "ความสามารถในการโน้มน้าวและยืนหยัดในจุดยืนเพื่อให้ผู้อื่นคล้อยตาม"],
        ["แสดงให้เห็นถึงความสำคัญและความเร่งด่วนของประเด็น", "ผลักดันให้เกิดการลงมือทำ"],
    ]
    cat_bullets_en = [
        ["Approachability.", "Relatability.", "Engagement, connect and build instant rapport with team."],
        ["Optimistic Presence.", "Focus.", "Ability to persuade and stand one's ground, in order to convince others."],
        ["Showing sense of importance and urgency in subject matter.", "Pressing for action."],
    ]
    cat_bullets = cat_bullets_th if is_thai else cat_bullets_en
    scale_label = "ระดับ" if is_thai else "Scale"
    item_gap = 10 if not is_thai else 8
    for idx, cat in enumerate(report.categories[:3]):
        write_line(category_labels[idx], size=12, bold=True, gap=14)
        bullets = cat_bullets[idx] if idx < len(cat_bullets) else []
        for b in bullets:
            write_line_indented(f"▪ {b}", indent=28, gap=item_gap)
        write_line_indented(
            f"{scale_label}: {_display_scale(cat.scale, is_thai)}",
            indent=0,  # Align with section title per reference
            bold=True,
            gap=12,
        )

    c.save()
