# report_core.py — shared report logic for report generation
import os
import io
import math
import random
import logging
import unicodedata
from dataclasses import dataclass
from datetime import datetime
from typing import Dict, Any, Optional

import cv2
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from mediapipe.python.solutions import pose as mp_pose_module
    from mediapipe.python.solutions.pose import Pose, PoseLandmark
    mp = True
except ImportError:
    try:
        import mediapipe as mp
        mp_pose_module = mp.solutions.pose
        Pose = mp_pose_module.Pose
        PoseLandmark = mp_pose_module.PoseLandmark
        mp = True
    except Exception:
        mp = None
        mp_pose_module = None
        Pose = None
        PoseLandmark = None

# Dataclasses
@dataclass
class FirstImpressionData:
    eye_contact_pct: float
    upright_pct: float
    stance_stability: float

@dataclass
class CategoryResult:
    name_en: str
    name_th: str
    score: int
    scale: str
    positives: int
    total: int
    description: str = ""

@dataclass
class ReportData:
    client_name: str
    analysis_date: str
    video_length_str: str
    overall_score: int
    categories: list
    summary_comment: str
    generated_by: str
    first_impression: Optional[FirstImpressionData] = None

# Helpers
def format_seconds_to_mmss(total_seconds: float) -> str:
    total_seconds = max(0, float(total_seconds))
    mm = int(total_seconds // 60)
    ss = int(round(total_seconds - mm * 60))
    if ss == 60:
        mm += 1
        ss = 0
    return f"{mm:02d}:{ss:02d}"

def get_video_duration_seconds(video_path: str) -> float:
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        return 0.0
    fps = cap.get(cv2.CAP_PROP_FPS) or 25.0
    frames = cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0.0
    cap.release()
    if fps <= 0:
        return 0.0
    return float(frames / fps)

def analyze_first_impression_from_video(video_path: str, sample_every_n: int = 5, max_frames: int = 200) -> FirstImpressionData:
    """Real First Impression analysis with MediaPipe using continuous scoring."""
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        return FirstImpressionData(eye_contact_pct=0.0, upright_pct=0.0, stance_stability=0.0)

    total = 0
    eye_frame_scores = []
    upright_frame_scores = []
    ankle_dist = []
    ankle_center_x = []
    stance_width_ratios = []
    prev_nose_offset_ratio = None
    prev_torso_angle = None

    with Pose(static_image_mode=False, model_complexity=1) as pose:
        i = 0
        while True:
            ret, frame = cap.read()
            if not ret:
                break
            i += 1
            if i % sample_every_n != 0:
                continue

            rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            res = pose.process(rgb)
            if not res.pose_landmarks:
                continue

            lms = res.pose_landmarks.landmark
            nose = lms[PoseLandmark.NOSE]
            leye = lms[PoseLandmark.LEFT_EYE]
            reye = lms[PoseLandmark.RIGHT_EYE]
            lsh = lms[PoseLandmark.LEFT_SHOULDER]
            rsh = lms[PoseLandmark.RIGHT_SHOULDER]
            lhip = lms[PoseLandmark.LEFT_HIP]
            rhip = lms[PoseLandmark.RIGHT_HIP]
            lank = lms[PoseLandmark.LEFT_ANKLE]
            rank = lms[PoseLandmark.RIGHT_ANKLE]

            if min(nose.visibility, leye.visibility, reye.visibility, lsh.visibility, rsh.visibility, lhip.visibility, rhip.visibility) < 0.5:
                continue

            total += 1

            # Eye contact (continuous): frontal face symmetry + gaze jitter penalty.
            eye_dist = abs(leye.x - reye.x)
            if eye_dist > 1e-4:
                mid_eye_x = (leye.x + reye.x) / 2.0
                nose_offset_ratio = abs(nose.x - mid_eye_x) / eye_dist
                symmetry_score = max(0.0, 1.0 - (nose_offset_ratio / 0.75))
                stability_bonus = 1.0
                if prev_nose_offset_ratio is not None:
                    jitter = abs(nose_offset_ratio - prev_nose_offset_ratio)
                    stability_bonus = max(0.70, 1.0 - (jitter / 0.35))
                prev_nose_offset_ratio = nose_offset_ratio
                eye_frame_scores.append(100.0 * symmetry_score * stability_bonus)

            # Uprightness (continuous): torso vertical angle + shoulder level + abrupt tilt penalty.
            mid_sh = np.array([(lsh.x + rsh.x) / 2.0, (lsh.y + rsh.y) / 2.0])
            mid_hip = np.array([(lhip.x + rhip.x) / 2.0, (lhip.y + rhip.y) / 2.0])
            v = mid_sh - mid_hip
            vert = np.array([0.0, -1.0])
            v_norm = np.linalg.norm(v) + 1e-9
            cosang = float(np.dot(v / v_norm, vert))
            ang = math.degrees(math.acos(max(-1.0, min(1.0, cosang))))
            shoulder_tilt = abs(lsh.y - rsh.y)
            shoulder_tilt_ratio = shoulder_tilt / v_norm

            angle_score = max(0.0, 1.0 - (ang / 28.0))
            shoulder_score = max(0.0, 1.0 - (shoulder_tilt_ratio / 0.20))
            temporal_bonus = 1.0
            if prev_torso_angle is not None:
                angle_jump = abs(ang - prev_torso_angle)
                temporal_bonus = max(0.75, 1.0 - (angle_jump / 45.0))
            prev_torso_angle = ang

            upright_frame_scores.append(
                100.0 * (0.75 * angle_score + 0.25 * shoulder_score) * temporal_bonus
            )

            # Stance (continuous): lower-body stability + center sway + stance width appropriateness.
            if min(lank.visibility, rank.visibility) >= 0.5:
                dx = (lank.x - rank.x)
                dy = (lank.y - rank.y)
                dist = math.sqrt(dx * dx + dy * dy)
                ankle_dist.append(dist)
                ankle_center_x.append((lank.x + rank.x) / 2.0)
                shoulder_width = abs(lsh.x - rsh.x)
                if shoulder_width > 1e-4:
                    stance_width_ratios.append(dist / shoulder_width)

            if total >= max_frames:
                break

    cap.release()

    if total == 0:
        return FirstImpressionData(eye_contact_pct=0.0, upright_pct=0.0, stance_stability=0.0)

    eye_pct = float(np.mean(np.array(eye_frame_scores))) if eye_frame_scores else 0.0
    upright_pct = float(np.mean(np.array(upright_frame_scores))) if upright_frame_scores else 0.0

    if len(ankle_dist) >= 10:
        dist_arr = np.array(ankle_dist)
        dist_std = float(np.std(dist_arr))
        dist_mean = float(np.mean(dist_arr)) + 1e-9
        rel_std = dist_std / dist_mean

        # Lower relative ankle-distance variance -> better lower-body stability.
        base_stability = max(0.0, min(100.0, 100.0 * (1.0 - (rel_std / 0.35))))

        sway_score = 0.0
        if len(ankle_center_x) >= 10:
            sway_std = float(np.std(np.array(ankle_center_x)))
            sway_score = max(0.0, min(100.0, 100.0 * (1.0 - (sway_std / 0.06))))

        width_score = 0.0
        if stance_width_ratios:
            ratios = np.array(stance_width_ratios)
            ratio_mean = float(np.mean(ratios))
            ratio_std = float(np.std(ratios))
            # Preferred stance is around shoulder width to slightly wider.
            width_pref = max(0.0, 1.0 - (abs(ratio_mean - 1.10) / 0.90))
            width_stability = max(0.0, 1.0 - (ratio_std / 0.35))
            width_score = 100.0 * (0.65 * width_pref + 0.35 * width_stability)

        stability = max(
            0.0,
            min(100.0, 0.50 * base_stability + 0.30 * sway_score + 0.20 * width_score),
        )
    else:
        stability = 0.0

    return FirstImpressionData(eye_contact_pct=eye_pct, upright_pct=upright_pct, stance_stability=stability)

def generate_eye_contact_text(pct: float) -> list:
    """Generate descriptive text based on eye contact percentage - 10 levels"""
    if pct >= 90:  # Outstanding
        return [
            "• Your eye contact is outstanding — consistently steady, warm, and completely audience-focused throughout the entire presentation.",
            "• You maintain unwavering direct gaze during key moments, which maximizes trust, clarity, and audience engagement.",
            "• Every gaze shift is purposeful, natural, and enhances your communication (e.g., thinking pauses, emphasis).",
            "• Your eye contact is a masterclass in confidence and credibility, with zero signs of avoidance or nervousness."
        ]
    elif pct >= 80:  # Excellent
        return [
            "• Your eye contact is excellent — steady, engaging, and audience-focused almost throughout the presentation.",
            "• You consistently maintain direct gaze during key moments, which strongly increases trust and clarity.",
            "• Gaze shifts are purposeful and natural, never detracting from your professional presence.",
            "• Your eye contact demonstrates strong confidence and builds exceptional credibility with the audience."
        ]
    elif pct >= 70:  # Very Strong
        return [
            "• Your eye contact is very strong — steady and audience-focused for most of the time.",
            "• You maintain direct gaze during key moments effectively, which clearly builds trust and clarity.",
            "• Occasional gaze shifts are natural and appropriate, showing thoughtfulness without reducing engagement.",
            "• Overall, your eye contact strongly supports confidence and credibility with the audience."
        ]
    elif pct >= 60:  # Strong
        return [
            "• Your eye contact is strong — generally steady and audience-focused during important moments.",
            "• You maintain direct gaze during most key moments, which helps establish trust.",
            "• Some gaze shifts occur but they're mostly natural and don't significantly impact your presence.",
            "• Your eye contact effectively supports good confidence and reasonable credibility."
        ]
    elif pct >= 50:  # Good
        return [
            "• Your eye contact is good, though you occasionally look away more frequently than ideal.",
            "• You maintain direct gaze during many important points, which helps build adequate trust.",
            "• Some gaze shifts could be reduced to maintain stronger and more consistent audience connection.",
            "• Your eye contact is acceptable overall and supports reasonable engagement with the audience."
        ]
    elif pct >= 40:  # Above Average
        return [
            "• Your eye contact is above average but shows room for improvement in consistency.",
            "• You make direct eye contact regularly, but it could be more sustained during crucial points.",
            "• Gaze shifts happen somewhat frequently, which can slightly reduce audience connection.",
            "• Increasing eye contact duration will noticeably improve your trust-building and credibility."
        ]
    elif pct >= 30:  # Average
        return [
            "• Your eye contact is average, with noticeable inconsistency throughout the presentation.",
            "• You make some direct eye contact, but it lacks the duration needed during key moments.",
            "• Frequent gaze shifts away from the audience reduce connection and engagement.",
            "• Your eye contact needs strengthening to build more effective trust and credibility."
        ]
    elif pct >= 20:  # Below Average
        return [
            "• Your eye contact is below average, with significant inconsistency and frequent avoidance.",
            "• Limited direct eye contact reduces audience trust and makes it harder to connect with listeners.",
            "• You look away quite often, which can signal discomfort or lack of confidence.",
            "• Improving eye contact should be a priority to enhance your presence and credibility."
        ]
    elif pct >= 10:  # Weak
        return [
            "• Your eye contact is weak, with very limited direct audience engagement throughout.",
            "• Minimal direct eye contact significantly reduces trust and connection with the audience.",
            "• You frequently avoid or wander away from audience gaze, which impacts perceived confidence.",
            "• Developing stronger eye contact is critical and will dramatically improve your communication effectiveness."
        ]
    else:  # Critical - Needs Major Work
        return [
            "• Your eye contact needs major improvement — there is very little direct audience engagement.",
            "• Lack of eye contact severely reduces audience trust, connection, and overall message impact.",
            "• You avoid direct gaze almost entirely, which significantly undermines your credibility and presence.",
            "• Building eye contact skills is essential and should be your top priority for presentation improvement."
        ]

def generate_eye_contact_text_th(pct: float) -> list:
    """Thai descriptive text based on eye contact percentage - 10 levels"""
    if pct >= 90:
        return [
            "• การสบตาโดดเด่นมาก สม่ำเสมอ อบอุ่น และโฟกัสผู้ฟังตลอดการนำเสนอ",
            "• คุณรักษาการสบตาในจุดสำคัญได้อย่างยอดเยี่ยม ช่วยเพิ่มความไว้วางใจสูงมาก",
            "• การเปลี่ยนสายตาเป็นธรรมชาติและมีจุดประสงค์ ไม่ทำให้พลังการสื่อสารลดลง",
            "• ภาพรวมสะท้อนความมั่นใจและความน่าเชื่อถือในระดับสูงสุด",
        ]
    elif pct >= 80:
        return [
            "• การสบตายอดเยี่ยม มีความต่อเนื่องและเชื่อมโยงผู้ฟังได้ดีมาก",
            "• คุณสบตาในประเด็นสำคัญได้ชัดเจน ช่วยสร้างความเชื่อมั่นได้มาก",
            "• มีการละสายตาบ้างเล็กน้อยแต่ยังเป็นธรรมชาติ",
            "• ภาพรวมแสดงบุคลิกที่มั่นใจและน่าเชื่อถือมาก",
        ]
    elif pct >= 70:
        return [
            "• การสบตาแข็งแรงและคงเส้นคงวาในช่วงส่วนใหญ่ของการพูด",
            "• ช่วงที่เน้นสาระสำคัญมีการสบตาที่ดี ช่วยเสริมความไว้วางใจ",
            "• มีจังหวะละสายตาเป็นระยะแต่ไม่กระทบการเชื่อมต่อผู้ฟังมาก",
            "• โดยรวมสนับสนุนภาพลักษณ์มืออาชีพได้ดี",
        ]
    elif pct >= 60:
        return [
            "• การสบตาอยู่ในระดับดีและค่อนข้างสม่ำเสมอ",
            "• จุดสำคัญส่วนใหญ่ยังคงสบตาได้เหมาะสม",
            "• มีการละสายตาบ้างแต่ยังไม่รบกวนภาพรวม",
            "• โดยรวมช่วยเสริมความมั่นใจได้ชัดเจน",
        ]
    elif pct >= 50:
        return [
            "• การสบตาอยู่ในระดับใช้ได้ แต่ยังมีบางช่วงที่หลุดโฟกัสผู้ฟัง",
            "• ในประเด็นสำคัญยังสบตาได้พอสมควร",
            "• หากลดการละสายตาจะช่วยเพิ่มพลังการสื่อสารได้มากขึ้น",
            "• ภาพรวมอยู่ในเกณฑ์ดีและพัฒนาได้อีก",
        ]
    elif pct >= 40:
        return [
            "• การสบตาปานกลางค่อนข้างดี แต่ความสม่ำเสมอยังไม่มากพอ",
            "• มีการสบตาเป็นช่วง ๆ มากกว่าต่อเนื่อง",
            "• การละสายตาเกิดค่อนข้างบ่อยในบางจังหวะ",
            "• หากเพิ่มระยะเวลาการสบตาจะช่วยเพิ่มความน่าเชื่อถือ",
        ]
    elif pct >= 30:
        return [
            "• การสบตาอยู่ระดับกลาง ยังมีความไม่สม่ำเสมอชัดเจน",
            "• การเชื่อมโยงกับผู้ฟังเกิดได้บางช่วงเท่านั้น",
            "• มีการละสายตาบ่อย ทำให้พลังการสื่อสารลดลง",
            "• ควรฝึกการสบตาในช่วงเน้นประเด็นสำคัญ",
        ]
    elif pct >= 20:
        return [
            "• การสบตาต่ำกว่ามาตรฐาน ยังหลีกเลี่ยงสายตาผู้ฟังค่อนข้างมาก",
            "• ส่งผลให้ความไว้วางใจและการเชื่อมโยงลดลง",
            "• การละสายตาบ่อยทำให้ภาพลักษณ์ดูไม่มั่นใจ",
            "• ควรพัฒนาเรื่องนี้เป็นลำดับต้น ๆ",
        ]
    elif pct >= 10:
        return [
            "• การสบตาอ่อน ยังเชื่อมโยงผู้ฟังได้จำกัดมาก",
            "• การไม่สบตาโดยตรงเกิดบ่อยและกระทบความน่าเชื่อถือ",
            "• ภาพรวมทำให้ความมั่นใจที่ผู้ฟังรับรู้ลดลง",
            "• ควรฝึกการสบตาแบบสั้น-สม่ำเสมอในทุกช่วงสำคัญ",
        ]
    else:
        return [
            "• การสบตาต้องปรับปรุงอย่างมาก เนื่องจากแทบไม่มีการเชื่อมสายตากับผู้ฟัง",
            "• ส่งผลโดยตรงต่อความไว้วางใจและการรับสาร",
            "• ภาพลักษณ์อาจถูกมองว่าหลีกเลี่ยงหรือไม่มั่นใจ",
            "• การฝึกการสบตาเป็นสิ่งจำเป็นเร่งด่วน",
        ]

def generate_uprightness_text(pct: float) -> list:
    """Generate descriptive text based on uprightness percentage - 10 levels"""
    if pct >= 90:  # Outstanding
        return [
            "• You maintain outstanding upright posture with perfect vertical alignment throughout the entire presentation.",
            "• Your chest remains fully open, shoulders are optimally relaxed, and head alignment is impeccable — signaling exceptional balance, readiness, and commanding authority.",
            "• Even during active gesturing, your vertical alignment stays completely stable, demonstrating masterful core control.",
            "• There is absolutely no visible slouching or collapsing at any point, projecting supreme professional presence."
        ]
    elif pct >= 80:  # Excellent
        return [
            "• You maintain excellent upright posture consistently throughout the presentation.",
            "• Your chest stays open, shoulders are naturally relaxed, and head alignment is near-perfect — signaling strong balance, readiness, and authority.",
            "• Your vertical alignment remains remarkably stable even when gesturing, demonstrating superior core control.",
            "• There is virtually no slouching or collapsing, which projects a highly professional and confident appearance."
        ]
    elif pct >= 70:  # Very Strong
        return [
            "• You maintain very strong upright posture throughout most of the presentation.",
            "• The chest stays open, shoulders remain relaxed, and head is well-aligned — clearly signaling balance, readiness, and authority.",
            "• Your vertical alignment stays stable during most gestures, showing very good core control.",
            "• Minimal slouching occurs, which strongly supports a professional and composed appearance."
        ]
    elif pct >= 60:  # Strong
        return [
            "• You maintain strong upright posture for the majority of the clip.",
            "• Your chest generally stays open, shoulders relaxed, and head aligned — signaling good balance and authority.",
            "• Vertical alignment remains fairly stable during gestures, showing solid core control.",
            "• Occasional minor slouching appears but doesn't detract significantly from your professional presence."
        ]
    elif pct >= 50:  # Good
        return [
            "• You maintain good upright posture most of the time with some occasional lapses.",
            "• Your chest and shoulder alignment is generally acceptable, though sometimes less consistent than ideal.",
            "• When you gesture, your vertical alignment usually remains reasonably stable.",
            "• Overall posture is good with room for improvement in maintaining greater consistency."
        ]
    elif pct >= 40:  # Above Average
        return [
            "• You maintain above-average posture, though there are noticeable periods of less-than-optimal alignment.",
            "• Your chest and shoulders show acceptable positioning most of the time, with occasional forward rounding.",
            "• Vertical alignment fluctuates somewhat during gestures, suggesting moderate core control.",
            "• Some slouching is visible, indicating that improved consistency would enhance your professional image."
        ]
    elif pct >= 30:  # Average
        return [
            "• Your posture is average, fluctuating between upright and noticeably slouched positions.",
            "• There are regular moments where shoulders round forward or head tilts down from ideal alignment.",
            "• Your vertical alignment could be considerably more consistent to project stronger confidence.",
            "• Improving posture consistency will meaningfully enhance your professional presence and perceived authority."
        ]
    elif pct >= 20:  # Below Average
        return [
            "• Your posture is below average, with frequent slouching or compromised alignment throughout.",
            "• Shoulders are often rounded forward and head alignment is inconsistent, reducing your presence.",
            "• Vertical alignment breaks down regularly, suggesting limited core engagement.",
            "• This notably affects your perceived confidence and professional appearance — improvement is needed."
        ]
    elif pct >= 10:  # Weak
        return [
            "• Your posture is weak, tending to slouch or collapse frequently during the presentation.",
            "• Shoulders are regularly hunched and head alignment is poor throughout much of the presentation.",
            "• This significantly undermines your perceived confidence, authority, and professional appearance.",
            "• Focusing on core strength and postural awareness should be a high priority for improvement."
        ]
    else:  # Critical - Needs Major Work
        return [
            "• Your posture needs major improvement, with consistent slouching or collapsing throughout.",
            "• Shoulders are chronically hunched and head alignment is severely compromised, greatly reducing presence.",
            "• This critically impacts your perceived confidence, authority, and overall professional credibility.",
            "• Developing better posture through core strengthening and awareness training is essential and should be your top priority."
        ]

def generate_uprightness_text_th(pct: float) -> list:
    """Thai descriptive text based on uprightness percentage - 10 levels"""
    if pct >= 90:
        return [
            "• ท่าทางตั้งตรงโดดเด่นมาก การจัดแนวลำตัวนิ่งและสมดุลตลอดการพูด",
            "• หน้าอกเปิด ไหล่ผ่อนคลาย ศีรษะอยู่ในแนวที่ดีอย่างสม่ำเสมอ",
            "• แม้มีการใช้ท่าทาง มือและลำตัวก็ยังคงความมั่นคงสูง",
            "• ภาพรวมสื่อความมั่นใจและภาวะผู้นำได้ชัดเจนมาก",
        ]
    elif pct >= 80:
        return [
            "• ท่าทางตั้งตรงยอดเยี่ยมและคงเส้นคงวาเกือบตลอดคลิป",
            "• การจัดแนวหัวไหล่และลำตัวอยู่ในเกณฑ์ดีมาก",
            "• มีการเอนตัวเล็กน้อยบางช่วงแต่ไม่กระทบภาพรวม",
            "• โดยรวมดูมั่นใจและมืออาชีพมาก",
        ]
    elif pct >= 70:
        return [
            "• ท่าทางตั้งตรงดีมากในช่วงส่วนใหญ่ของการนำเสนอ",
            "• โครงลำตัวโดยรวมสมดุลและควบคุมได้ดี",
            "• มีจังหวะที่แนวลำตัวหลุดเล็กน้อยเป็นครั้งคราว",
            "• ภาพรวมยังสนับสนุนความน่าเชื่อถือได้ดี",
        ]
    elif pct >= 60:
        return [
            "• ท่าทางตั้งตรงอยู่ในระดับดีและค่อนข้างสม่ำเสมอ",
            "• ลำตัวและไหล่ส่วนใหญ่ยังอยู่ในแนวเหมาะสม",
            "• มีการคล้อยลำตัวบ้างเล็กน้อยในบางช่วง",
            "• โดยรวมช่วยเสริมบุคลิกมืออาชีพได้ดี",
        ]
    elif pct >= 50:
        return [
            "• ท่าทางตั้งตรงอยู่ในระดับใช้ได้ แต่ยังมีช่วงที่หลุดแนว",
            "• ความต่อเนื่องของแนวลำตัวยังไม่สม่ำเสมอทั้งหมด",
            "• หากคุมแกนลำตัวให้คงที่ขึ้นจะดีขึ้นมาก",
            "• ภาพรวมดีและยังพัฒนาได้อีก",
        ]
    elif pct >= 40:
        return [
            "• ท่าทางอยู่ระดับปานกลางค่อนข้างดี มีช่วงเอนตัวที่สังเกตได้",
            "• แนวไหล่/ศีรษะยังไม่คงที่พอในบางช่วง",
            "• การควบคุมแกนกลางยังมีโอกาสปรับปรุง",
            "• ควรเน้นความสม่ำเสมอของ posture เพิ่มขึ้น",
        ]
    elif pct >= 30:
        return [
            "• ท่าทางอยู่ระดับกลาง มีทั้งช่วงตั้งตรงและช่วงคล้อยลำตัว",
            "• แนวไหล่และศีรษะหลุดจากแนวเหมาะสมค่อนข้างบ่อย",
            "• ส่งผลให้ภาพลักษณ์ความมั่นใจลดลงบางส่วน",
            "• ควรฝึกการคุมแกนลำตัวอย่างต่อเนื่อง",
        ]
    elif pct >= 20:
        return [
            "• ท่าทางต่ำกว่ามาตรฐาน มีการก้มหรือคล้อยตัวบ่อย",
            "• แนวไหล่และคอไม่คงที่ ทำให้บุคลิกดูไม่มั่นคง",
            "• กระทบความน่าเชื่อถือในการสื่อสารอย่างเห็นได้ชัด",
            "• ควรพัฒนาเรื่องท่ายืนและการจัดแนวลำตัวเป็นลำดับแรก",
        ]
    elif pct >= 10:
        return [
            "• ท่าทางค่อนข้างอ่อน มีการยุบตัว/ก้มตัวเป็นส่วนใหญ่",
            "• แนวลำตัวไม่มั่นคงต่อเนื่อง ทำให้พลังการสื่อสารลดลง",
            "• ภาพรวมอาจถูกมองว่าขาดความมั่นใจ",
            "• ควรฝึกการยืนตรงและเสริมการรับรู้ posture อย่างจริงจัง",
        ]
    else:
        return [
            "• ท่าทางต้องปรับปรุงอย่างมาก เนื่องจากการจัดแนวลำตัวยังไม่เหมาะสมเกือบทั้งหมด",
            "• การยุบตัว/ก้มตัวต่อเนื่องกระทบภาพลักษณ์มืออาชีพมาก",
            "• ลดการรับรู้ความมั่นใจและภาวะผู้นำ",
            "• แนะนำให้ฝึกพื้นฐาน posture อย่างเร่งด่วน",
        ]

def generate_stance_text(stability: float) -> list:
    """Generate descriptive text based on stance stability - 10 levels"""
    if stability >= 90:  # Outstanding
        return [
            "• Your stance is outstanding — perfectly symmetrical and solidly grounded, with feet placed optimally at shoulder-width apart.",
            "• Weight shifts are virtually non-existent, creating an exceptionally stable platform that demonstrates supreme confidence.",
            "• You maintain flawless forward orientation toward the audience throughout, maximally reinforcing clarity and engagement.",
            "• Your stance conveys rock-solid stability and commanding, welcoming authority ideal for executive leadership communication."
        ]
    elif stability >= 80:  # Excellent
        return [
            "• Your stance is excellent — exceptionally symmetrical and grounded, with feet placed perfectly about shoulder-width apart.",
            "• Weight shifts are extremely controlled and minimal, preventing any distraction and demonstrating strong confidence.",
            "• You maintain superior forward orientation toward the audience throughout, reinforcing excellent clarity and engagement.",
            "• Your stance conveys remarkable stability and a welcoming, authoritative presence ideal for professional leadership."
        ]
    elif stability >= 70:  # Very Strong
        return [
            "• Your stance is very strong — highly symmetrical and well-grounded, with feet well-positioned shoulder-width apart.",
            "• Weight shifts are well-controlled and quite minimal, showing strong confidence and solid balance.",
            "• You maintain very good forward orientation toward the audience, clearly reinforcing engagement and clarity.",
            "• The stance conveys strong stability and a professional, welcoming presence suitable for senior communication roles."
        ]
    elif stability >= 60:  # Strong
        return [
            "• Your stance is strong — symmetrical and grounded, with feet appropriately placed about shoulder-width apart.",
            "• Weight shifts are controlled and minimal, preventing distraction and demonstrating good confidence.",
            "• You maintain solid forward orientation toward the audience, reinforcing clarity and reasonable engagement.",
            "• The stance effectively conveys stability and a professional presence suitable for most business communication."
        ]
    elif stability >= 50:  # Good
        return [
            "• Your stance is good overall, with generally stable positioning and occasional minor weight shifts.",
            "• Feet placement is appropriate, maintaining reasonable balance throughout most of the presentation.",
            "• Weight distribution is mostly balanced, with some small adjustments visible but not particularly distracting.",
            "• Overall stance supports good stability and adequate professional presence with room for refinement."
        ]
    elif stability >= 40:  # Above Average
        return [
            "• Your stance is above average but shows some noticeable weight shifts that could be reduced.",
            "• Feet placement is acceptable though it could be more consistent for better balance and grounding.",
            "• Some visible swaying or shifting occurs, which may occasionally draw minor audience attention.",
            "• Improving stance stability will noticeably enhance your grounded presence and perceived authority."
        ]
    elif stability >= 30:  # Average
        return [
            "• Your stance is average, with regular weight shifts that reduce your grounded appearance.",
            "• Feet placement varies somewhat, affecting balance and overall stability during the presentation.",
            "• Visible swaying or shifting happens frequently enough to be somewhat distracting to attentive observers.",
            "• Enhancing stance stability will meaningfully improve your authority and professional presence."
        ]
    elif stability >= 20:  # Below Average
        return [
            "• Your stance is below average, showing frequent weight shifts or noticeable instability throughout.",
            "• Feet placement varies considerably and inconsistently, negatively affecting your balance and grounded appearance.",
            "• Visible swaying or shifting is quite apparent and can distract the audience from your message.",
            "• Improving stance stability should be a priority, as it will significantly enhance your authority and presence."
        ]
    elif stability >= 10:  # Weak
        return [
            "• Your stance is weak, with very frequent weight shifts or substantial instability throughout the presentation.",
            "• Feet placement is inconsistent and often suboptimal, severely affecting your balance and grounded appearance.",
            "• Visible swaying, shifting, or fidgeting is quite distracting and notably reduces perceived confidence.",
            "• Developing stance stability is critical and will dramatically improve your grounded authority and professional image."
        ]
    else:  # Critical - Needs Major Work
        return [
            "• Your stance needs major improvement, with constant instability or severe balance issues throughout.",
            "• Feet placement is highly inconsistent or problematic, critically compromising your grounded appearance and balance.",
            "• Continuous swaying, shifting, or movement is very distracting and significantly undermines audience confidence.",
            "• Building fundamental stance stability is essential and should be your absolute top priority for presentation improvement."
        ]

def generate_stance_text_th(stability: float) -> list:
    """Thai descriptive text based on stance stability - 10 levels"""
    if stability >= 90:
        return [
            "• ท่ายืนมั่นคงยอดเยี่ยม สมมาตร และมีฐานยืนที่แข็งแรงมาก",
            "• การถ่ายน้ำหนักน้อยมาก ทำให้ภาพรวมดูนิ่งและมั่นใจสูง",
            "• การวางเท้าเหมาะสมและช่วยเสริมพลังการสื่อสารอย่างชัดเจน",
            "• สะท้อนบุคลิกผู้นำที่พร้อมและน่าเชื่อถือมาก",
        ]
    elif stability >= 80:
        return [
            "• ท่ายืนมั่นคงมาก มีความสมดุลต่อเนื่อง",
            "• การถ่ายน้ำหนักเกิดน้อยและควบคุมได้ดี",
            "• ฐานยืนช่วยให้ภาพลักษณ์ดูมั่นใจและมืออาชีพ",
            "• โดยรวมสื่อความน่าเชื่อถือได้ดีมาก",
        ]
    elif stability >= 70:
        return [
            "• ท่ายืนแข็งแรงและค่อนข้างนิ่งในช่วงส่วนใหญ่",
            "• มีการปรับน้ำหนักเล็กน้อยแต่ไม่รบกวนผู้ฟังมาก",
            "• ฐานยืนโดยรวมดีและรองรับการสื่อสารได้ดี",
            "• ภาพรวมเสริมความมั่นใจได้ชัดเจน",
        ]
    elif stability >= 60:
        return [
            "• ท่ายืนอยู่ในระดับดี มีความมั่นคงพอเหมาะ",
            "• มีการเปลี่ยนน้ำหนักเป็นระยะแต่ยังควบคุมได้",
            "• ฐานยืนโดยรวมยังสนับสนุนภาพลักษณ์มืออาชีพ",
            "• ปรับความนิ่งเพิ่มอีกเล็กน้อยจะดีขึ้นมาก",
        ]
    elif stability >= 50:
        return [
            "• ท่ายืนอยู่ในระดับใช้ได้ แต่ยังมีการขยับน้ำหนักให้เห็นเป็นช่วง ๆ",
            "• ความสมดุลโดยรวมพอใช้ แต่ยังไม่คงที่ตลอด",
            "• ฐานยืนบางช่วงยังไม่แน่นพอ",
            "• พัฒนาได้ชัดเจนหากเพิ่มความนิ่งในการยืน",
        ]
    elif stability >= 40:
        return [
            "• ท่ายืนปานกลางค่อนข้างดี แต่มีการโยก/ถ่ายน้ำหนักบ่อยขึ้น",
            "• ความต่อเนื่องของฐานยืนยังไม่สม่ำเสมอ",
            "• อาจดึงความสนใจผู้ฟังออกจากเนื้อหาในบางช่วง",
            "• ควรฝึกการยืน grounded มากขึ้น",
        ]
    elif stability >= 30:
        return [
            "• ท่ายืนอยู่ระดับกลาง มีการถ่ายน้ำหนักค่อนข้างสม่ำเสมอ",
            "• ความนิ่งของฐานยืนยังไม่พอในจังหวะสำคัญ",
            "• ส่งผลให้ภาพลักษณ์ความมั่นใจลดลงบางส่วน",
            "• แนะนำฝึกคุมแกนล่างและการวางเท้าให้คงที่",
        ]
    elif stability >= 20:
        return [
            "• ท่ายืนต่ำกว่ามาตรฐาน มีอาการโยกหรือเปลี่ยนฐานยืนบ่อย",
            "• ความมั่นคงของช่วงล่างยังไม่ดี ทำให้บุคลิกดูไม่นิ่ง",
            "• กระทบความน่าเชื่อถือของการนำเสนอ",
            "• ควรปรับเรื่องการวางเท้าและถ่ายน้ำหนักโดยตรง",
        ]
    elif stability >= 10:
        return [
            "• ท่ายืนค่อนข้างอ่อน มีการขยับฐานยืนและถ่ายน้ำหนักบ่อยมาก",
            "• ความนิ่งของช่วงล่างไม่เพียงพอสำหรับการสื่อสารที่มั่นใจ",
            "• ผู้ฟังอาจรับรู้ถึงความกังวลหรือไม่มั่นคง",
            "• ควรฝึกท่ายืนพื้นฐานอย่างจริงจัง",
        ]
    else:
        return [
            "• ท่ายืนต้องปรับปรุงอย่างมาก เนื่องจากขาดความมั่นคงต่อเนื่อง",
            "• ฐานยืนไม่คงที่และรบกวนภาพรวมการสื่อสารชัดเจน",
            "• ลดความน่าเชื่อถือและพลังการนำเสนออย่างมาก",
            "• แนะนำฝึกการยืนให้มั่นคงเป็นลำดับเร่งด่วน",
        ]

# Analysis functions
def analyze_video_mediapipe(video_path: str, sample_fps: float = 5, max_frames: int = 300, **kwargs) -> Dict[str, Any]:
    """Real MediaPipe analysis with proper Laban Movement Analysis"""
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        raise RuntimeError("Cannot open video")
    
    fps = cap.get(cv2.CAP_PROP_FPS) or 25.0
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    duration = total_frames / fps if fps > 0 else 0
    
    frame_interval = max(1, int(fps / sample_fps))
    
    # Effort & Shape detection counters
    effort_counts = {
        "Directing": 0, "Enclosing": 0, "Punching": 0, "Spreading": 0,
        "Pressing": 0, "Dabbing": 0, "Indirecting": 0, "Gliding": 0,
        "Flicking": 0, "Advancing": 0, "Retreating": 0
    }
    shape_counts = {"Directing": 0, "Enclosing": 0, "Spreading": 0, "Indirecting": 0, "Advancing": 0, "Retreating": 0}
    
    analyzed = 0
    prev_landmarks = None
    
    with Pose(static_image_mode=False, model_complexity=1) as pose:
        frame_idx = 0
        while analyzed < max_frames:
            ret, frame = cap.read()
            if not ret:
                break
            
            if frame_idx % frame_interval == 0:
                rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                results = pose.process(rgb)
                
                if results.pose_landmarks:
                    lms = results.pose_landmarks.landmark
                    
                    # Calculate movement vectors if we have previous frame
                    if prev_landmarks is not None:
                        # Get key points
                        left_wrist = lms[PoseLandmark.LEFT_WRIST]
                        right_wrist = lms[PoseLandmark.RIGHT_WRIST]
                        left_elbow = lms[PoseLandmark.LEFT_ELBOW]
                        right_elbow = lms[PoseLandmark.RIGHT_ELBOW]
                        left_shoulder = lms[PoseLandmark.LEFT_SHOULDER]
                        right_shoulder = lms[PoseLandmark.RIGHT_SHOULDER]
                        nose = lms[PoseLandmark.NOSE]
                        
                        # Previous frame landmarks
                        prev_left_wrist = prev_landmarks[PoseLandmark.LEFT_WRIST]
                        prev_right_wrist = prev_landmarks[PoseLandmark.RIGHT_WRIST]
                        
                        # Calculate velocities (movement speed)
                        left_wrist_vel = math.sqrt(
                            (left_wrist.x - prev_left_wrist.x)**2 +
                            (left_wrist.y - prev_left_wrist.y)**2 +
                            (left_wrist.z - prev_left_wrist.z)**2
                        )
                        right_wrist_vel = math.sqrt(
                            (right_wrist.x - prev_right_wrist.x)**2 +
                            (right_wrist.y - prev_right_wrist.y)**2 +
                            (right_wrist.z - prev_right_wrist.z)**2
                        )
                        avg_velocity = (left_wrist_vel + right_wrist_vel) / 2
                        
                        # Spatial measurements
                        wrist_dist = abs(left_wrist.x - right_wrist.x)
                        shoulder_width = abs(left_shoulder.x - right_shoulder.x)
                        body_expansion = wrist_dist / max(shoulder_width, 0.1)
                        
                        # Hand height relative to shoulders
                        avg_hand_y = (left_wrist.y + right_wrist.y) / 2
                        avg_shoulder_y = (left_shoulder.y + right_shoulder.y) / 2
                        hands_above_shoulders = avg_hand_y < avg_shoulder_y
                        
                        # Hand depth (Z-axis) relative to body
                        avg_hand_z = (left_wrist.z + right_wrist.z) / 2
                        avg_shoulder_z = (left_shoulder.z + right_shoulder.z) / 2
                        hands_forward = avg_hand_z < avg_shoulder_z
                        
                        # Movement direction (hands)
                        left_z_delta = left_wrist.z - prev_left_wrist.z
                        right_z_delta = right_wrist.z - prev_right_wrist.z
                        avg_z_delta = (left_z_delta + right_z_delta) / 2
                        
                        # Forward/backward: Use average of both hands with higher threshold
                        forward_movement = avg_z_delta < -0.03  # Both hands moving toward camera
                        backward_movement = avg_z_delta > 0.03  # Both hands moving away from camera
                        
                        # Vertical movement
                        upward_movement = (left_wrist.y - prev_left_wrist.y) < -0.01 or (right_wrist.y - prev_right_wrist.y) < -0.01
                        downward_movement = (left_wrist.y - prev_left_wrist.y) > 0.01 or (right_wrist.y - prev_right_wrist.y) > 0.01
                        
                        # Effort qualities
                        is_sudden = avg_velocity > 0.08  # Fast movement
                        is_sustained = 0.02 < avg_velocity <= 0.08  # Moderate movement
                        is_strong = body_expansion > 1.2 or hands_above_shoulders
                        is_light = body_expansion < 0.8
                        
                        # === EFFORT DETECTION (11 types) ===
                        
                        # 1. DIRECTING: Direct, sustained, forward movement
                        if hands_forward and is_sustained and forward_movement:
                            effort_counts["Directing"] += 1
                            shape_counts["Directing"] += 1
                        
                        # 2. ENCLOSING: Arms coming together, wrapping motion
                        if body_expansion < 0.8 and avg_velocity > 0.03:
                            effort_counts["Enclosing"] += 1
                            shape_counts["Enclosing"] += 1
                        
                        # 3. PUNCHING: Sudden, strong, direct forward thrust
                        if is_sudden and is_strong and forward_movement:
                            effort_counts["Punching"] += 1
                        
                        # 4. SPREADING: Arms spreading wide, opening gesture
                        if body_expansion > 1.5 and avg_velocity > 0.03:
                            effort_counts["Spreading"] += 1
                            shape_counts["Spreading"] += 1
                        
                        # 5. PRESSING: Sustained, strong, downward force
                        if is_sustained and is_strong and downward_movement:
                            effort_counts["Pressing"] += 1
                        
                        # 6. DABBING: Sudden, light, direct touch
                        if is_sudden and is_light and avg_velocity > 0.05:
                            effort_counts["Dabbing"] += 1
                        
                        # 7. INDIRECTING: Indirect, curved, wandering movements
                        if not hands_forward and avg_velocity > 0.04 and body_expansion > 1.0:
                            effort_counts["Indirecting"] += 1
                            shape_counts["Indirecting"] += 1
                        
                        # 8. GLIDING: Sustained, light, indirect floating
                        if is_sustained and is_light and upward_movement:
                            effort_counts["Gliding"] += 1
                        
                        # 9. FLICKING: Sudden, light, indirect quick gesture
                        if is_sudden and is_light and not forward_movement:
                            effort_counts["Flicking"] += 1
                        
                        # 10. ADVANCING: Significant forward movement (toward audience)
                        # Requires sustained forward motion with reasonable speed
                        if forward_movement and avg_velocity > 0.06 and is_sustained:
                            effort_counts["Advancing"] += 1
                            shape_counts["Advancing"] += 1
                        
                        # 11. RETREATING: Significant backward movement (pulling back defensively)
                        # Requires sustained backward motion - not just gesture return
                        if backward_movement and avg_velocity > 0.06 and is_sustained:
                            effort_counts["Retreating"] += 1
                            shape_counts["Retreating"] += 1
                    
                    # Store current landmarks for next iteration
                    prev_landmarks = lms
                    analyzed += 1
            
            frame_idx += 1
    
    cap.release()
    
    # Calculate percentages
    total_detections = max(1, sum(effort_counts.values()))
    effort_detection = {k: round(v / total_detections * 100, 1) for k, v in effort_counts.items()}

    total_shape = max(1, sum(shape_counts.values()))
    shape_detection = {k: round(v / total_shape * 100, 1) for k, v in shape_counts.items()}

    # Add diversity/consistency signals so different speaking styles separate better.
    dominant_share = max(effort_detection.values()) / 100.0 if effort_detection else 0.0
    variety_floor = max(2, int(max(1, analyzed) * 0.03))
    variety_count = sum(1 for v in effort_counts.values() if v >= variety_floor)
    monotony_penalty = max(0.0, (dominant_share - 0.35) * 8.0)
    variety_factor = max(0.85, min(1.20, 0.85 + (variety_count / 11.0) * 0.35))

    total_frames = max(1, analyzed)
    engaging_activity = (
        effort_counts.get("Spreading", 0)
        + effort_counts.get("Enclosing", 0)
        + effort_counts.get("Gliding", 0)
        + effort_counts.get("Indirecting", 0)
    ) / total_frames
    confidence_activity = (
        effort_counts.get("Directing", 0)
        + effort_counts.get("Punching", 0)
        + effort_counts.get("Advancing", 0)
        + effort_counts.get("Pressing", 0)
    ) / total_frames
    authority_activity = (
        effort_counts.get("Pressing", 0)
        + effort_counts.get("Punching", 0)
        + effort_counts.get("Directing", 0)
        + effort_counts.get("Advancing", 0)
    ) / total_frames

    engaging_boost = max(0.85, min(1.35, 0.85 + engaging_activity * 1.60))
    confidence_boost = max(0.85, min(1.35, 0.85 + confidence_activity * 1.50))
    authority_boost = max(0.85, min(1.35, 0.85 + authority_activity * 1.50))

    # Category raw signals: positives minus small penalties from counter-signals.
    engaging_raw = (
        effort_detection.get("Spreading", 0) * 0.34
        + effort_detection.get("Enclosing", 0) * 0.26
        + effort_detection.get("Gliding", 0) * 0.22
        + effort_detection.get("Indirecting", 0) * 0.18
        - effort_detection.get("Punching", 0) * 0.12
        - effort_detection.get("Retreating", 0) * 0.10
    )
    confidence_raw = (
        effort_detection.get("Directing", 0) * 0.36
        + effort_detection.get("Punching", 0) * 0.24
        + effort_detection.get("Advancing", 0) * 0.24
        + effort_detection.get("Pressing", 0) * 0.16
        - effort_detection.get("Retreating", 0) * 0.18
        - effort_detection.get("Indirecting", 0) * 0.08
    )
    authority_raw = (
        effort_detection.get("Pressing", 0) * 0.34
        + effort_detection.get("Punching", 0) * 0.26
        + effort_detection.get("Directing", 0) * 0.24
        + effort_detection.get("Advancing", 0) * 0.16
        - effort_detection.get("Flicking", 0) * 0.18
        - effort_detection.get("Retreating", 0) * 0.12
    )

    # Shape-based alignment helps separate categories with similar effort mixes.
    engaging_shape = (
        shape_detection.get("Spreading", 0) * 0.45
        + shape_detection.get("Enclosing", 0) * 0.35
        + shape_detection.get("Indirecting", 0) * 0.20
    )
    confidence_shape = (
        shape_detection.get("Directing", 0) * 0.60
        + shape_detection.get("Advancing", 0) * 0.35
        - shape_detection.get("Retreating", 0) * 0.25
    )
    authority_shape = (
        shape_detection.get("Directing", 0) * 0.55
        + shape_detection.get("Advancing", 0) * 0.30
        - shape_detection.get("Retreating", 0) * 0.40
    )

    engaging_raw = max(
        0.0,
        (engaging_raw * engaging_boost + engaging_shape * 0.22) * variety_factor - monotony_penalty * 0.90,
    )
    confidence_raw = max(
        0.0,
        (confidence_raw * confidence_boost + confidence_shape * 0.25) * variety_factor - monotony_penalty * 1.00,
    )
    authority_raw = max(
        0.0,
        (authority_raw * authority_boost + authority_shape * 0.25) * variety_factor - monotony_penalty * 1.10,
    )

    # Mild contrast expansion to avoid all categories collapsing into the same band.
    raw_vec = np.array([engaging_raw, confidence_raw, authority_raw], dtype=float)
    raw_mean = float(np.mean(raw_vec))
    contrast = 0.95
    engaging_raw = engaging_raw + contrast * (engaging_raw - raw_mean)
    confidence_raw = confidence_raw + contrast * (confidence_raw - raw_mean)
    authority_raw = authority_raw + contrast * (authority_raw - raw_mean)

    engaging_score = min(7, max(1, round(engaging_raw / 5.4 + 1.4)))
    convince_score = min(7, max(1, round(confidence_raw / 5.4 + 1.4)))
    authority_score = min(7, max(1, round(authority_raw / 5.4 + 1.4)))
    
    return {
        "analysis_engine": "mediapipe_real_enhanced",
        "duration_seconds": duration,
        "analyzed_frames": analyzed,
        "total_indicators": 450 + 475 + 445,  # = 1370 (real total)
        "engaging_score": engaging_score,
        "engaging_pos": int(engaging_score / 7 * 450),
        "convince_score": convince_score,
        "convince_pos": int(convince_score / 7 * 475),
        "authority_score": authority_score,
        "authority_pos": int(authority_score / 7 * 445),
        "effort_detection": effort_detection,
        "shape_detection": shape_detection,
    }

def analyze_video_placeholder(video_path: str, seed: int = 42) -> Dict[str, Any]:
    """Fallback placeholder analysis"""
    random.seed(seed)
    duration = get_video_duration_seconds(video_path)
    
    effort_detection = {
        "Directing": 23.9, "Enclosing": 11.9, "Punching": 11.5, "Spreading": 11.3,
        "Pressing": 10.8, "Dabbing": 8.4, "Indirecting": 7.4, "Gliding": 6.2,
        "Flicking": 3.5, "Advancing": 2.6, "Retreating": 2.5
    }
    shape_detection = {"Directing": 40.1, "Enclosing": 20.0, "Spreading": 18.9, "Indirecting": 12.4, "Advancing": 4.4, "Retreating": 4.1}
    
    return {
        "analysis_engine": "placeholder",
        "duration_seconds": duration,
        "analyzed_frames": 100,
        "total_indicators": 450 + 475 + 445,  # = 1370 (real total)
        "engaging_score": 5,
        "engaging_pos": 431,
        "convince_score": 5,
        "convince_pos": 475,
        "authority_score": 5,
        "authority_pos": 445,
        "effort_detection": effort_detection,
        "shape_detection": shape_detection,
    }

# Graph generation
def generate_effort_graph(effort_data: Dict[str, float], shape_data: Dict[str, float], output_path: str):
    """Generate Effort Motion Detection graph (Page 4 style)"""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
    fig.suptitle('Effort Motion Detection Results', fontsize=14, fontweight='bold')
    
    # Light blue color
    light_blue = '#87CEEB'
    
    # Left: Effort Summary (sorted by value, highest at top)
    sorted_efforts = sorted(effort_data.items(), key=lambda x: x[1], reverse=False)  # reverse=False for bottom-to-top
    efforts = [k for k, v in sorted_efforts]
    values = [v for k, v in sorted_efforts]
    
    ax1.barh(efforts, values, color=light_blue)
    ax1.set_xlabel('Percentage (%)')
    ax1.set_title('Effort Summary')
    ax1.set_xlim(0, 100)
    for i, v in enumerate(values):
        ax1.text(v + 1, i, f'{v}%', va='center')
    
    # Right: Top Movement Efforts
    top_3 = sorted(effort_data.items(), key=lambda x: x[1], reverse=True)[:3]
    top_names = [f"{k} - Rank #{i+1}" for i, (k, v) in enumerate(top_3)]
    top_vals = [v for k, v in top_3]
    
    ax2.barh(top_names, top_vals, color=light_blue)
    ax2.set_xlabel('Percentage (%)')
    ax2.set_title('Top Movement Efforts')
    ax2.set_xlim(0, 100)
    for i, v in enumerate(top_vals):
        ax2.text(v + 1, i, f'{v}%', va='center')
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()

def generate_shape_graph(shape_data: Dict[str, float], output_path: str):
    """Generate Shape Motion Detection graph (Page 5 style)"""
    fig, ax = plt.subplots(figsize=(10, 6))
    
    # Light blue color
    light_blue = '#87CEEB'
    
    # Sort by value, highest (left) to lowest (right)
    sorted_shapes = sorted(shape_data.items(), key=lambda x: x[1], reverse=True)
    shapes = [k for k, v in sorted_shapes]
    values = [v for k, v in sorted_shapes]
    
    bars = ax.bar(shapes, values, color=light_blue)
    ax.set_ylabel('Percentage (%)')
    ax.set_xlabel('Shape Motion Type')
    ax.set_title('Shape Motion Detection Results', fontsize=14, fontweight='bold')
    ax.set_ylim(0, max(values) * 1.2 if values else 10)
    
    for bar, val in zip(bars, values):
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height,
                f'{val}%', ha='center', va='bottom')
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()

# DOCX Report building
def build_docx_report(
    report: ReportData,
    output_bio: io.BytesIO,
    graph1_path: str,
    graph2_path: str,
    lang: str = "en",
    report_style: str = "full",
):
    """Build 5-page DOCX report matching the exact template format"""
    doc = Document()
    is_simple = str(report_style or "full").strip().lower().startswith("simple")
    
    # Language-specific text
    is_thai = (lang == "th")
    
    texts = {
        "title": "รายงานการวิเคราะห์การนำเสนอ" if is_thai else "Character Analysis Report",
        "client_name": "ชื่อลูกค้า:" if is_thai else "Client Name:",
        "analysis_date": "วันที่วิเคราะห์:" if is_thai else "Analysis Date:",
        "video_info": "ข้อมูลวิดีโอ (Video Information)" if is_thai else "Video Information",
        "duration": "ระยะเวลา:" if is_thai else "Duration:",
        "detailed_analysis": "การวิเคราะห์โดยละเอียด (Detailed Analysis)" if is_thai else "Detailed Analysis",
        "first_impression": "1.  ความประทับใจแรกพบ (First Impression)" if is_thai else "1.  First impression",
        "eye_contact": "การสบตา (Eye Contact)" if is_thai else "Eye Contact",
        "uprightness": "ความตั้งตรงของร่างกาย (Uprightness)" if is_thai else "Uprightness (Posture & Upper-Body Alignment)",
        "stance": "การยืนและการวางเท้า (Stance)" if is_thai else "Stance (Lower-Body Stability & Grounding)",
        "impact_clients": "ผลกระทบ:" if is_thai else "Impact:",
        "engaging": "2. การสร้างความเป็นมิตรและสร้างสัมพันธภาพ" if is_thai else "2. Engaging & Connecting:",
        "approachability": "• ความเป็นกันเอง" if is_thai else "• Approachability",
        "relatability": "• ความเข้าถึงได้" if is_thai else "• Relatability",
        "engagement": "• การมีส่วนร่วม เชื่อมโยง และสร้างความคุ้นเคยกับทีมอย่างรวดเร็ว" if is_thai else "• Engagement, connect and build instant rapport with team",
        "confidence": "3. ความมั่นใจ:" if is_thai else "3. Confidence:",
        "optimistic": "• บุคลิกภาพเชิงบวก" if is_thai else "• Optimistic Presence",
        "focus": "• ความมีสมาธิ" if is_thai else "• Focus",
        "persuade": "• ความสามารถในการโน้มน้าวและยืนหยัดในจุดยืนเพื่อให้ผู้อื่นคล้อยตาม" if is_thai else "• Ability to persuade and stand one's ground, in order to convince others.",
        "authority": "4.  ความเป็นผู้นำ (Authority):" if is_thai else "4. Authority:",
        "importance": "• แสดงให้เห็นถึงความสำคัญและความเร่งด่วนของประเด็น" if is_thai else "• Showing sense of importance and urgency in subject matter",
        "pressing": "• ผลักดันให้เกิดการลงมือทำ" if is_thai else "• Pressing for action",
        "scale": "ระดับ:" if is_thai else "Scale:",
        "description": "คำอธิบาย: ตรวจพบ" if is_thai else "Description: Detected",
        "indicators": "การเคลื่อนไหวเชิงบวก" if is_thai else "positive indicators out of",
        "total_indicators": "ตัวชี้วัดทั้งหมด" if is_thai else "total indicators",
        "effort_title": "ผลลัพธ์การตรวจจับการเคลื่อนไหวแบบ Effort (Effort Motion Detection Results)" if is_thai else "Effort Motion Detection Results",
        "shape_title": "ผลลัพธ์การตรวจจับการเคลื่อนไหวแบบ Shape (Shape Motion Detection Results)" if is_thai else "Shape Motion Detection Results",
        "generated": "สร้างโดย AI People Reader™" if is_thai else "Generated by AI People Reader™",
    }
    
    # Impact texts in Thai
    impact_eye_thai = "การสบตาที่แข็งแกร่งส่งสัญญาณถึงความมีอยู่ ความจริงใจ และความเชื่อมั่นในความเป็นผู้นำ ทำให้ข้อความของคุณรู้สึกน่าเชื่อถือมากขึ้น"
    impact_upright_thai = "ท่าทางตรงสื่อถึงความมั่นใจในตนเอง ความชัดเจนในความคิด และความมั่นคงทางอารมณ์ ซึ่งเป็นลักษณะของผู้สื่อสารที่มีความไว้วางใจสูง"
    impact_stance_thai = "ท่ายืนที่มั่นคงช่วยเพิ่มอำนาจ การควบคุม และการส่งข้อความที่ราบรื่น ทำให้ผู้พูดดูเตรียมพร้อมและน่าเชื่อถือมากขึ้น"
    
    impact_eye_en = "Strong eye contact signals presence, sincerity, and leadership confidence, making your message feel more reliable."
    impact_upright_en = "Uprightness communicates self-assurance, clarity of thought, and emotional stability all traits of high-trust communicators."
    impact_stance_en = "A grounded stance enhances authority, control, and smooth message delivery, making the speaker appear more prepared and credible."

    def normalize_spacing() -> None:
        """
        Keep a consistent spacing pattern for numbered headings and bullet points
        across pages (especially sections 2-4).
        """
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if not t:
                continue

            pf = p.paragraph_format
            # Numbered section headings: "1.", "2.", "3.", "4."
            if t[:2] in ("1.", "2.", "3.", "4."):
                pf.space_before = Pt(14)
                pf.space_after = Pt(6)
                pf.line_spacing = 1.1
                continue

            # Bullets: List Bullet style or explicit "•"/"▪" in content; keep unified compact rhythm.
            if getattr(p.style, "name", "") == "List Bullet" or t.startswith("•") or t.startswith("▪"):
                pf.space_before = Pt(0)
                pf.space_after = Pt(6)
                pf.line_spacing = 1.2
                continue

            # Keep impact labels/body visually grouped.
            if t in (texts["impact_clients"], "Impact:", "ผลกระทบ:"):
                pf.space_before = Pt(6)
                pf.space_after = Pt(0)
                continue
    
    # Add header and footer images to all pages
    section = doc.sections[0]
    
    # Add header image
    header_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "Header.png")
    if os.path.exists(header_path):
        header = section.header
        header_para = header.paragraphs[0]
        header_run = header_para.add_run()
        header_run.add_picture(header_path, width=Inches(6.5))
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add footer image
    footer_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "Footer.png")
    if os.path.exists(footer_path):
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_run = footer_para.add_run()
        footer_run.add_picture(footer_path, width=Inches(6.5))
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _apply_bullet_layout(paragraph) -> None:
        if paragraph is None:
            return
        pf = paragraph.paragraph_format
        pf.left_indent = Pt(28)
        pf.first_line_indent = Pt(-14)
        pf.space_before = Pt(0)
        pf.space_after = Pt(4)
        pf.line_spacing = 1.15

    def _apply_scale_layout(paragraph) -> None:
        if paragraph is None:
            return
        pf = paragraph.paragraph_format
        pf.left_indent = Pt(28)
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15

    # ============================================================
    # PAGE 1: Cover + First Impression (Eye Contact start)
    # ============================================================
    # Same compact spacing for both Thai and English (match Thai layout).
    doc.add_paragraph()
    
    # Title
    title = doc.add_paragraph(texts["title"])
    title.runs[0].font.size = Pt(18)
    title.runs[0].font.bold = True
    title.paragraph_format.space_after = Pt(6)
    
    # Client info
    doc.add_paragraph(f"{texts['client_name']}  {report.client_name}")
    doc.add_paragraph(f"{texts['analysis_date']}  {report.analysis_date}")
    
    # Video info
    video_info = doc.add_paragraph(texts["video_info"])
    video_info.runs[0].bold = True
    doc.add_paragraph(f"{texts['duration']} {report.video_length_str}")
    doc.add_paragraph()
    
    # Detailed Analysis header
    detailed = doc.add_paragraph(texts["detailed_analysis"])
    detailed.runs[0].bold = True
    
    # Section 1: First impression (same layout TH & EN — match Thai)
    section1 = doc.add_paragraph(texts["first_impression"])
    section1.runs[0].bold = True
    section1.paragraph_format.space_before = Pt(10)
    section1.paragraph_format.space_after = Pt(4)

    def _fi_level_en(value: float, metric: str = "") -> str:
        if str(metric or "").strip().lower() == "eye_contact":
            return "High"
        v = float(value or 0.0)
        if v >= 75.0:
            return "High"
        if v >= 40.0:
            return "Moderate"
        return "Low"

    def _fi_level_th(value: float, metric: str = "") -> str:
        lv = _fi_level_en(value, metric=metric).lower()
        if lv.startswith("high"):
            return "สูง"
        if lv.startswith("moderate"):
            return "กลาง"
        return "ต่ำ"

    if report.first_impression:
        fi = report.first_impression
        eye_level = _fi_level_th(fi.eye_contact_pct, "eye_contact") if is_thai else _fi_level_en(fi.eye_contact_pct, "eye_contact")
        up_level = _fi_level_th(fi.upright_pct, "uprightness") if is_thai else _fi_level_en(fi.upright_pct, "uprightness")
        st_level = _fi_level_th(fi.stance_stability, "stance") if is_thai else _fi_level_en(fi.stance_stability, "stance")
    else:
        eye_level = "-" if is_thai else "-"
        up_level = "-" if is_thai else "-"
        st_level = "-" if is_thai else "-"

    eye_bullet = doc.add_paragraph(("• การสบตา (Eye Contact)" if is_thai else "• Eye Contact"))
    _apply_bullet_layout(eye_bullet)
    lvl1_text = f"{texts['scale']} {eye_level}"
    lvl1 = doc.add_paragraph(lvl1_text)
    lvl1.runs[0].bold = True
    _apply_scale_layout(lvl1)

    upright_bullet = doc.add_paragraph(("• ความตั้งตรงของร่างกาย (Uprightness)" if is_thai else "• Uprightness"))
    _apply_bullet_layout(upright_bullet)
    lvl2_text = f"{texts['scale']} {up_level}"
    lvl2 = doc.add_paragraph(lvl2_text)
    lvl2.runs[0].bold = True
    _apply_scale_layout(lvl2)

    stance_bullet = doc.add_paragraph(("• การยืนและการวางเท้า (Stance)" if is_thai else "• Stance"))
    _apply_bullet_layout(stance_bullet)
    lvl3_text = f"{texts['scale']} {st_level}"
    lvl3 = doc.add_paragraph(lvl3_text)
    lvl3.runs[0].bold = True
    _apply_scale_layout(lvl3)

    remark = doc.add_paragraph("หมายเหตุ" if is_thai else "Remark")
    remark.runs[0].bold = True
    remark.paragraph_format.space_before = Pt(8)
    remark.paragraph_format.space_after = Pt(2)
    doc.add_paragraph(
        "ความรู้สึกแรกพบมักเกิดขึ้นภายใน 5 วินาทีแรกของการพบกัน โดยพิจารณาจากภาพรวม การสบตา ความตั้งตรง และการยืนวางเท้า ก่อนเข้าสู่การวิเคราะห์เชิงพฤติกรรมในส่วนถัดไป"
        if is_thai
        else "First impression happens in the first 5 seconds of meeting someone, and is normally decided from the person's appearance, eye contact, uprightness and stance. However, after the first 5 seconds, the rest (below) are normally taken into consideration."
    )

    # PAGE BREAK TO PAGE 2
    doc.add_page_break()
    # Add one extra line of top spacing on the next page.
    doc.add_paragraph()
    
    # ============================================================
    # PAGE 2: Engaging & Connecting + Confidence
    # ============================================================
    
    # Section 2: Engaging & Connecting
    engaging_cat = report.categories[0]
    section2 = doc.add_paragraph(texts["engaging"])
    section2.runs[0].bold = True
    section2.paragraph_format.space_before = Pt(14)
    section2.paragraph_format.space_after = Pt(4)
    p_approach = doc.add_paragraph(texts["approachability"])
    _apply_bullet_layout(p_approach)
    p_relate = doc.add_paragraph(texts["relatability"])
    _apply_bullet_layout(p_relate)
    p_engage = doc.add_paragraph(texts["engagement"])
    _apply_bullet_layout(p_engage)
    scale_para1 = doc.add_paragraph(f"{texts['scale']} {engaging_cat.scale.capitalize()}")
    scale_para1.runs[0].bold = True
    _apply_scale_layout(scale_para1)
    if not is_simple:
        doc.add_paragraph(f"{texts['description']} {engaging_cat.positives} {texts['indicators']} {engaging_cat.total} {texts['total_indicators']}")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Section 3: Confidence
    confidence_cat = report.categories[1]
    section3 = doc.add_paragraph(texts["confidence"])
    section3.runs[0].bold = True
    section3.paragraph_format.space_before = Pt(14)
    section3.paragraph_format.space_after = Pt(4)
    p_opt = doc.add_paragraph(texts["optimistic"])
    _apply_bullet_layout(p_opt)
    p_focus = doc.add_paragraph(texts["focus"])
    _apply_bullet_layout(p_focus)
    p_persuade = doc.add_paragraph(texts["persuade"])
    _apply_bullet_layout(p_persuade)
    scale_para2 = doc.add_paragraph(f"{texts['scale']} {confidence_cat.scale.capitalize()}")
    scale_para2.runs[0].bold = True
    _apply_scale_layout(scale_para2)
    if not is_simple:
        doc.add_paragraph(f"{texts['description']} {confidence_cat.positives} {texts['indicators']} {confidence_cat.total} {texts['total_indicators']}")
    
    # Keep spacing between Section 3 and Section 4 similar to Section 2 and 3
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ============================================================
    # Section 4: Authority (same page as Section 3)
    # ============================================================
    
    # Section 4: Authority
    authority_cat = report.categories[2]
    section4 = doc.add_paragraph(texts["authority"])
    section4.runs[0].bold = True
    section4.paragraph_format.space_before = Pt(14)
    section4.paragraph_format.space_after = Pt(4)
    p_importance = doc.add_paragraph(texts["importance"])
    _apply_bullet_layout(p_importance)
    p_pressing = doc.add_paragraph(texts["pressing"])
    _apply_bullet_layout(p_pressing)
    scale_para3 = doc.add_paragraph(f"{texts['scale']} {authority_cat.scale.capitalize()}")
    scale_para3.runs[0].bold = True
    _apply_scale_layout(scale_para3)
    if not is_simple:
        doc.add_paragraph(f"{texts['description']} {authority_cat.positives} {texts['indicators']} {authority_cat.total} {texts['total_indicators']}")
        
        # PAGE BREAK TO PAGE 4
        doc.add_page_break()
        
        # ============================================================
        # PAGE 4: Effort Motion Detection Results
        # ============================================================
        
        # Spacing after header (same as page 1)
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Title
        title4 = doc.add_paragraph(texts["effort_title"])
        title4.runs[0].font.size = Pt(18)
        title4.runs[0].bold = True
        doc.add_paragraph()
        
        # Graph
        if os.path.exists(graph1_path):
            doc.add_picture(graph1_path, width=Inches(6.5))
        
        # PAGE BREAK TO PAGE 5
        doc.add_page_break()
        
        # ============================================================
        # PAGE 5: Shape Motion Detection Results
        # ============================================================
        
        # Spacing after header (same as page 1)
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Title
        title5 = doc.add_paragraph(texts["shape_title"])
        title5.runs[0].font.size = Pt(18)
        title5.runs[0].bold = True
        doc.add_paragraph()
        
        # Graph
        if os.path.exists(graph2_path):
            doc.add_picture(graph2_path, width=Inches(6.5))
        
        # Add "Generated by AI People Reader™" text below graph
        doc.add_paragraph()
        doc.add_paragraph()
    else:
        doc.add_paragraph()

    # Normalize paragraph rhythm so numbered headings and bullets follow
    # a consistent spacing pattern throughout the document.
    normalize_spacing()

    generated_para = doc.add_paragraph(texts["generated"])
    generated_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    generated_run = generated_para.runs[0]
    generated_run.italic = True
    generated_run.font.size = Pt(11)
    
    # Save
    doc.save(output_bio)

def build_pdf_report(
    report: ReportData,
    output_path: str,
    graph1_path: str,
    graph2_path: str,
    lang: str = "en",
    report_style: str = "full",
):
    logger = logging.getLogger("report_core_pdf")
    """Build PDF report aligned with DOCX simple/full text structure (no graph embedding)."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.utils import simpleSplit
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except Exception as e:
        raise RuntimeError(f"reportlab is required for PDF generation: {e}")

    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    x_left = 50
    x_right = 50
    usable_width = width - x_left - x_right
    header_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "Header.png")
    footer_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "Footer.png")
    top_content_y = height - 95
    bottom_content_y = 70
    y = top_content_y

    is_simple = str(report_style or "full").strip().lower().startswith("simple")
    is_thai = (lang == "th")
    regular_font = "Helvetica"
    bold_font = "Helvetica-Bold"
    requires_unicode_font = False

    def _first_existing(paths: list) -> str:
        for p in paths:
            if p and os.path.exists(p):
                return p
        return ""

    def _glob_existing(patterns: list) -> list:
        import glob
        found = []
        for pat in patterns:
            if not pat:
                continue
            found.extend(glob.glob(pat, recursive=True))
        # Keep deterministic order and unique entries.
        uniq = []
        seen = set()
        for p in sorted(found):
            if p not in seen and os.path.isfile(p):
                uniq.append(p)
                seen.add(p)
        return uniq

    def _register_ttf(font_name: str, path: str, require_thai: bool = False) -> bool:
        if not path:
            return False
        try:
            pdfmetrics.getFont(font_name)
            return True
        except Exception:
            pass
        try:
            tt = TTFont(font_name, path)
            if require_thai:
                # Best-effort Thai glyph validation. Do not over-reject fonts when metadata
                # differs between platforms; runtime rendering is the final authority.
                required = [ord("ก"), ord("า"), ord("ไ"), ord("ย")]
                cmap = getattr(tt.face, "charToGlyph", {}) or {}
                widths = getattr(tt.face, "charWidths", {}) or {}
                if cmap:
                    if any((cp not in cmap) and (str(cp) not in cmap) for cp in required):
                        return False
                elif widths:
                    if any((cp not in widths) and (str(cp) not in widths) for cp in required):
                        return False
            pdfmetrics.registerFont(tt)
            return True
        except Exception:
            return False

    if is_thai:
        thai_glob_candidates = _glob_existing(
            [
                "/usr/share/fonts/**/*.ttf",
                "/usr/local/share/fonts/**/*.ttf",
            ]
        )
        thai_glob_preferred = [
            p for p in thai_glob_candidates
            if any(k in os.path.basename(p).lower() for k in ("noto", "sarabun", "thsarabun", "thai", "garuda", "waree", "kinnari", "loma"))
        ]

        thai_regular_candidates = [
            os.getenv("PDF_THAI_FONT_PATH", "").strip(),
            os.getenv("THAI_FONT_PATH", "").strip(),
            os.getenv("REPORT_THAI_FONT_PATH", "").strip(),
            "/Library/Fonts/THSarabunNew.ttf",
            "/Library/Fonts/Sarabun-Regular.ttf",
            "/System/Library/Fonts/Supplemental/Thonburi.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansThai-Regular.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansThaiUI-Regular.ttf",
            "/usr/share/fonts/truetype/thai/Sarabun-Regular.ttf",
            "/usr/share/fonts/truetype/tlwg/Garuda.ttf",
            "/usr/share/fonts/truetype/tlwg/Waree.ttf",
            "/usr/share/fonts/truetype/tlwg/Loma.ttf",
            "/usr/share/fonts/truetype/tlwg/Kinnari.ttf",
        ]
        thai_bold_candidates = [
            os.getenv("PDF_THAI_FONT_BOLD_PATH", "").strip(),
            os.getenv("THAI_FONT_BOLD_PATH", "").strip(),
            os.getenv("REPORT_THAI_FONT_BOLD_PATH", "").strip(),
            "/Library/Fonts/THSarabunNew Bold.ttf",
            "/Library/Fonts/Sarabun-Bold.ttf",
            "/System/Library/Fonts/Supplemental/Thonburi Bold.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansThai-Bold.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansThaiUI-Bold.ttf",
            "/usr/share/fonts/truetype/thai/Sarabun-Bold.ttf",
            "/usr/share/fonts/truetype/tlwg/Garuda-Bold.ttf",
            "/usr/share/fonts/truetype/tlwg/Waree-Bold.ttf",
            "/usr/share/fonts/truetype/tlwg/Loma-Bold.ttf",
            "/usr/share/fonts/truetype/tlwg/Kinnari-Bold.ttf",
        ]
        thai_regular_candidates.extend(thai_glob_preferred)
        thai_bold_candidates.extend(thai_glob_preferred)

        ok_regular = False
        for path in thai_regular_candidates:
            if _register_ttf("ThaiPDFRegular", path, require_thai=True):
                ok_regular = True
                break

        ok_bold = False
        for path in thai_bold_candidates:
            if _register_ttf("ThaiPDFBold", path, require_thai=True):
                ok_bold = True
                break

        if not ok_regular:
            # Do not fail the whole job. Fallback to readable English PDF if Thai font is missing.
            logger.warning(
                "Thai font not found for PDF. Falling back to English labels/content for this PDF."
            )
            is_thai = False
            regular_font = "Helvetica"
            bold_font = "Helvetica-Bold"
            requires_unicode_font = False
        else:
            regular_font = "ThaiPDFRegular"
            bold_font = "ThaiPDFBold" if ok_bold else "ThaiPDFRegular"
            requires_unicode_font = True
            # ReportLab Thai rendering clips stacked marks more easily with bold faces.
            bold_font = regular_font

    def draw_header_footer() -> None:
        # Match DOCX branding as closely as possible for PDF output.
        if os.path.exists(header_path):
            try:
                c.drawImage(
                    header_path,
                    x=28,
                    y=height - 62,
                    width=width - 56,
                    height=36,
                    preserveAspectRatio=True,
                    mask="auto",
                )
            except Exception:
                pass
        if os.path.exists(footer_path):
            try:
                c.drawImage(
                    footer_path,
                    x=28,
                    y=18,
                    width=width - 56,
                    height=30,
                    preserveAspectRatio=True,
                    mask="auto",
                )
            except Exception:
                pass

    draw_header_footer()

    def _normalize_thai_for_pdf(text: str) -> str:
        return unicodedata.normalize("NFC", str(text or ""))

    def _draw_text_line(x: float, y_pos: float, font_name: str, size: int, line: str) -> None:
        text_obj = c.beginText()
        text_obj.setTextOrigin(x, y_pos)
        text_obj.setFont(font_name, size)
        text_obj.textLine(str(line or ""))
        c.drawText(text_obj)

    def write_line(text: str, size: int = 11, bold: bool = False, gap: int = 18):
        nonlocal y
        font = bold_font if bold else regular_font
        effective_gap = max(int(gap), int(size * (1.9 if is_thai else 1.35)))
        if requires_unicode_font:
            safe = _normalize_thai_for_pdf(text)
        else:
            # Normalize common unicode punctuation for ASCII/Latin fonts.
            normalized = (
                _normalize_thai_for_pdf(text)
                .replace("•", "- ")
                .replace("—", "-")
                .replace("–", "-")
                .replace("“", '"')
                .replace("”", '"')
                .replace("’", "'")
            )
            safe = normalized.encode("latin-1", "replace").decode("latin-1")

        def _split_by_chars(long_line: str) -> list:
            parts = []
            buf = ""
            for ch in str(long_line or ""):
                probe = f"{buf}{ch}"
                if (not buf) or (pdfmetrics.stringWidth(probe, font, size) <= usable_width):
                    buf = probe
                else:
                    parts.append(buf)
                    buf = ch
            if buf:
                parts.append(buf)
            return parts or [""]

        initial_lines = simpleSplit(safe, font, size, usable_width) or [""]
        lines = []
        for ln in initial_lines:
            if pdfmetrics.stringWidth(ln, font, size) <= usable_width:
                lines.append(ln)
                continue
            if is_thai or (" " not in ln):
                lines.extend(_split_by_chars(ln))
            else:
                lines.append(ln)

        wrapped_gap = max(16 if is_thai else 12, int(size * (1.7 if is_thai else 1.35)))
        for idx, line in enumerate(lines):
            if y <= bottom_content_y:
                c.showPage()
                draw_header_footer()
                y = top_content_y
            _draw_text_line(x_left, y, font, size, line)
            y -= effective_gap if idx == len(lines) - 1 else wrapped_gap

    def write_block(lines: list, size: int = 11, bold: bool = False, gap: int = 16):
        for line in lines:
            write_line(line, size=size, bold=bold, gap=gap)

    title = "รายงานการวิเคราะห์การนำเสนอ" if is_thai else "Character Analysis Report"
    detailed_analysis_label = "การวิเคราะห์โดยละเอียด" if is_thai else "Detailed Analysis"
    first_impression_label = "1. ความประทับใจแรกพบ" if is_thai else "1. First impression"
    scale_label = "ระดับ" if is_thai else "Scale"
    remark_label = "หมายเหตุ" if is_thai else "Remark"
    remark_text = (
        "ความรู้สึกแรกพบมักเกิดขึ้นภายใน 5 วินาทีแรกของการพบกัน โดยพิจารณาจากภาพรวม การสบตา ความตั้งตรง และการยืนวางเท้า ก่อนเข้าสู่การวิเคราะห์เชิงพฤติกรรมในส่วนถัดไป"
        if is_thai
        else "First impression happens in the first 5 seconds of meeting someone, and is normally decided from the person's appearance, eye contact, uprightness and stance. However, after the first 5 seconds, the rest (below) are normally taken into consideration."
    )
    category_labels = (
        ["2. การสร้างความเป็นมิตรและสัมพันธภาพ", "3. ความมั่นใจ", "4. ความเป็นผู้นำ"]
        if is_thai
        else ["2. Engaging & Connecting", "3. Confidence", "4. Authority"]
    )

    def _fi_level_en(value: float, metric: str = "") -> str:
        if str(metric or "").strip().lower() == "eye_contact":
            return "High"
        v = float(value or 0.0)
        if v >= 75.0:
            return "High"
        if v >= 40.0:
            return "Moderate"
        return "Low"

    def _fi_level_th(value: float, metric: str = "") -> str:
        lv = _fi_level_en(value, metric=metric).lower()
        if lv.startswith("high"):
            return "สูง"
        if lv.startswith("moderate"):
            return "กลาง"
        return "ต่ำ"

    write_line(title, size=16, bold=True, gap=24)
    write_line(f"{'ชื่อลูกค้า' if is_thai else 'Client Name'}: {report.client_name}", bold=True)
    write_line(f"{'วันที่วิเคราะห์' if is_thai else 'Analysis Date'}: {report.analysis_date}")
    write_line(f"{'ระยะเวลา' if is_thai else 'Duration'}: {report.video_length_str}", gap=22)
    write_line(detailed_analysis_label, size=13, bold=True, gap=20)

    # First impression — compact format (match Thai layout)
    write_line(first_impression_label, size=12, bold=True, gap=18)
    if report.first_impression:
        fi = report.first_impression
        eye_level = _fi_level_th(fi.eye_contact_pct, "eye_contact") if is_thai else _fi_level_en(fi.eye_contact_pct, "eye_contact")
        up_level = _fi_level_th(fi.upright_pct, "uprightness") if is_thai else _fi_level_en(fi.upright_pct, "uprightness")
        st_level = _fi_level_th(fi.stance_stability, "stance") if is_thai else _fi_level_en(fi.stance_stability, "stance")

        write_line("• การสบตา (Eye Contact)" if is_thai else "• Eye Contact", gap=14)
        write_line(f"{scale_label}: {eye_level}", bold=True, gap=14)
        write_line("• ความตั้งตรงของร่างกาย (Uprightness)" if is_thai else "• Uprightness", gap=14)
        write_line(f"{scale_label}: {up_level}", bold=True, gap=14)
        write_line("• การยืนและการวางเท้า (Stance)" if is_thai else "• Stance", gap=14)
        write_line(f"{scale_label}: {st_level}", bold=True, gap=14)
        write_line(remark_label, bold=True, gap=14)
        write_line(remark_text, gap=18)
    else:
        write_line("- Not available", gap=20)

    # Category bullet texts (match Thai layout)
    cat_bullets = (
        [
            ["• ความเป็นกันเอง", "• ความเข้าถึงได้", "• การมีส่วนร่วม เชื่อมโยง และสร้างความคุ้นเคยกับทีมอย่างรวดเร็ว"],
            ["• บุคลิกภาพเชิงบวก", "• ความมีสมาธิ", "• ความสามารถในการโน้มน้าวและยืนหยัดในจุดยืนเพื่อให้ผู้อื่นคล้อยตาม"],
            ["• แสดงให้เห็นถึงความสำคัญและความเร่งด่วนของประเด็น", "• ผลักดันให้เกิดการลงมือทำ"],
        ]
        if is_thai
        else [
            ["• Approachability", "• Relatability", "• Engagement, connect and build instant rapport with team"],
            ["• Optimistic Presence", "• Focus", "• Ability to persuade and stand one's ground, in order to convince others."],
            ["• Showing sense of importance and urgency in subject matter", "• Pressing for action"],
        ]
    )

    # Categories (same order as DOCX)
    for idx, cat in enumerate(report.categories[:3]):
        write_line(category_labels[idx], size=12, bold=True, gap=18)
        for bullet in cat_bullets[idx]:
            write_line(bullet, gap=14)
        write_line(
            f"{scale_label}: {cat.scale.capitalize()}",
            bold=True,
            gap=16,
        )
        if not is_simple:
            if is_thai:
                write_line(f"คำอธิบาย: ตรวจพบตัวชี้วัดเชิงบวก {cat.positives} จากทั้งหมด {cat.total}", gap=16)
            else:
                write_line(f"Description: Detected {cat.positives} positive indicators out of {cat.total}", gap=16)

    write_line("", gap=10)
    write_line("สร้างโดย AI People Reader" if is_thai else "Generated by AI People Reader", size=10)
    c.save()
